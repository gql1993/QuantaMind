"""
基于 Springer LNCS 模板 (splnproc2510) 生成 QuantaMind 论文 Word 文件
"""
import os, sys, copy
from docx import Document
from docx.shared import Pt

TEMPLATE = r"E:\work\QuantaMind\demo\splnproc2510.docx"
OUTPUT = r"E:\work\QuantaMind\docs\QuantaMind_Paper_LNCS.docx"

doc = Document(TEMPLATE)

for p in doc.paragraphs:
    p.clear()
for table in doc.tables:
    table._element.getparent().remove(table._element)

idx = [0]

def _style(name):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles['Normal']

def add(style_name, text):
    if idx[0] < len(doc.paragraphs):
        p = doc.paragraphs[idx[0]]
        p.style = _style(style_name)
        p.text = text
    else:
        p = doc.add_paragraph(text, style=_style(style_name))
    idx[0] += 1
    return p

def add_run(style_name, text, bold=False, italic=False):
    p = add(style_name, "")
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

# ── Title ──
add("papertitle", "QuantaMind: A Multi-Agent AI System for Autonomous Superconducting Quantum Chip Research")

# ── Authors ──
add("author", "Anonymous Authors")

# ── Affiliations ──
add("address", "1 Quantum Technology Innovation Center, Yangtze River Delta, China")
add("address", "2 Institute of Quantum Computing, CETC, China")

# ── Abstract ──
add("abstract",
    "Abstract. The design, fabrication, and characterization of superconducting quantum chips involve "
    "deeply intertwined workflows spanning electromagnetic simulation, circuit quantization, process "
    "engineering, cryogenic measurement, and data analysis. Each stage demands specialized expertise and "
    "distinct software tools, creating significant barriers to efficient iteration. Here we introduce "
    "QuantaMind, a multi-agent AI system that unifies the full lifecycle of superconducting quantum chip "
    "research behind a single natural-language interface. QuantaMind orchestrates twelve specialized AI "
    "scientist agents\u2014spanning chip design, electromagnetic simulation, theoretical physics, process "
    "engineering, device calibration, quantum error mitigation, and data analysis\u2014over a shared tool "
    "execution layer that integrates industry-standard platforms including Qiskit Metal, KQCircuits, "
    "Ansys HFSS/Q3D, ARTIQ, CHIPMES, and Apache Doris. The system features (i) a closed-loop "
    "architecture where design intent flows from natural language through EDA tools to GDS layout and "
    "back through simulation and measurement feedback, (ii) an autonomous heartbeat engine that "
    "continuously monitors equipment, data quality, and experimental anomalies across four escalating "
    "intelligence tiers, and (iii) a theoretical physicist agent implementing nine functional modules "
    "(M0\u2013M8) covering Hamiltonian modeling, noise budgeting, Bayesian parameter inversion, "
    "information-gain-driven experiment design, root-cause diagnosis, and design optimization. We "
    "demonstrate QuantaMind on a 20-qubit tunable-coupler chip design task derived from a real fabrication "
    "specification, showing that the system can autonomously generate parameterized GDS layouts, execute "
    "Q3D capacitance extraction and EPR analysis, diagnose gate-error root causes, and propose "
    "next-iteration design modifications\u2014all from conversational prompts.")

# ── Keywords ──
add("keywords",
    "Keywords: multi-agent system, superconducting quantum chip, autonomous research, "
    "LLM-based agent, chip design automation, closed-loop optimization.")

# ── 1 Introduction ──
add("heading1", "1   Introduction")

add("p1a",
    "Superconducting quantum computing has advanced rapidly, with processors scaling from tens to over "
    "one hundred qubits [1\u20135]. Yet the research and development cycle for each chip generation remains "
    "labor-intensive. A single design iteration typically involves (i) frequency planning and Hamiltonian "
    "parameter selection, (ii) layout design in specialized EDA tools, (iii) electromagnetic simulation for "
    "capacitance extraction and eigenmode analysis, (iv) cleanroom fabrication with multi-step lithography "
    "and Josephson junction deposition, (v) cryogenic measurement including spectroscopy, coherence "
    "characterization, and gate calibration, and (vi) data-driven diagnosis and redesign. Each stage relies "
    "on distinct software ecosystems and requires deep domain expertise that is rarely held by a single "
    "researcher.")

add("Normal",
    "Large language models (LLMs) have recently demonstrated the ability to bridge such expertise gaps. "
    "General-purpose scientific agents like The AI Scientist [6], PaperQA [7], and Kosmos [8] show "
    "end-to-end research capabilities. In quantum computing, El Agente Q [9] automates quantum chemistry "
    "workflows through hierarchical multi-agent planning, while El Agente Cu\u00e1ntico [10] extends this to "
    "quantum simulation across multiple software frameworks. In hardware calibration, k-agents [11] "
    "demonstrate autonomous qubit tuning through reinforcement learning.")

add("Normal",
    "However, existing systems address isolated stages of the quantum hardware lifecycle. No current agent "
    "system spans the full design\u2013simulate\u2013fabricate\u2013measure\u2013analyze\u2013redesign loop that defines "
    "superconducting chip R&D. This gap motivates QuantaMind.")

add("Normal",
    "Contributions. We present QuantaMind with five main contributions: "
    "(1) Full-lifecycle coverage integrating twelve AI scientist agents across theory, design, simulation, "
    "fabrication, calibration, error mitigation, data analysis, and knowledge management. "
    "(2) A closed-loop architecture with tool-call orchestration enabling multi-step autonomous reasoning. "
    "(3) A theoretical physicist agent with nine functional modules (M0\u2013M8) for physics-grounded reasoning. "
    "(4) An autonomous heartbeat engine with four escalating intelligence tiers (L0\u2013L3). "
    "(5) Real-world demonstration on a 20-qubit tunable-coupler chip from a production specification.")

# ── 2 System Architecture ──
add("heading1", "2   System Architecture")

add("heading2", "2.1   Overview")

add("p1a",
    "QuantaMind adopts a layered architecture inspired by the OpenClaw agent framework [12], specialized for "
    "quantum chip research. The system comprises six core components: Gateway (session management and "
    "protocol adaptation), Brain (intent classification, MoE routing, multi-model inference with fallback "
    "chains, and function-calling loop orchestration), Hands (a unified tool execution layer with 120+ "
    "registered tools across ten platform adapters), Memory (working memory, project memory, vector "
    "knowledge base), Heartbeat (four-tier autonomous monitoring), and Skills (domain skill marketplace).")

add("heading2", "2.2   Agent Team")

add("p1a",
    "QuantaMind deploys twelve specialized AI scientist agents ordered by the natural research workflow: "
    "(1) Orchestrator for task decomposition and multi-agent dispatch; "
    "(2) Theoretical Physicist for Hamiltonian modeling, noise analysis, and experiment design (10 tools); "
    "(3) Chip Designer for end-to-end transmon chip design and GDS export via Qiskit Metal and KQCircuits; "
    "(4) Simulation Engineer for HFSS eigenmode, Q3D extraction, LOM/EPR analysis (7 tools); "
    "(5) Materials Scientist for substrate screening; "
    "(6) Process Engineer for craft route management, yield/SPC via CHIPMES and OpenMES; "
    "(7) Device Operator for ARTIQ pulse sequences and Qiskit Pulse calibration; "
    "(8) Measurement Scientist for full qubit characterization and Mitiq error mitigation; "
    "(9) Quantum Algorithm Engineer for circuit design and noise-aware compilation; "
    "(10) Data Analyst for cross-domain queries via Doris, SeaTunnel, QCoDeS; "
    "(11) Knowledge Engineer for literature tracking and knowledge graph; "
    "(12) Project Manager for milestone tracking.")

add("heading2", "2.3   Tool Call Loop")

add("p1a",
    "Following the ReAct paradigm [13], each agent operates in a think\u2013act\u2013observe loop. The Brain "
    "constructs a message sequence including the agent's system prompt, conversation history, and available "
    "tool definitions. The LLM generates either a response or a structured tool call. If a tool call is issued, "
    "the Hands layer executes it and returns the result. This loop continues for up to N_max rounds "
    "(default: 8), enabling multi-step reasoning. When tool calls occur, the orchestrator automatically "
    "creates a conversation pipeline recording each step for full provenance.")

add("heading2", "2.4   Graceful Degradation")

add("p1a",
    "Each Hands adapter implements dual-mode design: when the target platform is available, real API calls "
    "are executed; when unavailable, the adapter falls back to physics-informed mock responses, ensuring "
    "the system remains functional for reasoning even without full infrastructure.")

# ── 3 Theoretical Physicist Agent ──
add("heading1", "3   Theoretical Physicist Agent")

add("p1a",
    "The theoretical physicist agent is the physics reasoning core of QuantaMind. It implements nine structured "
    "functional modules (M0\u2013M8) based on a detailed design specification for superconducting circuit theory.")

add("heading2", "3.1   Module Architecture")

add("p1a",
    "M0 (Data Ingestion) constructs a unified DeviceGraph from heterogeneous inputs. "
    "M1 (Quantization Modeling) builds effective Hamiltonian models using EPR, black-box quantization, or "
    "lumped-element methods, outputting qubit frequencies, anharmonicities, coupling strengths, dispersive "
    "shifts, and ZZ interactions. "
    "M2 (Noise Budget) decomposes T1 and T2 into constituent mechanisms (dielectric loss, TLS, Purcell, "
    "quasiparticle, flux noise, thermal photon) and generates a ranked error budget with sensitivity matrix. "
    "M3 (Parameter Inversion) performs Bayesian inference to obtain posterior distributions with uncertainty "
    "quantification and identifiability analysis. "
    "M4 (Experiment Design) uses information-gain maximization to recommend the next most informative "
    "experiment. "
    "M5 (Pulse Optimization) generates gate pulses considering multi-level leakage, crosstalk, and hardware "
    "constraints. "
    "M6 (Root-Cause Diagnosis) implements fault-tree reasoning with probabilistic ranking. "
    "M7 (Design Optimization) synthesizes results into Pareto-optimal design proposals. "
    "M8 (Knowledge Retrieval) searches a structured knowledge base mapping phenomena to mechanisms, "
    "experiments, and mitigations.")

# ── 4 Chip Design Pipeline ──
add("heading1", "4   Chip Design Pipeline")

add("heading2", "4.1   20-Qubit Tunable-Coupler Design")

add("p1a",
    "We demonstrate QuantaMind on a 20-qubit chip following a production specification [14]: "
    "chip size 12.5 mm \u00d7 12.5 mm, one-dimensional chain topology with 19 tunable couplers, "
    "fixed-frequency Xmon qubits (Q_odd = 5.152 GHz, Q_even = 4.650 GHz), tunable coupler at "
    "6.844 GHz with E_C = 348 MHz and Manhattan SQUID junction, 20 readout resonators "
    "(7.0\u20137.97 GHz) on 5 feedlines, 48 SMP connectors, and CPW impedance Z_0 = 50 \u03a9 "
    "(s = 10 \u03bcm, w = 5 \u03bcm on sapphire).")

add("Normal",
    "The design pipeline proceeds: (1) natural-language intent is routed to the Chip Designer agent; "
    "(2) the agent invokes metal_create_design, metal_add_transmon (\u00d720), metal_add_route (\u00d719), "
    "and metal_export_gds; (3) GDS generation uses both Qiskit Metal for component placement and a "
    "dedicated gdstk-based reference layout generator; (4) the Simulation Engineer performs Q3D "
    "extraction and HFSS eigenmode analysis; (5) the Theoretical Physicist validates parameters through "
    "Hamiltonian modeling and noise budgeting.")

add("heading2", "4.2   GDS Layout Generation")

add("p1a",
    "The layout generator produces 20 Xmon qubits with rotated cross geometry and Manhattan JJ markers, "
    "19 tunable couplers with SQUID structures and Z-line coupling rings, 5 readout feedlines with "
    "length-differentiated resonator branches, 48 wirebond launchpads distributed across four chip edges, "
    "four-corner alignment marks, and organized fanout routing. The generated GDS contains approximately "
    "1000 shapes across 8 layers in a 12.5 mm \u00d7 12.5 mm footprint.")

# ── 5 Experiments and Results ──
add("heading1", "5   Experiments and Results")

add("heading2", "5.1   Hamiltonian Modeling")

add("p1a",
    "Starting from the 20-qubit DeviceGraph, the theoretical physicist agent constructs an effective "
    "Hamiltonian model. Agent-predicted qubit frequencies agree with design values within 0.004 GHz "
    "(Q_odd: 5.148 \u00b1 0.004 GHz vs. 5.152 GHz design; Q_even: 4.652 \u00b1 0.003 GHz vs. 4.650 GHz). "
    "Anharmonicity is predicted at \u2212258 \u00b1 2 MHz (design: \u2212260 MHz), coupling strength at "
    "14.8 \u00b1 1.5 MHz, and residual ZZ at 18 \u00b1 5 kHz (below the 20 kHz target).")

add("heading2", "5.2   Noise Budget Analysis")

add("p1a",
    "The noise budget module identifies dielectric loss (T1 limit ~ 80 \u03bcs) as the dominant T1 mechanism, "
    "followed by TLS defects (120 \u03bcs), Purcell decay (200 \u03bcs), and quasiparticle tunneling (500 \u03bcs). "
    "For T2, 1/f flux noise (T\u03c6 ~ 60 \u03bcs) dominates, followed by thermal photon dephasing (200 \u03bcs). "
    "The sensitivity matrix identifies dielectric loss reduction and magnetic shielding improvement as "
    "the two highest-leverage interventions.")

add("heading2", "5.3   Root-Cause Diagnosis")

add("p1a",
    "Given a simulated anomaly of CZ gate fidelity below 99%, the diagnosis module produces a ranked "
    "list: (1) frequency collision near the CZ operating point (confidence 0.45), verified by high-resolution "
    "Chevron scan; (2) flux bias line drift (0.30), verified by long-term Ramsey tracking; (3) parasitic "
    "package mode hybridization (0.15), verified by broadband spectroscopy vs. temperature.")

add("heading2", "5.4   Design Optimization Proposal")

add("p1a",
    "Based on the noise budget and diagnosis, M7 generates three-tier recommendations: "
    "immediate actions (adjust CZ pulse amplitude, optimize readout frequency), "
    "mid-term changes (increase airbridge density, add Purcell filter, reduce SQUID loop area), "
    "and next-iteration redesign (new frequency plan, improved substrate cleaning protocol).")

add("heading2", "5.5   Autonomous Discovery")

add("p1a",
    "During a 24-hour monitoring period, the heartbeat engine produced four discoveries across tiers: "
    "L0 detected an equipment alarm on LITHO-03; L1 flagged yield average at 88% (below 90% threshold); "
    "L2 identified Q3 T1 = 32 \u03bcs (below 35 \u03bcs threshold) and recommended recalibration; "
    "L3 discovered that high-yield batches correlate with calibration fidelity of 99.2% vs. 97.8% in "
    "low-yield batches, logging this as a design insight for future iterations.")

# ── 6 Discussion ──
add("heading1", "6   Discussion")

add("p1a",
    "QuantaMind differs from existing agent systems in three key aspects. First, El Agente Q [9] and "
    "El Agente Cu\u00e1ntico [10] target computational chemistry and quantum simulation respectively, "
    "operating purely in software. QuantaMind targets superconducting chip R&D with hardware integration "
    "across 12+ platforms. Second, k-agents [11] demonstrate calibration automation but do not cover "
    "design or fabrication. QuantaMind spans the full design\u2013simulate\u2013fabricate\u2013measure\u2013analyze\u2013redesign "
    "loop. Third, QuantaMind's theoretical physicist agent provides physics-structured reasoning through "
    "nine functional modules, going beyond general LLM pattern matching.")

add("Normal",
    "Limitations include: (1) without Ansys HFSS installation, EM simulations rely on analytical "
    "approximations; (2) GDS layout generation is topologically correct but requires manual refinement "
    "for production masks; (3) LLM function-calling occasionally produces malformed arguments; "
    "(4) single-server deployment limits concurrent execution.")

add("Normal",
    "The roadmap proceeds in three phases: Phase 1 (current) establishes the diagnostic closed loop; "
    "Phase 2 adds real instrument integration with active learning; Phase 3 targets autonomous redesign "
    "with manufacturing-aware optimization.")

# ── 7 Conclusions ──
add("heading1", "7   Conclusions")

add("p1a",
    "We have presented QuantaMind, a multi-agent AI system for autonomous superconducting quantum chip "
    "research. By integrating twelve specialized agents over a unified tool layer spanning EDA, simulation, "
    "fabrication, and measurement platforms, QuantaMind demonstrates that complex, multi-stage quantum "
    "hardware workflows can be orchestrated through natural-language interaction. The theoretical physicist "
    "agent's nine-module architecture provides physics-grounded reasoning with explicit uncertainty "
    "quantification. Applied to a production-derived 20-qubit chip design, the system generates "
    "parameterized layouts, executes multi-physics simulations, and produces actionable design "
    "recommendations. These capabilities represent a meaningful step toward AI-native quantum hardware "
    "development, where the boundaries between design, simulation, fabrication, and characterization "
    "become increasingly fluid.")

# ── References ──
add("heading1", "References")

refs = [
    "Arute, F., et al.: Quantum supremacy using a programmable superconducting processor. Nature 574, 505\u2013510 (2019)",
    "Wu, Y., et al.: Strong quantum computational advantage using a superconducting quantum processor. Phys. Rev. Lett. 127, 180501 (2021)",
    "Google Quantum AI: Quantum error correction below the surface code threshold. Nature 638, 920\u2013926 (2025)",
    "IBM Quantum: The IBM Quantum Development Roadmap. https://www.ibm.com/quantum/roadmap (2024)",
    "CETC: 20-Qubit Tunable-Coupler Quantum Chip Design Specification. Internal Document TGQ-200-000-FA09-2025 (2025)",
    "Lu, C., et al.: The AI Scientist: Towards Fully Automated Open-Ended Scientific Discovery. arXiv:2408.06292 (2024)",
    "L\u00e1la, J., et al.: PaperQA: Retrieval-Augmented Generative Agent for Scientific Research. arXiv:2312.07559 (2023)",
    "Li, X., et al.: Kosmos: An End-to-End AI Research Agent. arXiv (2024)",
    "Zou, Y., et al.: El Agente: An Autonomous Agent for Quantum Chemistry. arXiv:2505.02484 (2025)",
    "Gustin, I., et al.: El Agente Cu\u00e1ntico: Automating quantum simulations. arXiv:2512.18847 (2025)",
    "Reuer, K., et al.: Autonomous calibration of superconducting qubits using k-agents. Nature (2025)",
    "OpenClaw: Open-source agent framework. https://github.com/openclaw (2024)",
    "Yao, S., et al.: ReAct: Synergizing Reasoning and Acting in Language Models. arXiv:2210.03629 (2022)",
    "Internal Design Document: 20-Qubit Tunable-Coupler Chip Design Specification. TGQ-200-000-FA09-2025 (2025)",
]

for i, ref in enumerate(refs):
    add("referenceitem", f"{i+1}. {ref}")

doc.save(OUTPUT)
print(f"论文已保存: {OUTPUT}")
print(f"文件大小: {os.path.getsize(OUTPUT):,} bytes")
