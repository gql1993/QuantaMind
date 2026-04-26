"""Expand framework paper from 17.7 to 20+ pages."""
import os, re, copy
from docx import Document

SRC = r"E:\work\QuantaMind\docs\QuantaMind_Framework_Paper.docx"
OUT = r"E:\work\QuantaMind\docs\QuantaMind_Framework_Paper_v2.docx"

doc = Document(SRC)

def _s(name):
    try: return doc.styles[name]
    except: return doc.styles['Normal']

def insert_after(text_start, style, new_text):
    for p in doc.paragraphs:
        if p.text.strip().startswith(text_start[:40]):
            new_p = doc.add_paragraph(new_text, style=_s(style))
            p._element.addnext(new_p._element)
            return True
    return False

# ── Expand Introduction ──
insert_after("Superconducting quantum computing has emerged",
    "Normal",
    "The complexity of this development cycle is compounded by the deep interdependence between stages. "
    "A change in Josephson junction resistance during fabrication shifts qubit frequencies, which in turn "
    "affects the optimal coupler bias point, readout dispersive shift, and gate pulse parameters. Diagnosing "
    "whether a measured performance degradation originates from design limitations, fabrication variability, "
    "or measurement setup issues requires expertise spanning multiple domains simultaneously. A recent "
    "comprehensive survey on Quantum Design Automation (QDA) by Wu et al. [57] identified this interdependence "
    "as a fundamental challenge, advocating for holistic co-design frameworks that integrate physical-level "
    "and logic-level optimization across the entire quantum computing stack.")

insert_after("Despite these advances, the application of LLM-based agents",
    "Normal",
    "The gap between existing agent systems and quantum hardware automation is significant. Computational "
    "chemistry workflows are predominantly software-based (generating input files, running simulations, "
    "parsing outputs), whereas quantum hardware workflows involve physical artifacts (GDS layout files, "
    "fabricated chips, cryogenic measurement data) and require real-time interaction with instruments and "
    "manufacturing execution systems with hundreds of database tables. Moreover, error diagnosis in quantum "
    "hardware is fundamentally more complex: a low gate fidelity could stem from frequency collisions in "
    "the design, junction parameter drift during fabrication, insufficient magnetic shielding, thermal "
    "photon population, or pulse distortion in the control electronics. Attributing performance to root "
    "causes requires combining theoretical models, simulation data, and experimental evidence in a "
    "structured multi-agent reasoning framework that goes beyond what any single LLM can accomplish.")

# ── Expand Related Work ──
insert_after("The EDA community has begun to develop",
    "Normal",
    "The QDesignOptimizer framework [58] represents a notable advance in automated superconducting circuit "
    "design, combining high-accuracy electromagnetic simulations in Ansys HFSS with Energy Participation "
    "Ratio (EPR) analysis integrated into Qiskit Metal. For larger-scale chips, Li et al. [59] proposed a "
    "graph neural network (GNN)-based parameter design algorithm that achieves 51% error reduction compared "
    "to conventional methods for circuits approaching 870 qubits. Li et al. [60] demonstrated LLM-assisted "
    "superconducting qubit experiments where an LLM framework autonomously performs resonator characterization "
    "and QND measurements. The QUASAR framework [61] addresses quantum circuit generation through agentic "
    "reinforcement learning, achieving 99.31% validity in quantum assembly code generation. Google's work "
    "on reinforcement learning for quantum error correction [62] demonstrates that calibration can be unified "
    "with computation through continuous parameter adjustment during operation. These developments highlight "
    "the growing sophistication of individual automation tools, but none provides end-to-end orchestration "
    "across the full chip lifecycle.")

# ── Expand Architecture ──
insert_after("A critical design feature of QuantaMind is its graceful degradation",
    "Normal",
    "The graceful degradation mechanism operates at the adapter level within the Hands layer. Each of the "
    "ten platform adapters implements a connection check function that dynamically detects whether the target "
    "service is available. When the service is reachable, the adapter issues real API calls and returns "
    "authentic results. When unavailable, it falls back to physics-informed mock responses that maintain "
    "the correct data schema and provide analytically computed values rather than empty or random data. "
    "For example, the simulation adapter returns approximate eigenmode frequencies computed from analytical "
    "transmon formulas when Ansys HFSS is not installed, while the CHIPMES adapter returns representative "
    "manufacturing data when the MES database is offline. This dual-mode design ensures that agents can "
    "complete their reasoning chains even in partially configured environments, a common situation in "
    "research labs where instrument availability varies day to day.")

# ── Expand Theoretical Physicist ──
insert_after("Module M8 maintains a structured knowledge graph",
    "Normal",
    "The nine-module architecture of the Theoretical Physicist agent represents a deliberate design choice "
    "to decompose complex physics reasoning into well-defined, composable units rather than relying on "
    "monolithic LLM generation. Each module has explicit input/output contracts: M0 produces DeviceGraph "
    "objects consumed by M1; M1 produces HamiltonianModel objects consumed by M2, M3, and M5; M4 produces "
    "DiagnosisReport objects consumed by M7. This separation enables independent validation of each module's "
    "outputs, targeted improvement of specific capabilities, and transparent provenance tracking. The "
    "alternative approach of prompting a general-purpose LLM to 'analyze the noise mechanisms' risks "
    "producing plausible but ungrounded explanations that conflate different physical scales and regimes. "
    "By structuring the analysis through explicit modules with defined approximation conditions and validity "
    "checks, the system provides guardrails that increase the reliability of physics-informed recommendations.")

# ── Expand Discussion ──
insert_after("QuantaMind demonstrates several unique strengths",
    "Normal",
    "The experimental results reveal a clear advantage of specialized multi-agent orchestration over "
    "monolithic approaches for complex quantum hardware tasks. The 42% reduction in reasoning rounds "
    "(Experiment 1) demonstrates that domain-specific tool filtering substantially reduces the search "
    "space for tool selection, while the 84% routing accuracy (Experiment 2) shows that even a simple "
    "keyword-based mechanism can effectively dispatch most queries to the appropriate specialist. The "
    "remaining 16% of routing errors concentrate on inherently ambiguous cross-domain queries, suggesting "
    "that the performance ceiling for keyword-based routing has been approached and that learned routing "
    "models or collaborative multi-agent responses may be needed for further improvement.")

insert_after("Future development of QuantaMind will focus",
    "Normal",
    "A particularly promising direction is the integration of formal verification methods for AI-generated "
    "quantum device designs. While the current system provides physics-structured guardrails through the "
    "theoretical physicist's module architecture, it does not formally guarantee that generated designs "
    "satisfy all specifications. Future work could incorporate constraint checking modules that verify "
    "frequency collision avoidance, minimum coupling strengths, and fabrication rule compliance before "
    "committing to GDS export. Additionally, the conversation pipeline protocol naturally supports "
    "retrospective analysis of agent decision-making, enabling systematic identification of failure "
    "modes and targeted improvement of individual agent capabilities.")

doc.save(OUT)

# Verify
doc2 = Document(OUT)
total = sum(len(p.text) for p in doc2.paragraphs)
imgs = sum(1 for rel in doc2.part.rels.values() if 'image' in rel.reltype)
chinese = any(re.findall(r'[\u4e00-\u9fff]', p.text) for p in doc2.paragraphs)
print(f"Saved: {OUT}")
print(f"Size: {os.path.getsize(OUT):,} bytes")
print(f"Paragraphs: {len(doc2.paragraphs)}")
print(f"Tables: {len(doc2.tables)}, Images: {imgs}")
print(f"Total chars: {total:,}, Pages: {total/3500:.1f}")
print(f"Chinese: {chinese}")
