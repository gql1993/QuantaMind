# -*- coding: utf-8 -*-
"""将 Markdown 转为 Word (.docx)，支持指定文件名。"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def strip_md(s):
    """去掉 Markdown 粗体标记，还原 HTML 实体"""
    if not s:
        return s
    s = re.sub(r'\*\*(.+?)\*\*', r'\1', s)
    s = s.replace('&gt;', '>').replace('&lt;', '<')
    return s

def add_paragraph(doc, text, style=None):
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for run in re.split(r'(\*\*.+?\*\*)', text):
        r = p.add_run(strip_md(run))
        if run.startswith('**') and run.endswith('**'):
            r.bold = True
    if style:
        p.style = style

def parse_table(lines):
    """解析 Markdown 表格，返回 (headers, rows)"""
    rows = []
    for line in lines:
        line = line.strip()
        if not line or line.startswith('|---'):
            continue
        cells = [c.strip() for c in line.split('|') if c.strip() or line.split('|').index(c) > 0]
        if cells:
            rows.append(cells)
    if not rows:
        return [], []
    headers = rows[0]
    data = rows[1:] if len(rows) > 1 else []
    return headers, data

def main():
    docs_dir = Path(__file__).resolve().parent
    if len(sys.argv) >= 2:
        name = sys.argv[1]
        if not name.endswith('.md'):
            name = name + '.md'
        md_path = docs_dir / name
        suffix = sys.argv[2] if len(sys.argv) >= 3 else ''
        out_path = docs_dir / (md_path.stem + suffix + '.docx')
    else:
        md_path = docs_dir / '长三角AI+量子融合产业创新平台-软件开发详细设计说明书.md'
        out_path = docs_dir / '长三角AI+量子融合产业创新平台-软件开发详细设计说明书.docx'
    if not md_path.exists():
        print('File not found:', md_path)
        sys.exit(1)
    text = md_path.read_text(encoding='utf-8')
    lines = text.split('\n')
    doc = Document()
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(10.5)
    i = 0
    in_code = False
    code_lines = []
    table_lines = []
    in_table = False
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('```'):
            if in_code:
                code_lines.append(line)
                code_text = '\n'.join(code_lines[1:-1])  # 去掉首尾 ```
                if code_text.strip():
                    p = doc.add_paragraph(code_text)
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    for run in p.runs:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                code_lines = [line]
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        # 表格：以 | 开头的行
        if line.strip().startswith('|') and '|' in line:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        if in_table and table_lines:
            headers, data = parse_table(table_lines)
            if headers:
                ncols = len(headers)
                table = doc.add_table(rows=1 + len(data), cols=ncols)
                table.style = 'Table Grid'
                for j, h in enumerate(headers):
                    table.rows[0].cells[j].text = strip_md(h)
                for r, row in enumerate(data):
                    for j, cell in enumerate(row):
                        if j < ncols:
                            table.rows[r + 1].cells[j].text = strip_md(cell)
                doc.add_paragraph()
            in_table = False
            table_lines = []
        # 标题
        if line.startswith('# ') and not line.startswith('## '):
            doc.add_heading(strip_md(line[2:].strip()), level=0)
            i += 1
            continue
        if line.startswith('## ') and not line.startswith('### '):
            doc.add_heading(strip_md(line[3:].strip()), level=1)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(strip_md(line[4:].strip()), level=2)
            i += 1
            continue
        # 水平线
        if line.strip() == '---':
            i += 1
            continue
        # 列表
        if line.strip().startswith('- '):
            add_paragraph(doc, line.strip()[2:], style='List Bullet')
            i += 1
            continue
        if re.match(r'^\d+\.\s', line.strip()):
            add_paragraph(doc, re.sub(r'^\d+\.\s', '', line.strip()), style='List Number')
            i += 1
            continue
        # 普通段落
        if line.strip():
            add_paragraph(doc, line)
        i += 1
    # 收尾表格
    if in_table and table_lines:
        headers, data = parse_table(table_lines)
        if headers:
            ncols = len(headers)
            table = doc.add_table(rows=1 + len(data), cols=ncols)
            table.style = 'Table Grid'
            for j, h in enumerate(headers):
                table.rows[0].cells[j].text = strip_md(h)
            for r, row in enumerate(data):
                for j, cell in enumerate(row):
                    if j < ncols:
                        table.rows[r + 1].cells[j].text = strip_md(cell)
    doc.save(out_path)
    print('Saved:', out_path)

if __name__ == '__main__':
    main()
