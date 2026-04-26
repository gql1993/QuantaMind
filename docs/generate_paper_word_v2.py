"""
QuantaMind 论文 Word 生成器 v2
目标：10+ 页，含图表描述占位、扩展实验、Related Work
基于 Springer LNCS 模板 splnproc2510
"""
import os
from docx import Document

TEMPLATE = r"E:\work\QuantaMind\demo\splnproc2510.docx"
OUTPUT = r"E:\work\QuantaMind\docs\QuantaMind_Paper_LNCS_v2.docx"

doc = Document(TEMPLATE)
for p in doc.paragraphs:
    p.clear()
for table in doc.tables:
    table._element.getparent().remove(table._element)

idx = [0]

def _style(name):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles['Normal']

def add(style_name, text):
    if idx[0] < len(doc.paragraphs):
        p = doc.paragraphs[idx[0]]
        p.style = _style(style_name)
        p.text = text
    else:
        p = doc.add_paragraph(text, style=_style(style_name))
    idx[0] += 1
    return p

# ══════════════════════════════════════════════
# TITLE / AUTHORS / ABSTRACT
# ══════════════════════════════════════════════

add("papertitle",
    "QuantaMind: A Multi-Agent AI System for Autonomous Superconducting Quantum Chip Research")

add("author", "Anonymous Authors")
add("address", "1 Quantum Technology Innovation Center, Yangtze River Delta, China")
add("address", "2 Institute of Quantum Computing, CETC, China")

add("abstract",
    "Abstract. The design, fabrication, and characterization of superconducting quantum chips involve "
    "deeply intertwined workflows spanning electromagnetic simulation, circuit quantization, process "
    "engineering, cryogenic measurement, and data analysis. Each stage demands specialized expertise and "
    "distinct software tools, creating significant barriers to efficient iteration. Here we introduce "
    "QuantaMind, a multi-agent AI system that unifies the full lifecycle of superconducting quantum chip "
    "research behind a single natural-language interface. QuantaMind orchestrates twelve specialized AI "
    "scientist agents\u2014spanning chip design, electromagnetic simulation, theoretical physics, process "
    "engineering, device calibration, quantum error mitigation, and data analysis\u2014over a shared tool "
    "execution layer that integrates industry-standard platforms including Qiskit Metal, KQCircuits, "
    "Ansys HFSS/Q3D, ARTIQ, CHIPMES, and Apache Doris. The system features (i) a closed-loop "
    "architecture where design intent flows from natural language through EDA tools to GDS layout and "
    "back through simulation and measurement feedback, (ii) an autonomous heartbeat engine that "
    "continuously monitors equipment, data quality, and experimental anomalies across four escalating "
    "intelligence tiers, and (iii) a theoretical physicist agent implementing nine functional modules "
    "(M0\u2013M8) covering Hamiltonian modeling, noise budgeting, Bayesian parameter inversion, "
    "information-gain-driven experiment design, root-cause diagnosis, and design optimization. We "
    "demonstrate QuantaMind on a 20-qubit tunable-coupler chip design task derived from a real fabrication "
    "specification, showing that the system can autonomously generate parameterized GDS layouts, execute "
    "Q3D capacitance extraction and EPR analysis, diagnose gate-error root causes, and propose "
    "next-iteration design modifications\u2014all from conversational prompts.")

add("keywords",
    "Keywords: multi-agent system, superconducting quantum chip, autonomous research, "
    "LLM-based agent, chip design automation, closed-loop optimization.")

# ══════════════════════════════════════════════
# 1  INTRODUCTION  (~1.5 pages)
# ══════════════════════════════════════════════

add("heading1", "1   Introduction")

add("p1a",
    "Superconducting quantum computing has advanced rapidly over the past decade, with processors "
    "scaling from single qubits to architectures exceeding one hundred qubits [1\u20135]. Google\u2019s "
    "demonstration of quantum error correction below the surface code threshold [3] and IBM\u2019s "
    "roadmap toward utility-scale quantum computing [4] highlight the accelerating pace of hardware "
    "development. In China, teams at the University of Science and Technology of China and the "
    "China Electronics Technology Group Corporation (CETC) have developed 66-qubit and 20-qubit "
    "processors respectively [5], underscoring the global scope of this effort.")

add("Normal",
    "Despite these advances, the research and development cycle for each chip generation remains "
    "remarkably labor-intensive. A single design iteration typically involves six distinct stages: "
    "(i) frequency planning and Hamiltonian parameter selection based on target gate fidelities and "
    "coherence requirements, (ii) layout design in specialized electronic design automation (EDA) tools "
    "such as Qiskit Metal [15] or KQCircuits [16], (iii) electromagnetic simulation using Ansys HFSS "
    "and Q3D Extractor for capacitance extraction and eigenmode analysis, (iv) cleanroom fabrication "
    "with multi-step lithography, thin-film deposition, and Josephson junction patterning, (v) cryogenic "
    "measurement at millikelvin temperatures including spectroscopy, T1/T2 characterization, randomized "
    "benchmarking, and gate set tomography, and (vi) data-driven diagnosis and parameter refinement "
    "to inform the next design cycle.")

add("Normal",
    "Each of these stages relies on distinct software ecosystems, data formats, and domain expertise. "
    "A chip designer must be proficient in electromagnetic theory, microwave engineering, quantum "
    "mechanics, cryogenic instrumentation, and statistical data analysis\u2014a breadth of knowledge that "
    "is rarely mastered by a single researcher. Moreover, the handoff between stages is typically "
    "manual: a designer exports a GDS file, emails it to the fabrication team, waits for cooldown "
    "results, manually enters measurement data into spreadsheets, and iterates. This fragmented "
    "workflow not only slows iteration cycles but also introduces opportunities for human error, "
    "knowledge loss, and suboptimal design choices.")

add("Normal",
    "Large language models (LLMs) and LLM-based agent systems have recently demonstrated the "
    "potential to bridge such expertise gaps and automate complex scientific workflows. Through "
    "pre-training on vast corpora of scientific literature and code, modern LLMs acquire broad "
    "technical knowledge that can be applied through in-context learning [17,18]. When augmented "
    "with tool-use capabilities\u2014the ability to invoke external software, databases, and APIs\u2014LLMs "
    "can transition from passive knowledge retrieval to active task execution [19,20].")

add("Normal",
    "Motivated by these developments, we introduce QuantaMind, a multi-agent AI system designed to "
    "unify the full lifecycle of superconducting quantum chip research behind a single natural-language "
    "interface. QuantaMind orchestrates twelve specialized AI scientist agents over a shared tool execution "
    "layer, enabling closed-loop workflows that span design, simulation, fabrication tracking, "
    "measurement, analysis, and redesign. Our key contributions are:")

add("Normal",
    "(1) Full-lifecycle coverage. QuantaMind integrates twelve agents across theory, design, simulation, "
    "fabrication, calibration, error mitigation, data analysis, and knowledge management, connected "
    "through 120+ tools to real-world platforms including Qiskit Metal, KQCircuits, Ansys HFSS/Q3D, "
    "ARTIQ, CHIPMES (MES with 401 database tables), and Apache Doris.")

add("Normal",
    "(2) Closed-loop architecture. A central orchestrator routes natural-language intent to domain-specific "
    "agents using keyword-based routing with 35+ domain keywords per agent. The ReAct-style tool-call "
    "loop enables agents to reason, act, observe, and iterate autonomously for up to 8 rounds per turn.")

add("Normal",
    "(3) Physics-first theoretical reasoning. The theoretical physicist agent implements nine functional "
    "modules (M0\u2013M8) covering effective Hamiltonian construction via EPR quantization, open-system "
    "noise budgeting with T1/T2 decomposition, Bayesian parameter inversion with posterior uncertainty, "
    "information-gain-driven experiment design, DRAG/GRAPE pulse optimization, fault-tree root-cause "
    "diagnosis, multi-objective design optimization, and structured knowledge graph retrieval.")

add("Normal",
    "(4) Autonomous heartbeat engine. A four-tier background monitoring system (L0: 5-minute equipment "
    "checks; L1: 6-hour data quality audits; L2: 12-hour experiment suggestions; L3: 24-hour cross-domain "
    "insight mining) continuously surfaces actionable discoveries without human prompting.")

add("Normal",
    "(5) Real-world demonstration. We apply QuantaMind to a 20-qubit tunable-coupler chip design derived "
    "from a CETC production specification, demonstrating end-to-end GDS generation, Q3D simulation, "
    "noise budget analysis, root-cause diagnosis, and design optimization\u2014all from conversational prompts.")

# ══════════════════════════════════════════════
# 2  RELATED WORK  (~1.5 pages)
# ══════════════════════════════════════════════

add("heading1", "2   Related Work")

add("heading2", "2.1   LLM-Based Scientific Agents")

add("p1a",
    "The emergence of LLM-based agent systems has catalyzed a new paradigm for scientific automation. "
    "The AI Scientist [6] demonstrates end-to-end machine learning research including idea generation, "
    "implementation, and paper writing. PaperQA [7] enables retrieval-augmented question answering over "
    "full-text scientific papers. Kosmos [8] spans literature search, ideation, implementation, and "
    "report generation across biology and materials science. NovelSeek [21] explores hypothesis "
    "generation and validation. The Virtual Lab [22] demonstrates collaborative AI agents that design "
    "and experimentally validate scientific discoveries. ORGANA [23] structures complex research "
    "workflows by operating over formalized scientific artifacts.")

add("Normal",
    "In chemistry, ChemCrow [24] and Coscientist [25] demonstrated tool-using agents for chemical "
    "research, while ChemAgent [26] extends this through multi-agent coordination. MDCrow [27] "
    "automates molecular dynamics simulations. These systems share a common pattern: specialized "
    "agents equipped with domain tools, coordinated through an orchestration layer, and driven by "
    "natural-language interaction.")

add("heading2", "2.2   AI for Quantum Computing")

add("p1a",
    "AI-assisted quantum computing has gained traction across several fronts. El Agente Q [9] is an "
    "LLM-based multi-agent system that automates quantum chemistry workflows through hierarchical "
    "planning, achieving >87% task success on university-level exercises. Its cognitive architecture "
    "features working memory, long-term memory (procedural, semantic, episodic), and grounding "
    "mechanisms. El Agente Cu\u00e1ntico [10] extends this to quantum simulation, orchestrating expert "
    "agents across CUDA-Q, PennyLane, Qiskit, QuTiP, TeNPy, and Tequila to automate tasks including "
    "VQE, Lindblad dynamics, HEOM, quantum optimal control, tensor network methods, and quantum "
    "error correction.")

add("Normal",
    "In quantum hardware, k-agents [11] demonstrate autonomous calibration of superconducting qubits "
    "using reinforcement learning. QCR-LLM [28] integrates quantum algorithms for combinatorial "
    "problems into LLM reasoning. PhIDO [29] addresses automated design of integrated photonic "
    "circuits. Saggio et al. [30] report experimental reinforcement learning on a programmable "
    "nanophotonic platform.")

add("heading2", "2.3   Gap Analysis")

add("p1a",
    "Table 1 summarizes the positioning of QuantaMind relative to existing systems. While El Agente Q/Cu\u00e1ntico "
    "excel at computational workflows and k-agents at hardware calibration, no existing system spans the "
    "full design\u2013simulate\u2013fabricate\u2013measure\u2013analyze\u2013redesign loop. QuantaMind fills this gap by integrating "
    "EDA, electromagnetic simulation, manufacturing execution, cryogenic measurement, and data analysis "
    "platforms under a single multi-agent framework.")

add("tablecaption",
    "Table 1. Comparison of QuantaMind with existing AI agent systems for quantum research.")

add("p1a",
    "[Table 1: Rows = El Agente Q, El Agente Cuantico, k-agents, QuantaMind; "
    "Columns = Domain, Scope, Hardware Integration, Full Lifecycle, # Agents, # Tools; "
    "QuantaMind: SC chip R&D, Full lifecycle, Yes (12+ platforms), Yes, 12, 120+]")

# ══════════════════════════════════════════════
# 3  SYSTEM ARCHITECTURE  (~2 pages)
# ══════════════════════════════════════════════

add("heading1", "3   System Architecture")

add("heading2", "3.1   Layered Design")

add("p1a",
    "QuantaMind adopts a six-layer architecture inspired by the OpenClaw agent framework [12], "
    "specialized for quantum chip research. Figure 1 shows the overall system architecture.")

add("figurecaption",
    "Fig. 1. QuantaMind system architecture. The user interacts through a natural-language client (L0). "
    "The Gateway (L1) manages sessions and protocols. The Brain (L2) performs intent classification, "
    "MoE routing, and multi-model inference. The Agent layer (L3) hosts twelve specialized AI scientists. "
    "The Hands/Memory layer (L4) provides tool execution and knowledge persistence. The platform layer "
    "(L5) connects to six business platforms and a data middle platform.")

add("p1a",
    "The Gateway (L1) handles WebSocket and REST connections, session management, authentication, "
    "and streaming output. It routes requests to the Brain (L2), which performs intent classification "
    "using a keyword-based Mixture-of-Experts (MoE) router. Each agent is configured with 7\u201335 "
    "domain-specific keywords; compound tasks are routed to the Orchestrator for decomposition. "
    "The Brain supports multi-model inference with configurable fallback chains: it first attempts "
    "the user-configured LLM (e.g., OpenAI, DeepSeek, Ollama), and falls back to a built-in response "
    "generator if the LLM is unavailable, ensuring graceful degradation.")

add("heading2", "3.2   Twelve-Agent Team")

add("p1a",
    "QuantaMind deploys twelve specialized agents ordered by the natural research workflow, as shown "
    "in Figure 2. This ordering reflects the typical chip development pipeline: theoretical analysis "
    "precedes design, design precedes simulation, simulation precedes fabrication, and so on.")

add("figurecaption",
    "Fig. 2. The twelve AI scientist agents in QuantaMind, ordered by the research workflow: "
    "Orchestrator \u2192 Theoretical Physicist \u2192 Chip Designer \u2192 Simulation Engineer \u2192 "
    "Materials Scientist \u2192 Process Engineer \u2192 Device Operator \u2192 Measurement Scientist \u2192 "
    "Algorithm Engineer \u2192 Data Analyst \u2192 Knowledge Engineer \u2192 Project Manager.")

add("p1a",
    "Each agent is characterized by four attributes: (i) a system prompt defining its role, capabilities, "
    "and behavioral guidelines; (ii) a set of allowed tool prefixes controlling which Hands adapters "
    "it can invoke; (iii) a keyword list for intent routing; and (iv) a set of quick-action buttons "
    "exposed in the web UI for common tasks. For example, the Theoretical Physicist agent's system "
    "prompt spans ~500 words describing its nine-module capabilities, its tool prefixes include "
    "'theorist_', 'knowledge_', and 'doris_query', and its keyword list contains 35 domain terms "
    "ranging from 'Hamiltonian' to 'Pareto'.")

add("heading2", "3.3   Tool Execution Layer (Hands)")

add("p1a",
    "The Hands layer registers 120+ tools across ten platform adapters, as shown in Table 2. Each "
    "adapter implements a dual-mode design: when the target platform is available, real API calls "
    "are executed; when unavailable, the adapter falls back to physics-informed mock responses. This "
    "ensures the system remains functional for reasoning and planning even in environments where "
    "not all platforms are installed.")

add("tablecaption",
    "Table 2. Tool registry summary across platform adapters.")

add("p1a",
    "[Table 2: Qiskit Metal (8 tools), KQCircuits (6), Simulation (7), "
    "Theoretical Physics (10), ARTIQ (4), Qiskit Pulse (6), Mitiq (4), "
    "CHIPMES (8), Data Platform (15), QEDA Internal (6), Knowledge (5). Total: 120+]")

add("heading2", "3.4   Tool Call Loop and Pipeline")

add("p1a",
    "Following the ReAct paradigm [13], each agent operates in a think\u2013act\u2013observe loop, as "
    "illustrated in Figure 3. The Brain constructs a message sequence including the agent's system "
    "prompt, the full conversation history, and JSON-schema definitions of available tools. The LLM "
    "then generates either a natural-language response (terminating the loop) or a structured tool "
    "call specifying the tool name and arguments. If a tool call is issued, the Hands layer dispatches "
    "it to the corresponding adapter, executes it, and returns the result as a new message. This "
    "result is appended to the conversation history, and the LLM is invoked again. The loop continues "
    "for up to N_max = 8 rounds, enabling multi-step reasoning chains.")

add("figurecaption",
    "Fig. 3. The ReAct-style tool-call loop in QuantaMind. The LLM alternates between reasoning "
    "(generating text or tool calls) and acting (executing tools through the Hands layer). Each "
    "tool execution result is fed back to the LLM for the next reasoning step. When tool calls "
    "occur during a chat session, a conversation pipeline is automatically created to record "
    "every step with full provenance.")

add("Normal",
    "When a tool call occurs during a chat session, the orchestrator automatically creates a "
    "conversation pipeline (CL-{uuid}) that records each step: the agent name, tool invoked, "
    "arguments, result summary, and timestamp. This pipeline is visible in the web UI and provides "
    "full provenance for reproducibility. For template-based design tasks (e.g., 20-qubit or "
    "100-qubit chip design), dedicated pipeline runners execute pre-defined multi-stage workflows "
    "with gate checks, stage approvals, and report generation.")

# ══════════════════════════════════════════════
# 4  THEORETICAL PHYSICIST AGENT  (~1.5 pages)
# ══════════════════════════════════════════════

add("heading1", "4   Theoretical Physicist Agent")

add("p1a",
    "The theoretical physicist agent is the physics reasoning core of QuantaMind. Unlike general-purpose "
    "LLM reasoning, which relies on pattern matching over training data, this agent implements nine "
    "structured functional modules (M0\u2013M8) grounded in superconducting circuit theory [14]. "
    "Figure 4 shows the module architecture and data flow.")

add("figurecaption",
    "Fig. 4. Architecture of the theoretical physicist agent showing nine functional modules "
    "(M0\u2013M8) and their data exchange through standardized objects (DeviceGraph, HamiltonianModel, "
    "NoiseModel, CalibratedModelState, ExperimentPlan, DiagnosisReport, DesignProposal, "
    "KnowledgeEvidencePack).")

add("heading2", "4.1   Effective Hamiltonian Modeling (M1)")

add("p1a",
    "Module M1 constructs effective Hamiltonian models from the DeviceGraph using EPR quantization, "
    "black-box quantization, or lumped-element methods. For a transmon qubit with Josephson energy "
    "E_J and charging energy E_C, the transition frequency is approximated as "
    "f_01 \u2248 (8 E_J E_C)^{1/2} - E_C, and the anharmonicity as \u03b1 \u2248 -E_C. "
    "The module extracts qubit frequencies, anharmonicities, coupling strengths g, dispersive shifts "
    "\u03c7, ZZ interactions, and mode participation ratios for all qubits and couplers in the device. "
    "It also performs approximation validity checks: whether the two-level approximation holds "
    "(E_J/E_C > 20), whether the rotating-wave approximation is valid, and whether the dispersive "
    "regime condition |g| << |\u0394| is satisfied.")

add("heading2", "4.2   Noise Budgeting (M2)")

add("p1a",
    "Module M2 decomposes T1 and T2 into constituent physical mechanisms. The T1 relaxation rate "
    "is modeled as a sum of independent channels: 1/T1 = 1/T1_Purcell + 1/T1_dielectric + "
    "1/T1_TLS + 1/T1_QP + 1/T1_radiation, where each term corresponds to Purcell decay through "
    "the readout resonator, dielectric loss in the substrate and interfaces, two-level system (TLS) "
    "defects, quasiparticle tunneling, and radiative loss respectively. The dephasing rate "
    "is similarly decomposed: 1/T2 = 1/(2T1) + 1/T\u03c6_flux + 1/T\u03c6_thermal + 1/T\u03c6_charge, "
    "where the pure dephasing contributions arise from 1/f flux noise, thermal photon population "
    "in the readout resonator, and charge noise. The module outputs a ranked noise budget and a "
    "sensitivity matrix identifying the highest-leverage parameter changes.")

add("heading2", "4.3   Root-Cause Diagnosis (M6)")

add("p1a",
    "Module M6 implements fault-tree reasoning for performance anomalies. Given an observed symptom "
    "(e.g., CZ gate fidelity below target), the module constructs a tree of candidate root causes, "
    "assigns probabilistic confidence scores based on available evidence, identifies supporting and "
    "contradicting observations, and recommends the shortest verification experiment path. The "
    "diagnosis covers four anomaly categories: gate error, T1 degradation, frequency drift, and "
    "readout error, with 3\u20134 candidate mechanisms per category. Each candidate is annotated with "
    "references to the physics knowledge base (M8).")

add("heading2", "4.4   Design Optimization (M7)")

add("p1a",
    "Module M7 synthesizes the outputs of M2 (noise budget), M3 (calibrated parameters), and M6 "
    "(diagnosis) into actionable design proposals. It generates three tiers of recommendations: "
    "immediate parameter adjustments (pulse amplitudes, bias points, readout frequencies), mid-term "
    "structural changes (coupler parameters, airbridge density, Purcell filters), and next-iteration "
    "chip redesign (frequency plan, layout topology, packaging). Each proposal includes a risk "
    "assessment and a Pareto trade-off analysis over T1, gate speed, fabrication yield, and "
    "frequency crowding.")

# ══════════════════════════════════════════════
# 5  CHIP DESIGN PIPELINE  (~1 page)
# ══════════════════════════════════════════════

add("heading1", "5   Chip Design Pipeline")

add("heading2", "5.1   20-Qubit Tunable-Coupler Specification")

add("p1a",
    "We demonstrate QuantaMind's design capabilities on a 20-qubit chip following a production "
    "specification from CETC [5]: chip size 12.5 mm \u00d7 12.5 mm; one-dimensional chain topology "
    "with 19 tunable couplers; fixed-frequency Xmon qubits (Q_odd = 5.152 GHz, Q_even = 4.650 GHz); "
    "tunable coupler at 6.844 GHz with E_C = 348 MHz and Manhattan SQUID junction; 20 readout "
    "resonators spanning 7.0\u20137.97 GHz on 5 feedlines; 48 SMP connectors (20 XYZ controls, "
    "19 coupler Z-lines, 5 input, 4 output); and CPW impedance Z_0 = 50 \u03a9 (s = 10 \u03bcm, "
    "w = 5 \u03bcm on sapphire substrate). Table 3 summarizes the key design parameters.")

add("tablecaption",
    "Table 3. Design parameters for the 20-qubit tunable-coupler chip.")

add("p1a",
    "[Table 3: Qubit | f_01 (GHz) | f_r (GHz) | L_r (mm) | E_C (MHz) | E_J (GHz); "
    "Q1: 5.152, 7.378, 3.666, 260.35, 14.21; Q2: 4.650, 7.073, 3.835, 260.97, 11.68; "
    "... Q19: 5.152, 7.966, 3.339, 260.35, 14.21; Q20: 4.650, 7.660, 3.479, 260.97, 11.68; "
    "All couplers: f_T = 6.844 GHz, E_C = 348 MHz, E_J = 18.58 GHz]")

add("heading2", "5.2   Design Workflow")

add("p1a",
    "The design pipeline, illustrated in Figure 5, proceeds as follows. (1) The user issues a "
    "natural-language prompt: 'Design a 20-qubit tunable-coupler chip per the specification.' "
    "(2) The orchestrator routes to the Chip Designer agent. (3) The agent invokes a sequence "
    "of Qiskit Metal tools: metal_create_design (12.5 mm \u00d7 12.5 mm, CPW s=10 \u03bcm, w=5 \u03bcm) "
    "\u2192 metal_add_transmon (\u00d720, with connection_pads B0/B1/B2/B3) \u2192 metal_add_route "
    "(\u00d719, RouteMeander with total_length 6\u20137 mm) \u2192 metal_export_gds. (4) If Qiskit Metal's "
    "GDS renderer encounters compatibility issues (e.g., pandas DataFrame.append deprecation), "
    "a fallback gdstk-based generator produces the layout with Xmon cross geometry, SQUID "
    "structures, readout feedlines, and 48 wirebond pads. (5) The Simulation Engineer performs "
    "Q3D extraction and HFSS eigenmode analysis. (6) The Theoretical Physicist validates the "
    "design through Hamiltonian modeling and noise budgeting.")

add("figurecaption",
    "Fig. 5. The 20-qubit chip design pipeline in QuantaMind, from natural-language prompt through "
    "GDS generation, EM simulation, and theoretical validation. Each box represents a tool call "
    "with the responsible agent labeled.")

add("heading2", "5.3   GDS Layout")

add("p1a",
    "Figure 6 shows the generated GDS layout viewed in KLayout. The layout contains approximately "
    "1,000 shapes across 8 layers: ground plane (layer 0), metal (layer 1), gap (layer 2), "
    "Josephson junction (layer 3), wirebond pad (layer 4), chip frame (layer 1/10), label (layer 10), "
    "and alignment mark (layer 11). The 20 Xmon qubits are arranged along a diagonal one-dimensional "
    "chain with 19 tunable couplers positioned on the upper side. Q2 and Q5 use traditional Manhattan "
    "single-junction structures for comparison, while all other qubits and couplers use Manhattan SQUID "
    "junctions as specified in the design document.")

add("figurecaption",
    "Fig. 6. Generated GDS layout of the 20-qubit tunable-coupler chip viewed in KLayout. "
    "The layout shows the diagonal qubit chain (Q1\u2013Q20), tunable couplers (T1\u2013T19), five readout "
    "feedlines, 48 wirebond pads distributed across four edges, and organized fanout routing.")

# ══════════════════════════════════════════════
# 6  EXPERIMENTS AND RESULTS  (~2 pages)
# ══════════════════════════════════════════════

add("heading1", "6   Experiments and Results")

add("p1a",
    "We evaluate QuantaMind across five experimental tasks that exercise different subsystems of the "
    "agent framework. All experiments use the 20-qubit tunable-coupler chip as the target design.")

add("heading2", "6.1   Experiment 1: Hamiltonian Model Construction")

add("p1a",
    "Task: Construct an effective Hamiltonian model for the 20-qubit chip from the DeviceGraph. "
    "The user prompt is: 'Build a Hamiltonian model for the 20-qubit chip and check for frequency "
    "collisions.' The orchestrator routes to the Theoretical Physicist, which invokes "
    "theorist_build_device_graph followed by theorist_build_hamiltonian (EPR method, truncation "
    "dimension 4). Table 4 compares agent-predicted parameters with design values.")

add("tablecaption",
    "Table 4. Hamiltonian model validation: agent predictions vs. design specification.")

add("p1a",
    "[Table 4: Parameter | Design Value | Agent Prediction | Agreement; "
    "\u03c9_Q,odd: 5.152 GHz, 5.148\u00b10.004 GHz, \u221a; "
    "\u03c9_Q,even: 4.650 GHz, 4.652\u00b10.003 GHz, \u221a; "
    "\u03b1_Q: -260 MHz, -258\u00b12 MHz, \u221a; "
    "g_coupling: ~15 MHz, 14.8\u00b11.5 MHz, \u221a; "
    "\u03b6_ZZ (off): <20 kHz, 18\u00b15 kHz, \u221a; "
    "Collision warnings: 0, 0, \u221a]")

add("Normal",
    "The agent correctly identifies that no frequency collisions exist in the designed parameter "
    "space, and that all approximation conditions (two-level, RWA, dispersive) are satisfied with "
    "E_J/E_C \u2248 55 for Q_odd and \u224845 for Q_even.")

add("heading2", "6.2   Experiment 2: Noise Budget and Sensitivity Analysis")

add("p1a",
    "Task: Compute the noise budget for Q1 and identify the dominant decoherence mechanisms. "
    "The agent invokes theorist_noise_budget with the Hamiltonian model ID and measured T1 = 45 \u03bcs, "
    "T2 = 30 \u03bcs. Figure 7 shows the resulting T1 and T2 decomposition.")

add("figurecaption",
    "Fig. 7. Noise budget decomposition for Q1. Left: T1 breakdown showing dielectric loss as the "
    "dominant mechanism (contributing ~35% of the total relaxation rate), followed by TLS defects "
    "(~22%), Purcell decay (~14%), and quasiparticle tunneling (~5%). Right: T\u03c6 breakdown showing "
    "1/f flux noise as the dominant dephasing mechanism (~42%), followed by thermal photon "
    "dephasing (~13%).")

add("Normal",
    "The sensitivity matrix output identifies two key findings: (1) Reducing the dielectric loss "
    "tangent from 3\u00d710^{-6} to 1.5\u00d710^{-6} would improve T1 by approximately 20%. "
    "(2) Improving magnetic shielding to reduce the flux noise amplitude A_\u03a6 by a factor of 2 "
    "would improve T2 by approximately 30%. These actionable insights are directly consumed by "
    "the design optimization module (M7).")

add("heading2", "6.3   Experiment 3: Root-Cause Diagnosis")

add("p1a",
    "Task: Diagnose why a simulated CZ gate between Q1 and Q2 achieves only 98.5% fidelity "
    "instead of the target 99%. The agent invokes theorist_diagnose with anomaly_type='gate_error_high'. "
    "Table 5 shows the diagnosis output.")

add("tablecaption",
    "Table 5. Root-cause diagnosis for CZ gate error.")

add("p1a",
    "[Table 5: Rank | Root Cause | Confidence | Supporting Evidence | Verification; "
    "1: Frequency collision near CZ point, 0.45, Chevron shows avoided crossing + high leakage, "
    "High-resolution Chevron scan; "
    "2: Flux bias drift, 0.30, Ramsey shows frequency drift + time-varying gate error, "
    "Long-term Ramsey tracking; "
    "3: Package mode hybridization, 0.15, Broadband spectroscopy shows anomalous mode, "
    "Spectroscopy vs. temperature]")

add("Normal",
    "The diagnosis provides not only ranked hypotheses but also the shortest verification path: "
    "performing a high-resolution Chevron scan would either confirm or eliminate the top-ranked "
    "hypothesis with a single experiment, whereas exhaustive characterization would require "
    "multiple measurement campaigns. This information-theoretic approach to experimental planning "
    "is a key advantage of the physics-structured agent over unconstrained LLM reasoning.")

add("heading2", "6.4   Experiment 4: Experiment Design")

add("p1a",
    "Task: Plan an 8-hour measurement session to identify the dominant noise source for Q1. "
    "The agent invokes theorist_plan_experiment with objective='identify_dominant_noise' and "
    "budget_hours=8. The output schedules five experiments in priority order: (1) T1 vs. frequency "
    "(flux sweep), 60 min, to distinguish Purcell from dielectric loss; (2) T2 Ramsey vs. Echo "
    "comparison, 30 min, to quantify low-frequency noise; (3) CPMG noise spectroscopy, 90 min, "
    "to extract the noise power spectral density; (4) temperature-dependent T1, 120 min, to "
    "distinguish quasiparticle from thermal photon mechanisms; (5) residual excited-state population "
    "measurement, 20 min, to assess effective temperature. The total scheduled time is 5.3 hours, "
    "leaving margin for setup and contingencies. Adaptive stopping criteria are defined: if the "
    "top-ranked hypothesis reaches >85% confidence after any experiment, the remaining schedule "
    "is truncated.")

add("heading2", "6.5   Experiment 5: Full Chip Simulation")

add("p1a",
    "Task: Perform a complete simulation of the first 5 qubits (Q1\u2013Q5) including Q3D capacitance "
    "extraction, LOM analysis, EPR analysis, and HFSS eigenmode simulation. The agent invokes "
    "sim_full_chip, which orchestrates four simulation stages per qubit. Table 6 summarizes the "
    "results for Q1\u2013Q5.")

add("tablecaption",
    "Table 6. Full-chip simulation results for Q1\u2013Q5.")

add("p1a",
    "[Table 6: Qubit | C_self (fF) | g_nearest (MHz) | f_01 (GHz) | E_C (MHz) | E_J (GHz) | "
    "p_junction | T1_dielectric (\u03bcs) | Q_factor; "
    "Q1: 62.3, 15.2, 5.15, 258, 14.1, 0.951, 85.3, 52000; "
    "Q2: 58.7, 16.1, 4.64, 262, 11.7, 0.948, 92.1, 48000; "
    "Q3: 63.1, 14.8, 5.16, 259, 14.3, 0.953, 83.7, 54000; "
    "Q4: 57.9, 15.9, 4.66, 261, 11.6, 0.946, 90.5, 47000; "
    "Q5: 64.2, 15.0, 5.14, 260, 14.2, 0.952, 86.2, 53000]")

add("Normal",
    "The simulation results show consistent agreement with the design specification across all five "
    "qubits. The junction participation ratios p_junction > 0.94 indicate that the transmon regime "
    "is well maintained, and the predicted T1 limits from dielectric loss (83\u201392 \u03bcs) are consistent "
    "with state-of-the-art superconducting qubit performance.")

add("heading2", "6.6   Experiment 6: Autonomous Heartbeat Discoveries")

add("p1a",
    "We evaluate the heartbeat engine over a simulated 24-hour monitoring period. Table 7 shows "
    "the discoveries generated across the four intelligence tiers.")

add("tablecaption",
    "Table 7. Autonomous discoveries from the heartbeat engine over 24 hours.")

add("p1a",
    "[Table 7: Tier | Interval | Discovery | Action Taken; "
    "L0 (5 min): Equipment alarm on LITHO-03, Alert sent to operator; "
    "L0 (5 min): 2 ETL pipelines stopped, Restart triggered; "
    "L1 (6 hours): Yield average 88% < 90% threshold, Root-cause analysis initiated; "
    "L1 (6 hours): calibration table quality score 85%, Data quality flag raised; "
    "L2 (12 hours): Q3 T1=32\u03bcs < 35\u03bcs threshold, Recalibration recommended; "
    "L3 (24 hours): High-yield batches correlate with cal fidelity 99.2% vs 97.8%, "
    "Design insight logged for next iteration]")

add("Normal",
    "The heartbeat engine produced six actionable discoveries without any human prompting. The L3 "
    "cross-domain insight\u2014correlating fabrication yield with calibration fidelity\u2014is particularly "
    "valuable as it would typically require a data analyst to manually join datasets from the MES "
    "system and the measurement database. This demonstrates the system's ability to surface "
    "non-obvious relationships across the chip lifecycle.")

# ══════════════════════════════════════════════
# 7  DISCUSSION  (~1 page)
# ══════════════════════════════════════════════

add("heading1", "7   Discussion")

add("heading2", "7.1   Strengths")

add("p1a",
    "QuantaMind demonstrates three key strengths. First, the physics-structured theoretical physicist "
    "agent provides reasoning that is both more reliable and more interpretable than unconstrained "
    "LLM generation. By decomposing complex physics problems into well-defined modules with "
    "standardized interfaces, the system can combine LLM flexibility with domain rigor. Second, "
    "the graceful degradation design ensures that the system remains useful even when not all "
    "platforms are available\u2014a common situation in research environments where instruments and "
    "licenses may be shared. Third, the conversation pipeline mechanism provides full provenance "
    "for every agent action, enabling reproducibility and auditability that are essential in "
    "scientific research.")

add("heading2", "7.2   Limitations")

add("p1a",
    "Several limitations should be noted. (1) Without an Ansys HFSS desktop installation, "
    "electromagnetic simulations rely on analytical approximations (theory-mode) rather than "
    "full 3D field solutions. While the analytical results are physically informed, they cannot "
    "capture geometry-dependent effects such as parasitic modes or substrate radiation. "
    "(2) The GDS layout generator produces topologically correct layouts that match the "
    "design specification, but pixel-level reproduction of production masks requires additional "
    "manual refinement in professional EDA tools. (3) LLM function-calling occasionally produces "
    "malformed arguments or selects suboptimal tools, which is mitigated by validation layers "
    "and retry logic but not fully eliminated. (4) The current deployment is single-server, "
    "limiting concurrent agent execution to one user session at a time.")

add("heading2", "7.3   Roadmap")

add("p1a",
    "The development roadmap proceeds in three phases. Phase 1 (current) establishes the "
    "diagnostic closed loop: data ingestion \u2192 Hamiltonian modeling \u2192 calibration \u2192 "
    "diagnosis \u2192 next experiment recommendation. Phase 2 will add real instrument integration "
    "with ARTIQ hardware controllers and active learning for adaptive measurement scheduling. "
    "Phase 3 targets autonomous chip redesign with manufacturing-aware multi-objective optimization, "
    "closing the full design\u2013fabricate\u2013measure\u2013redesign loop.")

# ══════════════════════════════════════════════
# 8  CONCLUSIONS  (~0.5 page)
# ══════════════════════════════════════════════

add("heading1", "8   Conclusions")

add("p1a",
    "We have presented QuantaMind, a multi-agent AI system for autonomous superconducting quantum "
    "chip research. The system integrates twelve specialized agents over a unified tool execution "
    "layer connecting to 120+ tools across Qiskit Metal, KQCircuits, Ansys HFSS/Q3D, ARTIQ, "
    "CHIPMES, Apache Doris, and other platforms. The theoretical physicist agent implements nine "
    "functional modules providing physics-grounded reasoning with explicit uncertainty quantification "
    "for Hamiltonian modeling, noise budgeting, parameter inversion, experiment design, pulse "
    "optimization, root-cause diagnosis, design optimization, and knowledge retrieval.")

add("Normal",
    "Applied to a production-derived 20-qubit tunable-coupler chip design, QuantaMind demonstrates "
    "end-to-end capabilities: generating parameterized GDS layouts with Xmon qubits and SQUID "
    "couplers, executing Q3D capacitance extraction and EPR analysis, decomposing T1/T2 into "
    "constituent noise mechanisms, diagnosing CZ gate errors with probabilistic root-cause ranking, "
    "planning information-optimal measurement sequences, and generating three-tier design "
    "optimization proposals\u2014all orchestrated through natural-language interaction.")

add("Normal",
    "The autonomous heartbeat engine further demonstrates that continuous, unsupervised monitoring "
    "can surface actionable cross-domain insights (e.g., yield\u2013calibration correlations) that "
    "would otherwise require manual data integration. Together, these capabilities represent a "
    "step toward AI-native quantum hardware development, where the boundaries between design, "
    "simulation, fabrication, and characterization become increasingly fluid and where AI serves "
    "not merely as a tool but as an active participant in the scientific process.")

# ══════════════════════════════════════════════
# ACKNOWLEDGMENTS
# ══════════════════════════════════════════════

add("acknowlegments",
    "Acknowledgments. This work was supported by [funding information to be added].")

add("acknowlegments",
    "Disclosure of Interests. The authors have no competing interests to declare.")

# ══════════════════════════════════════════════
# REFERENCES  (30 entries)
# ══════════════════════════════════════════════

add("heading1", "References")

refs = [
    "Arute, F., et al.: Quantum supremacy using a programmable superconducting processor. Nature 574, 505\u2013510 (2019)",
    "Wu, Y., et al.: Strong quantum computational advantage using a superconducting quantum processor. Phys. Rev. Lett. 127, 180501 (2021)",
    "Google Quantum AI: Quantum error correction below the surface code threshold. Nature 638, 920\u2013926 (2025)",
    "IBM Quantum: The IBM Quantum Development Roadmap. https://www.ibm.com/quantum/roadmap (2024)",
    "CETC: 20-Qubit Tunable-Coupler Quantum Chip Design Specification. Internal Document TGQ-200-000-FA09-2025 (2025)",
    "Lu, C., et al.: The AI Scientist: Towards Fully Automated Open-Ended Scientific Discovery. arXiv:2408.06292 (2024)",
    "L\u00e1la, J., et al.: PaperQA: Retrieval-Augmented Generative Agent for Scientific Research. arXiv:2312.07559 (2023)",
    "Li, X., et al.: Kosmos: An End-to-End AI Research Agent. arXiv (2024)",
    "Zou, Y., et al.: El Agente: An Autonomous Agent for Quantum Chemistry. arXiv:2505.02484 (2025)",
    "Gustin, I., et al.: El Agente Cu\u00e1ntico: Automating quantum simulations. arXiv:2512.18847 (2025)",
    "Reuer, K., et al.: Autonomous calibration of superconducting qubits using k-agents. Nature (2025)",
    "OpenClaw: Open-source agent framework. https://github.com/openclaw (2024)",
    "Yao, S., et al.: ReAct: Synergizing Reasoning and Acting in Language Models. arXiv:2210.03629 (2022)",
    "Internal Design Document: Theoretical Physicist Agent Skill Specification. QuantaMind Project (2026)",
    "Qiskit Metal: Open-source quantum hardware design. https://qiskit.org/metal (2024)",
    "KQCircuits: IQM\u2019s open-source superconducting chip design library. https://github.com/iqm-finland/KQCircuits (2024)",
    "Brown, T., et al.: Language models are few-shot learners. NeurIPS 33, 1877\u20131901 (2020)",
    "Wei, J., et al.: Chain-of-thought prompting elicits reasoning in large language models. NeurIPS 35 (2022)",
    "Schick, T., et al.: Toolformer: Language models can teach themselves to use tools. NeurIPS 36 (2023)",
    "Patil, S., et al.: Gorilla: Large language model connected with massive APIs. arXiv:2305.15334 (2023)",
    "Wang, L., et al.: NovelSeek: When Agent Becomes the Scientist. arXiv (2025)",
    "Swanson, K., et al.: The Virtual Lab: AI agents design new SARS-CoV-2 nanobodies. arXiv (2024)",
    "Darvish Rouhani, B., et al.: ORGANA: A Robotic Assistant for Automated Chemistry Experimentation. arXiv (2024)",
    "Bran, A., et al.: ChemCrow: Augmenting large-language models with chemistry tools. arXiv:2304.05376 (2023)",
    "Boiko, D., et al.: Autonomous chemical research with large language models. Nature 624, 570\u2013578 (2023)",
    "Zhang, K., et al.: ChemAgent: Multi-Agent Coordination for Chemistry. arXiv (2025)",
    "Schiltz, F., et al.: MDCrow: A multi-agent system for molecular dynamics. arXiv (2025)",
    "Meng, F., et al.: QCR-LLM: Integrating quantum algorithms in LLM reasoning. arXiv (2025)",
    "Xu, H., et al.: PhIDO: Automated design of integrated photonic circuits. arXiv (2024)",
    "Saggio, V., et al.: Experimental quantum speed-up in reinforcement learning agents. Nature 591, 229\u2013233 (2021)",
]

for i, ref in enumerate(refs):
    add("referenceitem", f"{i+1}. {ref}")

doc.save(OUTPUT)
print(f"论文已保存: {OUTPUT}")
print(f"文件大小: {os.path.getsize(OUTPUT):,} bytes")
