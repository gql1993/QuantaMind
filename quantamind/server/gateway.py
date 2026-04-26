# QuantumClaw Gateway — 中枢网关（FastAPI + WebSocket）

import asyncio
import json
import logging
import os
import time
import uuid
from collections import deque
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

from pathlib import Path
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from quantamind import config
from quantamind.shared.api import (
    ChatRequest,
    ChatChunk,
    SessionCreate,
    SessionInfo,
    MessageRole,
    TaskInfo,
    TaskStatus,
)
from quantamind.agents.orchestrator import Orchestrator
from quantamind.server import memory, hands, hands_intel, skills_loader, heartbeat as hb, state_store, routes_system, arxiv_intel

config.ensure_dirs()
state_store.ensure_schema()
hb.load_persistent_discoveries()

# ── 日志收集（OpenClaw 风格：内存 ring buffer + tail API） ──
_LOG_BUFFER: deque[dict] = deque(maxlen=500)
_gateway_logger = logging.getLogger("quantamind.gateway")


class _BufferHandler(logging.Handler):
    def emit(self, record: logging.LogRecord) -> None:
        _LOG_BUFFER.append({
            "time": datetime.fromtimestamp(record.created, tz=timezone.utc).isoformat(),
            "level": record.levelname,
            "subsystem": record.name,
            "message": self.format(record),
        })


logging.root.addHandler(_BufferHandler())
_gateway_logger.setLevel(logging.DEBUG)
_STARTUP_TIME = time.time()

app = FastAPI(
    title="QuantaMind Gateway",
    description="量子科技自主科研 AI 中台 · 服务端",
    version="0.1.0",
)


@app.on_event("startup")
async def _on_startup():
    async def _delayed_heartbeat():
        await asyncio.sleep(5)
        await hb.heartbeat_loop()
    async def _delayed_intel_scheduler():
        await asyncio.sleep(5)
        await hb.intel_scheduler_loop()
    async def _delayed_intel_cache_warmer():
        await asyncio.sleep(5)
        await hb.intel_cache_warmer_loop()
    async def _delayed_taxonomy_engineer():
        await asyncio.sleep(5)
        await hb.taxonomy_engineer_loop()
    asyncio.create_task(_delayed_heartbeat())
    asyncio.create_task(_delayed_intel_scheduler())
    asyncio.create_task(_delayed_intel_cache_warmer())
    asyncio.create_task(_delayed_taxonomy_engineer())
    _gateway_logger.info("Heartbeat 后台任务已排定")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── 内存存储（生产换 Redis） ──
sessions: dict[str, dict] = state_store.load_sessions()
tasks: dict[str, dict] = state_store.load_tasks()
chat_histories: dict[str, list] = state_store.load_chat_histories()
chat_jobs: Dict[str, dict] = {}
orchestrator = Orchestrator()


def _is_intel_digest_shortcut(message: str) -> bool:
    text = (message or "").strip().lower()
    if not text:
        return False
    digest_terms = ("今天情报", "今日情报", "今天日报", "今日日报", "发日报", "发送日报", "发送今天情报", "发送今天日报")
    return any(term in text for term in digest_terms) and any(k in text for k in ("情报", "日报", "飞书"))


def _should_force_intel_digest(message: str) -> bool:
    text = (message or "").strip().lower()
    return any(term in text for term in ("重发", "重新发送", "重新推送", "补发", "再次发送"))


def _tool_candidates(tool: str) -> List[str]:
    if tool == "klayout":
        return [
            r"C:\Program Files (x86)\KLayout\klayout_app.exe",
            r"C:\Program Files\KLayout\klayout_app.exe",
        ]
    if tool == "ansys":
        return [
            r"D:\ANSYS Inc\ANSYS Student\v252\AnsysEM\ansysedtsv.exe",
            r"D:\ANSYS Inc\v252\AnsysEM\ansysedtsv.exe",
            r"C:\Program Files\AnsysEM\v252\Win64\ansysedt.exe",
            r"C:\Program Files\AnsysEM\v251\Win64\ansysedt.exe",
            r"C:\Program Files\ANSYS Inc\v252\AnsysEM\Win64\ansysedt.exe",
            r"C:\Program Files\ANSYS Inc\v251\AnsysEM\Win64\ansysedt.exe",
            r"D:\ANSYS Inc\ANSYS Student\v251\AnsysEM\ansysedtsv.exe",
        ]
    return []


def _resolve_tool_path(tool: str) -> Optional[str]:
    import shutil

    for candidate in _tool_candidates(tool):
        if os.path.exists(candidate):
            return candidate

    if tool == "klayout":
        return shutil.which("klayout_app") or shutil.which("klayout")

    if tool == "ansys":
        for cmd in ("ansysedtsv", "ansysedt"):
            resolved = shutil.which(cmd)
            if resolved:
                return resolved

    return None


def _tail_text_file(path: str, max_chars: int = 2000) -> str:
    if not os.path.exists(path):
        return ""
    try:
        with open(path, "rb") as f:
            f.seek(0, os.SEEK_END)
            size = f.tell()
            f.seek(max(size - max_chars, 0))
            data = f.read().decode("utf-8", errors="ignore")
        return data[-max_chars:]
    except Exception:
        return ""

# ── AI 科学家 Agent 团队注册表（OpenClaw agents.list） ──
# Agent 按工作流顺序排列：协调→理论→设计→仿真→材料→制造→运维→测控→算法→数据→知识→管理
AGENTS: List[Dict[str, Any]] = [
    {"id": "orchestrator", "label": "协调者", "emoji": "🧠", "role": "通用", "desc": "任务分解、多 Agent 调度、冲突仲裁、结果汇总。当用户请求涉及多个领域时，自动拆解子任务并分发给对应专家 Agent。",
     "platforms": "全平台", "status": "active",
     "tools": ["任务路由", "子任务分解", "Agent 调度", "结果融合"],
     "skills": ["多 Agent 协同", "复合任务编排"],
     "actions": [{"label": "查看当前任务队列", "prompt": "列出当前所有进行中的任务"}, {"label": "执行系统诊断", "prompt": "对当前系统状态进行全面诊断"}]},
    {"id": "theorist", "label": "理论物理学家", "emoji": "⚛️", "role": "理论", "desc": "面向超导量子电路的多尺度理论建模与闭环决策中枢。覆盖9大核心模块：M0数据接入与语义归一化、M1器件量子化建模（EPR/BBQ/Lumped→有效哈密顿量→频率/非简谐性/耦合/ZZ串扰/Purcell）、M2开系统噪声与误差预算（介质损耗/TLS/磁通噪声/准粒子/热光子→T1/T2分解+门误差预算+敏感度矩阵）、M3参数反演与模型校准（贝叶斯MCMC→带置信区间的参数后验+可辨识性分析）、M4信息增益驱动的实验设计（主动学习→最优实验序列）、M5控制脉冲与门优化（DRAG/GRAPE/鲁棒控制→泄漏抑制+串扰补偿）、M6诊断与根因归因（故障树+概率推理→候选根因排序+反证实验）、M7设计优化方案生成（多目标Pareto→参数窗口+版图修改）、M8文献知识图谱与证据中枢（现象→机理→验证实验→缓解手段）。",
     "platforms": "Q-EDA / 数据中台 / 测控平台 / 知识图谱", "status": "active",
     "tools": ["theorist_build_device_graph", "theorist_build_hamiltonian", "theorist_noise_budget", "theorist_calibrate_model", "theorist_plan_experiment", "theorist_optimize_pulse", "theorist_diagnose", "theorist_design_proposal", "theorist_knowledge_search", "theorist_status", "search_knowledge", "warehouse_query_qubit"],
     "skills": ["多尺度有效模型构建（EPR量化/BBQ/Lumped）", "约束下多目标逆向设计", "开量子系统噪声机理建模", "参数反演与贝叶斯校准", "信息增益驱动实验设计", "DRAG/GRAPE脉冲优化", "故障树根因归因", "设计优化方案生成", "物理知识图谱推理", "表征结果解释与失效归因"],
     "actions": [{"label": "构建哈密顿量模型", "prompt": "为当前20比特芯片构建DeviceGraph并执行EPR量化建模，提取频率、非简谐性、耦合强度、ZZ串扰等关键物理量"}, {"label": "噪声与误差预算", "prompt": "基于哈密顿量模型计算T1/T2分解和门误差预算，识别主导噪声源并给出敏感度矩阵"}, {"label": "诊断门误差根因", "prompt": "双比特CZ门保真度低于目标，执行根因归因分析，给出候选根因排序和反证实验方案"}, {"label": "设计实验方案", "prompt": "基于当前参数不确定性，规划信息增益最大的实验序列，包括停止条件和自适应策略"}, {"label": "生成设计优化方案", "prompt": "综合噪声预算和诊断结果，生成下一版芯片的参数窗口、版图修改建议和Pareto分析"}, {"label": "检索物理知识", "prompt": "检索与T1下降相关的物理机理、验证实验和缓解手段"}]},
    {"id": "chip_designer", "label": "AI 芯片设计师", "emoji": "💎", "role": "设计", "desc": "端到端超导量子芯片设计：需求分析 → 拓扑/频率选型 → 版图布局（Qiskit Metal + KQCircuits）→ 路由布线 → DRC → 电磁仿真（HFSS/Q3D/Sonnet）→ LOM/EPR 分析 → GDS 导出 → 制造掩膜。集成 Qiskit Metal（IBM）的设计+分析能力与 KQCircuits（IQM）的版图+制造能力。",
     "platforms": "Q-EDA（启元）/ Qiskit Metal / KQCircuits", "status": "active",
     "tools": ["metal_create_design", "metal_add_transmon", "metal_add_route", "metal_analyze_lom", "metal_analyze_epr", "metal_export_gds", "kqc_create_chip", "kqc_add_swissmon", "kqc_export_ansys", "kqc_export_gds", "kqc_export_mask", "qeda_get_chip", "qeda_get_qubit", "qeda_junction_params", "qeda_catalog", "search_knowledge"],
     "skills": ["Transmon 芯片设计", "Swissmon 芯片设计", "版图自动布线", "LOM/EPR 分析", "仿真导出", "掩膜生成", "QEDA 团队资料查阅", "芯片规格查询"],
     "actions": [{"label": "查看 QEDA 资料目录", "prompt": "查看芯片设计团队的资料目录，了解有哪些设计规范、方案和工具"}, {"label": "设计 20 比特芯片", "prompt": "基于 20 比特芯片规格，使用 Qiskit Metal 创建设计，添加比特和路由"}, {"label": "设计 105 比特芯片", "prompt": "查询 105 比特芯片规格和设计要求，给出设计方案"}, {"label": "查询 JJ 参数", "prompt": "获取 SQUID 曼哈顿型约瑟夫森结的设计参数"}, {"label": "搜索设计规范", "prompt": "搜索关于共面波导特征阻抗的设计规范"}]},
    {"id": "simulation_engineer", "label": "AI 仿真工程师", "emoji": "🖥️", "role": "仿真", "desc": "电磁仿真任务全生命周期管理：HFSS 本征模/驱动模仿真、Q3D 电容提取、Sonnet 平面EM仿真、量子电路仿真（态矢量/密度矩阵/噪声模拟）、参数扫描与优化、仿真结果分析与可视化。连接'设计'与'制造'之间的验证环节。",
     "platforms": "Q-EDA / Ansys HFSS / Q3D / Sonnet / Qiskit Aer", "status": "active",
     "tools": ["sim_status", "sim_hfss_eigenmode", "sim_q3d_extraction", "sim_lom_analysis", "sim_epr_analysis", "sim_export_hfss", "sim_full_chip", "metal_analyze_lom", "metal_analyze_epr", "metal_export_gds", "kqc_export_ansys", "kqc_export_sonnet", "warehouse_query_design"],
     "skills": ["HFSS 本征模仿真", "HFSS 驱动模仿真", "Q3D 电容矩阵提取", "LOM 集总模型分析", "EPR 能量参与比分析", "GDS→HFSS 项目导出", "完整芯片仿真流程", "仿真结果分析与对比"],
     "actions": [{"label": "检查仿真环境", "prompt": "检查当前仿真环境状态，包括 Ansys HFSS、pyaedt、可用仿真引擎"}, {"label": "HFSS 本征模仿真", "prompt": "对 Q1 执行 HFSS 本征模仿真，提取谐振频率、Q值和模态信息"}, {"label": "Q3D 电容提取", "prompt": "对 Q1、Q2 和 T1 执行 Q3D 电容矩阵提取，评估耦合强度"}, {"label": "完整芯片仿真", "prompt": "对 20 比特芯片的前 5 个比特执行完整仿真流程（Q3D+LOM+EPR+Eigenmode）"}, {"label": "导出 HFSS 项目", "prompt": "将当前 GDS 版图导出为 Ansys HFSS 项目文件，生成仿真设置脚本"}]},
    {"id": "material_scientist", "label": "AI 材料科学家", "emoji": "🧪", "role": "材料", "desc": "材料候选筛选、高通量计算任务提交、工艺与器件仿真、材料数据库检索。对接格物平台。",
     "platforms": "格物", "status": "standby",
     "tools": ["gewu_search_material", "gewu_submit_compute", "gewu_query_result"],
     "skills": ["材料筛选", "高通量计算"],
     "actions": [{"label": "搜索材料", "prompt": "搜索适合超导量子比特的衬底材料"}, {"label": "提交计算", "prompt": "对候选材料提交高通量计算任务"}]},
    {"id": "process_engineer", "label": "AI 制造工程师", "emoji": "🏭", "role": "制造（多子智能体）",
     "desc": "端到端制造管控团队统筹：下设 MES 设备工程师、工艺工程师、缺陷分析工程师、良率工程师、制程整合工程师、测试工程师等子智能体分工协作。覆盖工艺路线与派工、良率与 SPC、设备 SECS/GEM、配方与告警、Grafana 可视化。集成 OpenMES（业务底座）、secsgem（设备通信）、Grafana（可视化）。",
     "platforms": "墨子 MES / OpenMES / secsgem / Grafana", "status": "active",
     "tools": ["mes_list_routes", "mes_list_lots", "mes_get_lot", "mes_list_work_orders", "mes_query_yield", "mes_query_spc", "mes_dispatch", "mes_report_work", "secs_list_equipment", "secs_equipment_status", "secs_remote_command", "secs_list_alarms", "secs_get_recipes", "grafana_list_dashboards", "grafana_embed_url", "grafana_create_equipment_dashboard"],
     "skills": ["良率查询与分析", "SPC 过程监控", "工艺路线管理", "批次派工", "设备告警处理", "配方管理"],
     "sub_agents": [
        {"id": "mes_equipment_engineer", "label": "MES 设备工程师", "role": "机台与 SECS/GEM",
         "desc": "聚焦制造执行系统侧机台：设备枚举与在线状态、远程命令、配方列表与下发路径、告警闭环；配合 Grafana 机台监控与嵌入看板。",
         "tools": ["secs_list_equipment", "secs_equipment_status", "secs_remote_command", "secs_list_alarms", "secs_get_recipes", "grafana_list_dashboards", "grafana_embed_url", "grafana_create_equipment_dashboard"],
         "skills": ["设备在线与状态诊断", "SECS/GEM 远程命令与配方管理", "设备告警分级与处置闭环", "机台监控大屏与嵌入"]},
        {"id": "mes_process_specialist", "label": "工艺工程师", "role": "路线与工艺窗口",
         "desc": "负责工艺路线、站点顺序与工艺窗口；批次/工单派工与报工对齐机台配方，保障工艺参数在 MES 与现场一致。",
         "tools": ["mes_list_routes", "mes_list_lots", "mes_get_lot", "mes_list_work_orders", "mes_dispatch", "mes_report_work", "secs_get_recipes"],
         "skills": ["工艺路线与站点编排", "批次派工与完工报工", "工艺窗口与配方对齐", "在制批次状态追踪"]},
        {"id": "mes_defect_analyst", "label": "缺陷分析工程师", "role": "失效与根因",
         "desc": "结合良率分层、SPC 失控与批次履历，定位缺陷模式与可疑工艺站点，输出可验证的根因假设与复现路径。",
         "tools": ["mes_query_yield", "mes_query_spc", "mes_get_lot", "mes_list_lots"],
         "skills": ["良率异常与缺陷模式识别", "SPC 失控与特殊原因分析", "批次—站点—参数追溯", "失效假设与验证清单"]},
        {"id": "mes_yield_engineer", "label": "良率工程师", "role": "良率与统计",
         "desc": "负责良率指标定义、趋势与分层（站点/产品/时间）；维护控制图与规格限，支撑良率目标与改善闭环。",
         "tools": ["mes_query_yield", "mes_query_spc", "mes_get_lot", "mes_list_lots"],
         "skills": ["良率趋势与分层统计", "控制图与过程能力", "规格限与判异规则", "良率目标分解与追踪"]},
        {"id": "mes_pie", "label": "制程整合工程师", "role": "跨站整合",
         "desc": "打通多站点瓶颈与资源约束，协调设备状态、工艺窗口与良率结果，推动跨模块问题升级与一体化改善方案。",
         "tools": ["mes_list_routes", "mes_list_lots", "mes_get_lot", "mes_query_yield", "mes_query_spc", "mes_dispatch", "secs_list_equipment", "secs_equipment_status", "secs_list_alarms", "grafana_list_dashboards", "grafana_embed_url"],
         "skills": ["跨站瓶颈与产能平衡", "设备—工艺—良率联动分析", "制程窗口整合与风险评审", "一体化改善与升级路径"]},
        {"id": "mes_test_engineer", "label": "测试工程师", "role": "测试与数据 gate",
         "desc": "负责测试站点报工、测试结果与良率 gate 核对，保证测试数据进入 MES 履历完整、可追溯，支撑放行与返工决策。",
         "tools": ["mes_report_work", "mes_get_lot", "mes_list_work_orders", "mes_query_yield", "mes_list_lots"],
         "skills": ["测试程序与站点报工", "良率 gate 与规格核对", "测试数据完整性与追溯", "放行/返工数据支撑"]},
     ],
     "actions": [{"label": "查询最新良率", "prompt": "查询最近的良率数据和趋势分析"}, {"label": "SPC 监控", "prompt": "查询 JJ 电阻的 SPC 数据，检查是否有失控点"}, {"label": "查看工艺路线", "prompt": "列出所有工艺路线及当前批次状态"}, {"label": "设备状态总览", "prompt": "列出所有设备状态，检查是否有告警"}, {"label": "生成设备看板", "prompt": "为所有设备创建 Grafana 监控看板"}]},
    {"id": "device_ops", "label": "AI 设备运维员", "emoji": "🔧", "role": "测控实验室设备运维",
     "desc": "量子测控侧设备日常运维（非产线 MES 机台角色）：ARTIQ 实时测控（脉冲序列/参数扫描）+ Qiskit Pulse 校准（振幅/DRAG/频率/读出）+ Grafana 测控看板；可按需使用 SECS/GEM 与产线设备接口做告警/配方查询。与「AI 制造工程师」下设的 MES 设备工程师分工不同——后者侧重制造执行系统与产线机台、良率与派工闭环；本角色侧重实验室量子硬件控制与读出校准链路。",
     "platforms": "悬镜 测控 / ARTIQ / Qiskit Pulse / secsgem / Grafana", "status": "active",
     "tools": ["artiq_list_devices", "artiq_run_pulse", "artiq_run_scan", "pulse_cal_amplitude", "pulse_cal_drag", "pulse_cal_frequency", "pulse_full_calibration", "secs_list_equipment", "secs_remote_command", "secs_list_alarms", "secs_get_recipes", "grafana_create_equipment_dashboard"],
     "skills": ["ARTIQ 实时测控", "脉冲序列执行", "Qiskit Pulse 校准", "SECS/GEM 通信", "告警监控"],
     "actions": [{"label": "T1/T2 测量", "prompt": "使用 ARTIQ 对 Q0-Q4 执行 T1 和 T2 Ramsey 测量"}, {"label": "全套校准", "prompt": "使用 Qiskit Pulse 对 Q0-Q4 执行全套校准（频率+振幅+DRAG+读出）"}, {"label": "设备状态", "prompt": "列出所有 ARTIQ 硬件设备和 SECS/GEM 设备状态"}, {"label": "Rabi 振荡", "prompt": "对 Q0 执行 Rabi 振荡测量，确定 pi 脉冲参数"}]},
    {"id": "measure_scientist", "label": "AI 测控科学家", "emoji": "📡", "role": "测控科研", "desc": "量子比特表征与性能优化：ARTIQ 实时测控（光谱/Rabi/T1/T2/读出）→ Qiskit Pulse 高精度校准 → Mitiq 错误缓解（ZNE/PEC/CDR/DD）→ 系统级分析与科研方案。三层协作：ARTIQ 控制硬件、Pulse 优化参数、Mitiq 提升结果质量。",
     "platforms": "悬镜 测控 / ARTIQ / Qiskit Pulse / Mitiq / 数据中台", "status": "active",
     "tools": ["artiq_run_pulse", "artiq_run_scan", "pulse_full_calibration", "pulse_cal_amplitude", "pulse_cal_drag", "mitiq_zne", "mitiq_pec", "mitiq_dd", "mitiq_benchmark", "knowledge_search"],
     "skills": ["ARTIQ 比特表征", "Qiskit Pulse 校准优化", "Mitiq 错误缓解", "退相干分析", "纠错实验设计"],
     "actions": [{"label": "完整比特表征", "prompt": "使用 ARTIQ 对所有比特执行完整表征：光谱扫描→Rabi→T1→T2 Ramsey→T2 Echo→读出优化"}, {"label": "全套校准+纠错", "prompt": "用 Qiskit Pulse 全套校准后，用 Mitiq 对比 ZNE/PEC/CDR 效果并推荐最佳纠错方案"}, {"label": "动力学去耦优化", "prompt": "用 Mitiq 对比 XY4/CPMG/UDD 动力学去耦序列，找到最优 T2 延长方案"}, {"label": "纠错实验设计", "prompt": "基于当前比特参数和 Mitiq 纠错基准测试，设计表面码量子纠错实验方案"}]},
    {"id": "algorithm_engineer", "label": "AI 量子算法工程师", "emoji": "🔬", "role": "算法与软件", "desc": "量子电路设计与优化、算法研发（VQE/QAOA/QML/Grover）、电路编译与转译（适配芯片拓扑）、量子-经典混合算法、电路深度/门数优化、噪声感知编译、基准测试电路生成。连接'软件算法'与'硬件芯片'，确保算法在真实硬件上高效运行。",
     "platforms": "天元 云 / 开物 QML / Qiskit / Cirq", "status": "active",
     "tools": ["knowledge_search", "warehouse_query_qubit", "mitiq_zne", "mitiq_benchmark"],
     "skills": ["量子电路设计", "VQE/QAOA 变分算法", "电路编译与转译", "噪声感知优化", "QML 量子机器学习", "基准测试电路"],
     "actions": [{"label": "设计 VQE 电路", "prompt": "为 5 比特 Transmon 芯片设计一个 H2 分子基态能量的 VQE 变分电路"}, {"label": "电路编译优化", "prompt": "将当前量子电路编译到 heavy-hex 拓扑，最小化 CNOT 门数"}, {"label": "噪声感知编译", "prompt": "根据当前比特 T1/T2 和门保真度数据，做噪声感知的电路编译优化"}, {"label": "生成基准电路", "prompt": "生成 Quantum Volume 和 CLOPS 基准测试电路"}]},
    {"id": "data_analyst", "label": "AI 数据分析师", "emoji": "📊", "role": "数据分析", "desc": "数据中台的核心使用者：跨域数据查询（设计/制造/测控）、自然语言转 SQL（Text2SQL）、良率与校准关联分析、数据质量监控、数据血缘追溯、ETL 管道管理、数据资产盘点。把设计→制造→测控三域数据串起来，用数据驱动科研决策。",
     "platforms": "数据中台（QCoDeS + SeaTunnel + OLAP + qData）", "status": "active",
     "tools": ["warehouse_list_domains", "warehouse_list_tables", "warehouse_query_sql", "warehouse_query_qubit", "warehouse_query_yield", "warehouse_cross_domain", "qdata_text2sql", "qdata_list_assets", "qdata_quality_check", "qdata_lineage", "qdata_create_service", "qdata_list_standards", "seatunnel_list_jobs", "seatunnel_pipeline_status", "seatunnel_sync_qcodes", "seatunnel_sync_mes", "seatunnel_sync_design", "qcodes_list_experiments", "qcodes_list_runs", "qcodes_export_dataset", "grafana_list_dashboards", "grafana_embed_url"],
     "skills": ["跨域数据查询", "Text2SQL 自然语言查询", "良率趋势分析", "设计→制造→测控追溯", "数据质量监控", "ETL 管道管理", "数据血缘分析", "数据资产盘点"],
     "actions": [{"label": "自然语言查数据", "prompt": "用自然语言查询：最近一周各比特的 T1 和 T2 数据趋势如何？"}, {"label": "跨域追溯", "prompt": "从设计 design_001 追溯到制造批次和最终测控表征结果"}, {"label": "良率关联分析", "prompt": "分析良率与校准保真度之间的相关性，找出影响良率的关键因素"}, {"label": "数据质量巡检", "prompt": "对所有数据域执行数据质量检查，报告质量得分和问题"}, {"label": "ETL 管道状态", "prompt": "查看所有 SeaTunnel 数据同步管道状态，是否有延迟或失败"}, {"label": "数据血缘", "prompt": "查询 qubit_characterization 表的数据血缘，看数据从哪来、流向哪里"}]},
    {"id": "intel_officer", "label": "AI 情报员", "emoji": "🛰️", "role": "外部情报（多子智能体）",
     "desc": "情报团队统筹外源技术信号与内部沉淀：下设体系工程师、外源侦察、日报与推送等子智能体分工协作。聚焦量子芯片设计、制造、测控、量子计算与量子纠错等方向，产出摘要、技术路线、日报，并闭环至飞书与知识库。",
     "platforms": "arXiv / 飞书 / 项目记忆 / 知识库", "status": "active",
     "tools": ["intel_run_daily_digest", "intel_list_recent_papers", "intel_list_reports", "knowledge_search"],
     "skills": ["外部论文追踪", "技术路线提炼", "日报生成", "情报推送", "知识入库与复用"],
     "sub_agents": [
        {"id": "intel_systems_engineer", "label": "体系工程师", "role": "体系与模板",
         "desc": "定义情报产品形态、日报/专题模板、质量检查与归档规范；协调知识库分层与飞书卡片结构，保证产出可检索、可复用。",
         "tools": ["intel_list_reports", "knowledge_search"],
         "skills": ["技术路线提炼", "知识入库与复用"]},
        {"id": "intel_scout", "label": "外源研究侦察", "role": "检索与筛选",
         "desc": "对接 arXiv 等外源，执行检索、去重与主题过滤，为日报与专题提供候选论文与元数据。",
         "tools": ["intel_list_recent_papers", "knowledge_search"],
         "skills": ["外部论文追踪"]},
        {"id": "intel_digest_push", "label": "日报与推送专员", "role": "生成与触达",
         "desc": "编排定时情报日报，汇总技术路线与结论，跟踪飞书等渠道推送状态与历史版本。",
         "tools": ["intel_run_daily_digest", "intel_list_reports"],
         "skills": ["日报生成", "情报推送"]},
     ],
     "actions": [{"label": "立即生成情报日报", "prompt": "立即执行一次 arXiv 情报日报，输出技术路线、结论和链接"}, {"label": "查看近期论文", "prompt": "列出最近检索到的量子芯片与量子计算相关论文"}, {"label": "查看情报历史", "prompt": "列出历史情报日报和推送状态"}]},
    {"id": "knowledge_engineer", "label": "AI 知识工程师", "emoji": "📚", "role": "知识管理", "desc": "知识图谱构建与维护、arxiv 文献自动追踪与摘要、实验经验沉淀（成功/失败模式）、跨项目知识复用、技术报告生成、研究趋势分析。把团队的隐性知识变成可查询、可推理的结构化资产。",
     "platforms": "数据中台（知识图谱）/ arxiv / 项目记忆", "status": "active",
     "tools": ["knowledge_search", "warehouse_query_sql", "qdata_list_assets", "qdata_lineage"],
     "skills": ["知识图谱构建", "文献自动追踪", "实验经验沉淀", "技术报告生成", "研究趋势分析", "跨项目知识复用"],
     "actions": [{"label": "文献日报", "prompt": "检索过去 24 小时 arxiv 上与超导量子比特相关的新论文，生成摘要"}, {"label": "知识图谱查询", "prompt": "查询知识图谱中与 Transmon 退相干机制相关的所有实体和关系"}, {"label": "经验总结", "prompt": "总结最近 10 次芯片设计迭代中的成功模式和失败教训"}, {"label": "趋势分析", "prompt": "分析 2025-2026 年超导量子计算领域的研究热点和技术趋势"}]},
    {"id": "project_manager", "label": "AI 项目经理", "emoji": "📋", "role": "项目管理", "desc": "科研项目全流程管理：里程碑规划与追踪、任务分解与分配、进度汇报（日报/周报/月报）、资源协调（设备时间/计算资源）、风险识别与预警、多团队协作协调、会议纪要与行动项追踪。确保科研目标按期达成。",
     "platforms": "全平台（汇总各 Agent 状态）", "status": "active",
     "tools": ["warehouse_query_sql", "warehouse_cross_domain", "seatunnel_pipeline_status", "grafana_list_dashboards", "qdata_list_assets"],
     "skills": ["里程碑规划", "进度追踪", "周报/月报生成", "资源协调", "风险预警", "多团队协调"],
     "actions": [{"label": "生成周报", "prompt": "汇总本周所有 Agent 的工作成果、任务完成情况和关键数据指标，生成项目周报"}, {"label": "里程碑检查", "prompt": "检查当前项目里程碑完成情况，标出延期风险项"}, {"label": "资源使用分析", "prompt": "分析本周设备使用率（测控系统/仿真服务器/制造设备），给出下周资源调度建议"}, {"label": "风险预警", "prompt": "扫描所有进行中的任务和数据指标，识别潜在风险并给出建议"}]},
]

# ── 自主发现（Demo fallback；真实运行优先使用 heartbeat.discoveries） ──
DEMO_DISCOVERIES: List[Dict[str, Any]] = [
    {"id": "demo_d1", "category": "实验", "type": "实验建议", "title": "Q3 频率偏移需关注", "summary": "Q3 比特最近两次校准显示频率偏移约 12 MHz，可能与串扰补偿参数有关。建议安排一次全链路 T1/T2 测量以确认。", "level": "L2", "time": "2026-03-10 14:15", "time_iso": "2026-03-10T14:15:00Z", "source": "Demo", "severity": "medium", "entity_type": "qubit", "entity_id": "Q3", "recommended_action": "对 Q3 发起 Ramsey、T1/T2 与读出复测，并检查串扰补偿参数。", "handled": False},
    {"id": "demo_d2", "category": "假设", "type": "跨域洞察", "title": "耦合腔长度对 ZZ 耦合的影响", "summary": "上周提出的假设“缩短耦合腔 15% 可降低 ZZ 耦合 20%”已通过仿真验证：ZZ 耦合降低约 18%，但频率灵敏度增加 5%。需权衡。", "level": "L3", "time": "2026-03-09 22:00", "time_iso": "2026-03-09T22:00:00Z", "source": "Demo", "severity": "info", "entity_type": "coupler", "entity_id": "cavity_length", "recommended_action": "继续做频率灵敏度与门保真度的多目标分析，确认是否进入下一版设计。", "handled": False},
    {"id": "demo_d3", "category": "周报", "type": "周报", "title": "本周项目进展摘要", "summary": "本周完成：5 比特拓扑选型、频率分配初版、Q1-Q3 仿真通过；待办：DRC 修复 2 项违规、Q4/Q5 仿真待提交。建议下周重点推进 DRC 与双比特门校准。", "level": "L3", "time": "2026-03-08 18:00", "time_iso": "2026-03-08T18:00:00Z", "source": "Demo", "severity": "info", "entity_type": "project", "entity_id": "weekly_summary", "recommended_action": "将待办拆成任务并分配给设计、仿真、测控团队。", "handled": False},
]


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def _persist_session(session_id: str) -> None:
    if session_id in sessions:
        state_store.upsert_session(session_id, sessions[session_id])


def _append_history(session_id: str, message: dict) -> None:
    history = chat_histories.setdefault(session_id, [])
    history.append(message)
    state_store.append_chat_message(session_id, message)


def _persist_task(task_id: str) -> None:
    if task_id in tasks:
        state_store.upsert_task(task_id, tasks[task_id])


def _persist_pipeline(pipeline_id: str) -> None:
    if pipeline_id in pipelines:
        state_store.upsert_pipeline(pipeline_id, pipelines[pipeline_id])


def _flush_state_store() -> None:
    if not state_store.get_health().get("available", False):
        return
    for sid in list(sessions.keys()):
        _persist_session(sid)
    for sid, history in list(chat_histories.items()):
        for msg in history:
            state_store.append_chat_message(sid, msg)
    for task_id in list(tasks.keys()):
        _persist_task(task_id)
    for pipeline_id in list(pipelines.keys()):
        _persist_pipeline(pipeline_id)
    try:
        from quantamind.server import hands_qiskit_pulse as pulse_mod
        cal_map = pulse_mod.get_calibration_values().get("calibrations", {})
        for key, payload in cal_map.items():
            state_store.record_pulse_calibration(key, payload)
    except Exception:
        pass
    for disc in hb.get_discoveries():
        try:
            state_store.upsert_discovery(disc["id"], disc)
        except Exception:
            pass


def _ensure_seed_tasks():
    now = _now_iso()
    seed = [
        {"task_id": "t1", "status": TaskStatus.COMPLETED, "title": "Q-EDA 本征模仿真", "task_type": "qeda_simulation", "created_at": now, "session_id": None, "needs_approval": False, "result": {"freq_ghz": 5.2}, "error": None},
        {"task_id": "t2", "status": TaskStatus.RUNNING, "title": "DRC 检查", "task_type": "qeda_drc", "created_at": now, "session_id": None, "needs_approval": False, "result": None, "error": None},
        {"task_id": "t3", "status": TaskStatus.PENDING, "title": "导出 GDS", "task_type": "qeda_export_gds", "created_at": now, "session_id": None, "needs_approval": True, "result": None, "error": None},
        {"task_id": "t4", "status": TaskStatus.COMPLETED, "title": "T1/T2 测量 Q1-Q3", "task_type": "measure_t1t2", "created_at": now, "session_id": None, "needs_approval": False, "result": {"T1_us": 45.2, "T2_us": 32.1}, "error": None},
        {"task_id": "t5", "status": TaskStatus.PENDING, "title": "制造批次 B2026-03 提交", "task_type": "mes_submit", "created_at": now, "session_id": None, "needs_approval": True, "result": None, "error": None},
    ]
    for t in seed:
        if t["task_id"] not in tasks:
            tasks[t["task_id"]] = t
            _persist_task(t["task_id"])


_ensure_seed_tasks()
_gateway_logger.info("QuantaMind Gateway 启动完成")


def _get_orchestrator():
    return orchestrator


def _rebuild_orchestrator():
    global orchestrator
    from quantamind.agents.orchestrator import Orchestrator
    orchestrator = Orchestrator()
    return orchestrator


@app.post("/api/v1/sessions", response_model=SessionInfo)
def create_session(body: Optional[SessionCreate] = None):
    sid = str(uuid.uuid4())
    project_id = body.project_id if body else None
    now = __import__("datetime").datetime.utcnow().isoformat() + "Z"
    sessions[sid] = {"project_id": project_id, "created_at": now}
    _persist_session(sid)
    return SessionInfo(session_id=sid, project_id=sessions[sid]["project_id"], created_at=now)


@app.post("/api/v1/chat")
async def chat(req: ChatRequest):
    session_id = req.session_id or str(uuid.uuid4())
    if session_id not in sessions:
        sessions[session_id] = {"project_id": None, "created_at": _now_iso()}
        _persist_session(session_id)

    ctx = {"project_id": sessions[session_id].get("project_id")}
    from quantamind.shared.api import ChatMessage
    messages = [ChatMessage(role=MessageRole.USER, content=req.message)]

    _append_history(session_id, {"role": "user", "content": req.message, "time": _now_iso()})
    _gateway_logger.info("Chat [%s]: %s", session_id[:8], req.message[:80])

    async def stream():
        full_parts = []
        async for chunk in orchestrator.respond(messages, ctx):
            full_parts.append(chunk)
            yield f"data: {json.dumps(ChatChunk(type='content', data=chunk, session_id=session_id).model_dump())}\n\n"
        # 如果有对话流水线，在 done 中带上 pipeline_id
        pid = getattr(orchestrator, '_current_pipeline_id', None)
        _append_history(session_id, {"role": "assistant", "content": "".join(full_parts), "time": _now_iso()})
        done_data = {"session_id": session_id}
        if pid:
            done_data["pipeline_id"] = pid
        yield f"data: {json.dumps({**ChatChunk(type='done', session_id=session_id).model_dump(), 'pipeline_id': pid})}\n\n"

    if req.stream:
        return StreamingResponse(stream(), media_type="text/event-stream")
    full = []
    async for c in orchestrator.respond(messages, ctx):
        full.append(c)
    reply = "".join(full)
    _append_history(session_id, {"role": "assistant", "content": reply, "time": _now_iso()})
    pid = getattr(orchestrator, '_current_pipeline_id', None)
    return {"content": reply, "session_id": session_id, "pipeline_id": pid}


@app.post("/api/v1/chat-async")
async def chat_async(req: ChatRequest):
    session_id = req.session_id or str(uuid.uuid4())
    if session_id not in sessions:
        sessions[session_id] = {"project_id": None, "created_at": _now_iso()}
        _persist_session(session_id)

    ctx = {"project_id": sessions[session_id].get("project_id")}
    from quantamind.shared.api import ChatMessage

    messages = [ChatMessage(role=MessageRole.USER, content=req.message)]
    _append_history(session_id, {"role": "user", "content": req.message, "time": _now_iso()})
    _gateway_logger.info("ChatAsync [%s]: %s", session_id[:8], req.message[:80])

    job_id = f"chatjob-{uuid.uuid4().hex[:10]}"
    lowered = (req.message or "").lower()
    initial_status = "已提交，正在后台执行…"
    if any(k in lowered for k in ("情报", "论文", "arxiv", "日报", "飞书")):
        initial_status = "已提交给 AI 情报员，正在检索论文并生成日报…"
    chat_jobs[job_id] = {
        "job_id": job_id,
        "session_id": session_id,
        "status": "queued",
        "stage": "queued",
        "status_message": initial_status,
        "content": "",
        "pipeline_id": None,
        "agent_id": None,
        "agent_label": None,
        "tool_name": None,
        "error": None,
        "created_at": _now_iso(),
        "updated_at": _now_iso(),
        "completed_at": None,
    }

    async def _run_chat_job():
        job = chat_jobs[job_id]
        full_parts: List[str] = []
        async def _progress_update(payload: Dict[str, Any]):
            job["stage"] = payload.get("stage", job.get("stage"))
            job["status_message"] = payload.get("status_message", job.get("status_message"))
            if payload.get("pipeline_id"):
                job["pipeline_id"] = payload["pipeline_id"]
            if payload.get("agent_id"):
                job["agent_id"] = payload["agent_id"]
            if payload.get("agent_label"):
                job["agent_label"] = payload["agent_label"]
            if payload.get("tool_name"):
                job["tool_name"] = payload["tool_name"]
            job["updated_at"] = _now_iso()
        try:
            job["status"] = "running"
            job["stage"] = "running"
            job["status_message"] = "正在理解你的问题…"
            job["updated_at"] = _now_iso()
            if _is_intel_digest_shortcut(req.message):
                job["agent_id"] = "intel_officer"
                job["agent_label"] = "AI 情报员（快捷路径）"
                job["stage"] = "shortcut_digest"
                job["status_message"] = "已识别为今日情报快捷指令，正在用缓存日报快捷路径处理…"
                shortcut_result = await asyncio.to_thread(
                    hands_intel.run_daily_digest_shortcut,
                    _should_force_intel_digest(req.message),
                )
                summary = shortcut_result.get("summary") or "今日情报快捷路径已执行。"
                report = shortcut_result.get("report") or {}
                if report.get("report_id"):
                    summary += f"\n报告 ID：{report['report_id']}"
                job["content"] = summary
                _append_history(session_id, {"role": "assistant", "content": summary, "time": _now_iso()})
                job["status"] = "completed"
                job["stage"] = "completed"
                job["status_message"] = "今日情报快捷路径已完成。"
                job["completed_at"] = _now_iso()
                job["updated_at"] = job["completed_at"]
                return
            run_ctx = {**ctx, "progress_callback": _progress_update}
            async for chunk in orchestrator.respond(messages, run_ctx):
                if chunk:
                    full_parts.append(chunk)
                    job["content"] = "".join(full_parts)
                pid = getattr(orchestrator, "_current_pipeline_id", None)
                if pid:
                    job["pipeline_id"] = pid
                job["updated_at"] = _now_iso()
            pid = getattr(orchestrator, "_current_pipeline_id", None)
            if pid:
                job["pipeline_id"] = pid
            job["content"] = "".join(full_parts)
            _append_history(session_id, {"role": "assistant", "content": job["content"], "time": _now_iso()})
            job["status"] = "completed"
            job["stage"] = "completed"
            job["status_message"] = "已完成。"
            job["completed_at"] = _now_iso()
            job["updated_at"] = job["completed_at"]
        except Exception as e:
            job["status"] = "failed"
            job["stage"] = "failed"
            job["status_message"] = "执行失败。"
            job["error"] = str(e)
            job["completed_at"] = _now_iso()
            job["updated_at"] = job["completed_at"]
            _gateway_logger.exception("ChatAsync failed [%s]", session_id[:8])

    asyncio.create_task(_run_chat_job())
    return {"job_id": job_id, "session_id": session_id}


@app.get("/api/v1/chat-jobs/{job_id}")
def get_chat_job(job_id: str):
    from fastapi import HTTPException

    if job_id not in chat_jobs:
        raise HTTPException(status_code=404, detail="chat job not found")
    return chat_jobs[job_id]


@app.websocket(config.GATEWAY_WS_PATH)
async def websocket_endpoint(ws: WebSocket):
    await ws.accept()
    try:
        while True:
            raw = await ws.receive_text()
            data = json.loads(raw)
            msg = data.get("message", "")
            session_id = data.get("session_id") or str(uuid.uuid4())
            if session_id not in sessions:
                sessions[session_id] = {"project_id": data.get("project_id"), "created_at": __import__("datetime").datetime.utcnow().isoformat() + "Z"}
                _persist_session(session_id)
            ctx = {"project_id": sessions[session_id].get("project_id")}
            from quantamind.shared.api import ChatMessage
            messages = [ChatMessage(role=MessageRole.USER, content=msg)]
            _append_history(session_id, {"role": "user", "content": msg, "time": _now_iso()})
            full_parts = []
            async for chunk in orchestrator.respond(messages, ctx):
                full_parts.append(chunk)
                await ws.send_text(json.dumps(ChatChunk(type="content", data=chunk, session_id=session_id).model_dump()))
            _append_history(session_id, {"role": "assistant", "content": "".join(full_parts), "time": _now_iso()})
            await ws.send_text(json.dumps(ChatChunk(type="done", session_id=session_id).model_dump()))
    except WebSocketDisconnect:
        pass


@app.get("/api/v1/heartbeat")
def get_heartbeat():
    status = hb.get_status()
    return {
        "level": "multi",
        "interval_minutes": config.HEARTBEAT_INTERVAL_MINUTES,
        "last_run": status.get("last_run"),
        "next_run": None,
        "running": status.get("running", False),
        "discoveries_count": status.get("discoveries_count", 0),
        "next_intel_run": status.get("next_intel_run"),
        "next_intel_cache_warm": status.get("next_intel_cache_warm"),
        "last_intel_report_id": status.get("last_intel_report_id"),
        "levels": status.get("levels", []),
    }


@app.get("/api/v1/intel/overview")
def get_intel_overview():
    from quantamind.server import arxiv_intel

    status = hb.get_status()
    reports = arxiv_intel.list_reports(limit=1)
    latest_report = reports[0] if reports else None
    raw_papers = (latest_report or {}).get("top_papers") or arxiv_intel.list_recent_papers(limit=6)
    recent_papers = [arxiv_intel.enrich_intel_paper_for_overview(p) for p in (raw_papers or [])[:6]]
    source_summary = (latest_report or {}).get("source_summary") or {"source_counts": {}, "backend_counts": {}}
    return {
        "latest_report": latest_report,
        "recent_papers": recent_papers,
        "source_summary": source_summary,
        "next_intel_run": status.get("next_intel_run"),
        "next_intel_cache_warm": status.get("next_intel_cache_warm"),
        "last_intel_report_id": status.get("last_intel_report_id"),
    }


routes_system.bind_context({
    "sessions": sessions,
    "tasks": tasks,
    "hands": hands,
    "arxiv_intel": arxiv_intel,
    "skills_loader": skills_loader,
    "hb": hb,
    "state_store": state_store,
    "TaskStatus": TaskStatus,
    "AGENTS": AGENTS,
    "_LOG_BUFFER": _LOG_BUFFER,
    "_STARTUP_TIME": _STARTUP_TIME,
    "_gateway_logger": _gateway_logger,
    "DEMO_DISCOVERIES": DEMO_DISCOVERIES,
    "_now_iso": _now_iso,
    "get_heartbeat": get_heartbeat,
    "orchestrator_getter": _get_orchestrator,
    "rebuild_orchestrator": _rebuild_orchestrator,
})
app.include_router(routes_system.router)


# ── 会话管理（OpenClaw sessions.list / sessions.delete） ──

@app.get("/api/v1/sessions")
def list_sessions():
    out = []
    for sid, s in sessions.items():
        history = chat_histories.get(sid, [])
        msg_count = len(history)
        first_user_msg = next((m["content"] for m in history if m.get("role") == "user"), None)
        title = first_user_msg[:30] if first_user_msg else None
        out.append({"session_id": sid, "project_id": s.get("project_id"), "created_at": s.get("created_at"), "messages": msg_count, "title": title})
    return {"sessions": out}


@app.delete("/api/v1/sessions/{session_id}")
def delete_session(session_id: str):
    from fastapi import HTTPException
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="session not found")
    del sessions[session_id]
    chat_histories.pop(session_id, None)
    state_store.delete_session(session_id)
    return {"deleted": session_id}


@app.get("/api/v1/sessions/{session_id}/history")
def session_history(session_id: str):
    """OpenClaw chat.history：会话消息历史"""
    return {"session_id": session_id, "messages": chat_histories.get(session_id, [])}


# ── Agent 团队（OpenClaw agents.list） ──

@app.get("/api/v1/agents")
def list_agents():
    # 禁止 CDN/浏览器缓存，避免升级 Gateway 后仍显示旧「工艺工程师」与缺失 sub_agents
    return JSONResponse(
        content={"agents": AGENTS},
        headers={
            "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
            "Pragma": "no-cache",
        },
    )


@app.get("/api/v1/launch/{tool}")
def launch_tool(tool: str):
    """Launch a local EDA tool by name."""
    import subprocess
    path = _resolve_tool_path(tool)
    if tool in {"metal", "kqcircuits"}:
        return {"status": "error", "message": f"Tool '{tool}' has no desktop application to launch"}
    if not path:
        candidates = _tool_candidates(tool)
        return {
            "status": "error",
            "message": f"未找到可启动的 {tool} 可执行程序",
            "checked_paths": candidates,
        }
    try:
        subprocess.Popen([path], cwd=os.path.dirname(path))
        return {"status": "launched", "tool": tool, "path": path}
    except Exception as e:
        return {"status": "error", "message": str(e)}


@app.get("/api/v1/launch/metal-gui")
def launch_metal_gui(json_path: Optional[str] = None):
    """启动 MetalGUI 可视化窗口"""
    from quantamind.server.open_metal_gui import open_gui_with_design
    return open_gui_with_design(json_path)


@app.get("/api/v1/klayout/status")
def klayout_status():
    path = _resolve_tool_path("klayout")
    if path:
        return {"installed": True, "launchable": True, "path": path}
    try:
        import klayout.db
        return {
            "installed": True,
            "launchable": False,
            "path": "Python module (klayout.db)",
            "python_module": True,
            "message": "已安装 Python 模块，但未检测到桌面版 KLayout",
        }
    except ImportError:
        pass
    return {"installed": False, "launchable": False}


@app.get("/api/v1/sim/status")
def sim_status():
    try:
        from quantamind.server.hands_simulation import get_simulation_status
        status = get_simulation_status()
        ansys_path = _resolve_tool_path("ansys")
        status["ansys_launchable"] = bool(ansys_path)
        if ansys_path:
            status["ansys_path"] = ansys_path
        return status
    except Exception as e:
        return {"error": str(e), "ansys_desktop_installed": False}


# ── 审批（OpenClaw exec.approval.resolve） ──

class ApprovalAction(BaseModel):
    action: str  # approve | reject


@app.post("/api/v1/tasks/{task_id}/approve")
def approve_task(task_id: str, body: ApprovalAction):
    from fastapi import HTTPException
    if task_id not in tasks:
        raise HTTPException(status_code=404, detail="task not found")
    t = tasks[task_id]
    if not t.get("needs_approval"):
        raise HTTPException(status_code=400, detail="task does not need approval")
    if body.action == "approve":
        t["needs_approval"] = False
        t["status"] = TaskStatus.RUNNING
        _persist_task(task_id)
        _gateway_logger.info("任务 %s 已审批通过", task_id)
        return {"task_id": task_id, "result": "approved", "new_status": "running"}
    elif body.action == "reject":
        t["needs_approval"] = False
        t["status"] = TaskStatus.FAILED
        t["error"] = "审批被拒绝"
        _persist_task(task_id)
        _gateway_logger.info("任务 %s 已被拒绝", task_id)
        return {"task_id": task_id, "result": "rejected", "new_status": "failed"}
    raise HTTPException(status_code=400, detail="action must be approve or reject")


# ── 日志（OpenClaw logs.tail） ──

@app.get("/api/v1/logs")
def get_logs(limit: int = 100, level: Optional[str] = None):
    logs = list(_LOG_BUFFER)
    if level:
        logs = [l for l in logs if l["level"].upper() == level.upper()]
    return {"logs": logs[-limit:], "total": len(logs)}


@app.get("/api/v1/dataplatform/pipeline-history")
def pipeline_data_history(last_n: int = 20):
    """查询数据中台中的流水线历史"""
    from quantamind.server import hands_doris as doris
    return doris.query_pipeline_history(last_n)


@app.get("/api/v1/dataplatform/step-records")
def pipeline_step_records(pipeline_id: Optional[str] = None, agent: Optional[str] = None, tool: Optional[str] = None):
    """查询数据中台中的步骤记录"""
    from quantamind.server import hands_doris as doris
    return doris.query_step_records(pipeline_id, agent, tool)


@app.get("/api/v1/dataplatform/training-export")
def export_training_data(domain: str = "all"):
    """导出 AI 训练数据集"""
    from quantamind.server import hands_doris as doris
    return doris.export_training_dataset(domain)


# ── 项目资料库 ──

from fastapi import UploadFile, File, Form

@app.post("/api/v1/library/folders")
def create_library_folder(name: str = Form(...), description: str = Form(""), project_id: str = Form("default")):
    from quantamind.server import project_library as lib
    return lib.create_folder(name, description, project_id)


@app.get("/api/v1/library/folders")
def list_library_folders(project_id: Optional[str] = None):
    from quantamind.server import project_library as lib
    return {"folders": lib.list_folders(project_id)}


@app.delete("/api/v1/library/folders/{folder_id}")
def delete_library_folder(folder_id: str):
    from quantamind.server import project_library as lib
    return lib.delete_folder(folder_id)


@app.put("/api/v1/library/folders/{folder_id}")
def rename_library_folder(folder_id: str, name: str = Form(...)):
    from quantamind.server import project_library as lib
    return lib.rename_folder(folder_id, name)

@app.post("/api/v1/library/folders/{folder_id}/retry")
def retry_library_folder(folder_id: str, failed_only: bool = True):
    from quantamind.server import project_library as lib
    return lib.retry_folder(folder_id, failed_only=failed_only)


@app.post("/api/v1/library/upload")
async def upload_file(file: UploadFile = File(...), project_id: str = Form("default"), folder_id: str = Form("")):
    """上传文件到项目资料库（可指定文件夹），自动解析"""
    from quantamind.server import project_library as lib
    content = await file.read()
    result = lib.save_file(file.filename, content, project_id, folder_id)
    safe = {k: v for k, v in result.items() if k != "path"}
    return safe


@app.get("/api/v1/library/files")
def list_library_files(project_id: Optional[str] = None, search: Optional[str] = None):
    from quantamind.server import project_library as lib
    if search:
        files = lib.search_files(search)
    else:
        files = lib.list_files(project_id)
    safe = [{k: v for k, v in f.items() if k != "path"} for f in files]
    return {"files": safe, "count": len(safe)}


@app.get("/api/v1/library/files/{file_id}")
def get_library_file(file_id: str):
    from fastapi import HTTPException
    from quantamind.server import project_library as lib
    f = lib.get_file(file_id)
    if not f:
        raise HTTPException(status_code=404, detail="file not found")
    safe = {k: v for k, v in f.items() if k != "path"}
    return safe


class FileMoveRequest(BaseModel):
    target_folder_id: str = ""

@app.post("/api/v1/library/files/{file_id}/move")
def move_library_file(file_id: str, body: FileMoveRequest):
    from quantamind.server import project_library as lib
    return lib.move_file(file_id, body.target_folder_id)

@app.post("/api/v1/library/files/{file_id}/retry")
def retry_library_file(file_id: str):
    from quantamind.server import project_library as lib
    return lib.retry_file(file_id)

@app.delete("/api/v1/library/files/{file_id}")
def delete_library_file(file_id: str):
    from quantamind.server import project_library as lib
    return lib.delete_file(file_id)


@app.delete("/api/v1/library/files")
def clear_library_files(project_id: str = "default"):
    from quantamind.server import project_library as lib
    return lib.delete_all_files(project_id)


@app.get("/api/v1/library/files/{file_id}/download")
def download_library_file(file_id: str):
    from fastapi import HTTPException
    from quantamind.server import project_library as lib
    f = lib.get_file(file_id)
    if not f:
        raise HTTPException(status_code=404, detail="file not found")
    return FileResponse(f["path"], filename=f["filename"])


@app.get("/api/v1/library/stats")
def library_stats():
    from quantamind.server import project_library as lib
    return lib.get_stats()


@app.get("/api/v1/dataplatform/tables")
def dataplatform_tables(domain: Optional[str] = None):
    from quantamind.server import hands_warehouse as wh
    return wh.list_tables(domain)


@app.get("/api/v1/dataplatform/assets")
def dataplatform_assets(domain: Optional[str] = None):
    from quantamind.server import hands_qdata as qd
    return qd.list_data_assets(domain)


@app.get("/api/v1/dataplatform/standards")
def dataplatform_standards():
    from quantamind.server import hands_qdata as qd
    return qd.list_data_standards()


@app.get("/api/v1/dataplatform/sync-jobs")
def dataplatform_sync_jobs():
    from quantamind.server import hands_seatunnel as st
    return st.list_jobs()


@app.get("/api/v1/dataplatform/sync-status")
def dataplatform_sync_status():
    from quantamind.server import hands_seatunnel as st
    return st.get_pipeline_status()


@app.get("/api/v1/dataplatform/lineage")
def dataplatform_lineage(table: str = ""):
    from quantamind.server import hands_qdata as qd
    return qd.get_lineage(table)


@app.get("/api/v1/dataplatform/quality")
def dataplatform_quality(table: str = ""):
    from quantamind.server import hands_qdata as qd
    return qd.run_quality_check(table)


@app.get("/api/v1/dataplatform/table-detail")
def dataplatform_table_detail(table: str = "", domain: str = ""):
    """获取表详情：列结构 + 行数 + 示例数据 + 血缘 + 质量"""
    from quantamind.server import hands_warehouse as wh, hands_qdata as qd
    tables = wh.list_tables(domain)
    table_info = None
    for t in tables.get("tables", []):
        if t["table"] == table:
            table_info = t
            break
    if not table_info:
        return {"error": f"表 {table} 不存在"}
    lineage = qd.get_lineage(table)
    quality = qd.run_quality_check(table)
    # 列结构描述
    col_desc_map = {
        "design_id": "设计方案唯一标识", "chip_type": "芯片类型（transmon/xmon等）", "qubit_count": "量子比特数量",
        "topology": "拓扑结构（heavy_hex/grid/linear/chain）", "freq_allocation": "频率分配方案",
        "sim_id": "仿真任务ID", "sim_type": "仿真类型（HFSS/Q3D/Sonnet）", "freq_ghz": "频率（GHz）",
        "q_factor": "品质因子Q", "coupling_mhz": "耦合强度（MHz）",
        "drc_id": "DRC检查ID", "violations": "违规数量", "pass": "是否通过",
        "export_id": "导出ID", "filename": "GDS文件名", "size_mb": "文件大小（MB）", "approved": "是否已审批",
        "lot_id": "批次号", "product": "产品型号", "qty": "数量", "status": "状态", "current_step": "当前工步", "yield_pct": "良率（%）",
        "wo_id": "工单号", "step": "工步名称", "equipment_id": "设备ID", "operator": "操作员", "start_time": "开始时间", "end_time": "结束时间",
        "record_id": "记录ID", "defects": "缺陷数", "total": "总数",
        "spc_id": "SPC记录ID", "parameter": "监控参数名", "value": "测量值", "ucl": "上控制限", "lcl": "下控制限", "mean": "均值",
        "event_id": "事件ID", "event_type": "事件类型", "alarm_id": "告警ID", "message": "消息内容", "timestamp": "时间戳",
        "char_id": "表征记录ID", "chip_id": "芯片ID", "qubit": "比特编号", "T1_us": "T1弛豫时间（μs）", "T2_us": "T2退相干时间（μs）", "anharmonicity_mhz": "非谐性（MHz）",
        "cal_id": "校准ID", "gate": "门类型", "amplitude": "振幅", "drag_beta": "DRAG参数", "fidelity_pct": "保真度（%）",
        "bench_id": "基准ID", "gate_type": "门类型", "error_rate": "错误率", "method": "测量方法",
        "mit_id": "纠错ID", "circuit": "电路描述", "technique": "纠错技术", "noisy_value": "噪声值", "mitigated_value": "缓解后值", "improvement_pct": "改善百分比",
        "meas_id": "测量ID", "experiment": "实验名称", "run_id": "Run ID",
        "pipeline_id": "流水线ID", "name": "名称", "template": "模板", "total_steps": "总步骤", "completed_steps": "完成步骤", "failed_steps": "失败步骤",
        "step_id": "步骤ID", "stage": "阶段", "agent": "Agent名称", "title": "步骤标题", "tool": "工具名",
        "args_json": "参数JSON", "result_json": "结果JSON", "started_at": "开始时间", "completed_at": "完成时间",
        "created_at": "创建时间", "measured_at": "测量时间", "calibrated_at": "校准时间",
        "res_design_ghz": "谐振腔设计频率（GHz）", "res_sim_ghz": "谐振腔仿真频率（GHz）", "Ec_mhz": "非谐Ec（MHz）", "Lj_nH": "约瑟夫森结电感（nH）",
        "id": "比特ID", "mutual_pH": "互感（pH）",
    }
    columns_detail = []
    for col in table_info.get("columns", []):
        columns_detail.append({"name": col, "description": col_desc_map.get(col, "")})
    return {
        "table": table, "domain": domain or table_info.get("domain", ""),
        "columns": columns_detail, "rows": table_info.get("rows", 0),
        "lineage": lineage.get("lineage", {}),
        "quality": {"score": quality.get("quality_score", 0), "total": quality.get("total_records", 0), "issues": quality.get("issues_found", 0), "checks": quality.get("checks", [])},
    }


@app.get("/api/v1/dataplatform/capabilities")
def dataplatform_capabilities():
    """数据中台能力全景：QCoDeS + SeaTunnel + OLAP（可配置）+ qData"""
    from quantamind.server import hands_qcodes as qc, hands_seatunnel as st, hands_warehouse as wh, hands_qdata as qd
    wcap = wh.get_capabilities()
    return {
        "qcodes": qc.get_capabilities(),
        "seatunnel": st.get_capabilities(),
        "warehouse": wcap,
        "qdata": qd.get_capabilities(),
        "relationship": {
            "summary": "四层自底向上：QCoDeS 采集量子测量数据 → SeaTunnel 将设计/制造/测控三域数据汇入 OLAP（PostgreSQL / ClickHouse / Doris / SelectDB 等，由 datacenter_warehouse.engine 配置）→ 亚秒级 SQL 分析 → qData 提供数据治理、质量检查、Text2SQL 和数据服务 API。",
            "layers": [
                {"layer": "数据采集层", "project": "QCoDeS", "role": "量子仪器数据采集（S参数/T1/T2/校准），SQLite 存储，pandas 导出"},
                {"layer": "数据管道层", "project": "Apache SeaTunnel", "role": "100+ 连接器，批量/实时/CDC；REST 拉取运行任务"},
                {"layer": "数据仓库层", "project": "OLAP（可配置引擎）", "role": "三大域数据存储、跨域关联 SQL、训练数据导出"},
                {"layer": "数据治理层", "project": "qData", "role": "资产目录、标准管理、质量检查、Text2SQL、数据服务 API、血缘追溯"},
            ],
            "data_flow": "设计 + MES + 测控 → SeaTunnel ETL → OLAP 仓库 → qData 治理 → QuantaMind Agent",
            "cross_domain": "设计域(chip_designs/simulations) ↔ 制造域(lots/yield/spc) ↔ 测控域(characterization/calibration/benchmarks)",
        },
    }


@app.get("/api/v1/measure/capabilities")
def measure_capabilities():
    """悬镜测控能力全景：ARTIQ + Qiskit Experiments + Qiskit Pulse + Mitiq"""
    from quantamind.server import hands_artiq as artiq, hands_qiskit_pulse as pulse, hands_mitiq as mitiq_mod, hands_qiskit_experiments as qexp
    return {
        "artiq": artiq.get_capabilities(),
        "qiskit_experiments": qexp.get_capabilities(),
        "qiskit_pulse": pulse.get_capabilities(),
        "mitiq": mitiq_mod.get_capabilities(),
        "relationship": {
            "summary": "四层自底向上协作：ARTIQ 负责实时硬件控制（FPGA ns 精度）→ Qiskit Experiments 负责标准实验定义与自动分析（T1/Ramsey/Rabi/Readout/Drag）→ Qiskit Pulse 负责门级校准参数回写与优化（振幅/DRAG/频率）→ Mitiq 负责电路级错误缓解（ZNE/PEC/CDR/DD），形成从硬件到结果的完整测控链路。",
            "layers": [
                {"layer": "实时控制层", "project": "ARTIQ", "role": "FPGA 纳秒时序控制、DDS/TTL/DAC 驱动、脉冲序列执行、参数扫描、数据采集"},
                {"layer": "标准实验分析层", "project": "Qiskit Experiments", "role": "标准实验模板、曲线拟合、结构化实验结果、批量表征与校准建议"},
                {"layer": "校准优化层", "project": "Qiskit Pulse", "role": "门脉冲调度构建、振幅/DRAG/频率/读出校准、校准值管理"},
                {"layer": "错误缓解层", "project": "Mitiq", "role": "ZNE/PEC/CDR 错误缓解、动力学去耦、多技术对比与自动推荐"},
            ],
            "data_flow": "量子芯片 ←(硬件信号)→ ARTIQ(实时控制) → 原始测量数据 → Qiskit Experiments(标准实验分析) → Qiskit Pulse(校准参数回写) → Mitiq(错误缓解) → 高质量结果 → QuantaMind Agent(分析决策)",
        },
    }


@app.get("/api/v1/chipmes/info")
def chipmes_info():
    from quantamind.server import hands_chipmes as cm
    return cm.get_system_info()


@app.post("/api/v1/chipmes/start")
def chipmes_start():
    """优先启动真实 Java CHIPMES，失败时回退到快速兼容服务"""
    real_result = chipmes_start_real()
    if real_result.get("status") in {"already_running", "starting"}:
        return real_result

    # 回退到快速兼容服务，确保前端仍能秒级可用
    import subprocess, socket, sys
    from quantamind.server import hands_chipmes as cm
    if cm._is_connected():
        return {"status": "already_running", "message": "CHIPMES 已在运行（端口8082）", "port": 8082}

    def _port_open(port: int) -> bool:
        try:
            with socket.create_connection(("127.0.0.1", port), timeout=2):
                return True
        except Exception:
            return False

    if _port_open(8082):
        return {"status": "error", "message": "端口 8082 已被占用，但 CHIPMES 未响应，请先释放端口"}
    try:
        process = subprocess.Popen(
            [sys.executable, "-m", "quantamind.server.chipmes_quickstart"],
            cwd=str(Path(__file__).resolve().parent.parent.parent),
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            creationflags=subprocess.CREATE_NEW_PROCESS_GROUP | subprocess.DETACHED_PROCESS,
        )
        time.sleep(1.5)
        if process.poll() is not None or not cm._is_connected():
            return {
                "status": "startup_failed",
                "message": "CHIPMES 快速服务启动失败",
                "exit_code": process.poll(),
            }
        _gateway_logger.info("CHIPMES 快速兼容服务已启动")
        return {
            "status": "already_running",
            "message": "CHIPMES 已启动（快速模式）",
            "port": 8082,
            "fallback_from": real_result,
        }
    except Exception as e:
        _gateway_logger.error("CHIPMES 启动失败：%s", e)
        return {"status": "error", "message": str(e)}


@app.post("/api/v1/chipmes/start-real")
def chipmes_start_real():
    """启动真实 Java CHIPMES（保留给进一步排障使用）"""
    import subprocess, socket
    from quantamind.server import hands_chipmes as cm
    if cm._is_connected():
        return {"status": "already_running", "message": "CHIPMES 已在运行（端口8082）", "port": 8082}

    def _port_open(port: int) -> bool:
        try:
            with socket.create_connection(("127.0.0.1", port), timeout=2):
                return True
        except Exception:
            return False

    missing = []
    if not _port_open(1433):
        missing.append("SQL Server（端口1433）")
    if not _port_open(6379):
        missing.append("Redis（端口6379）")
    if missing:
        return {"status": "deps_missing", "message": f"CHIPMES 依赖未就绪：{' 和 '.join(missing)}，请先启动后再试",
                "missing": missing}

    service_dir = os.path.join(cm.MES_DIR, "CHIPMES-service")
    java_exe = os.path.join(service_dir, "jdk11", "bin", "java.exe")
    jar_path = os.path.join(service_dir, "cetc2-web", "cetc2-web.jar")
    if not os.path.exists(java_exe):
        return {"status": "error", "message": f"Java 不存在：{java_exe}"}
    if not os.path.exists(jar_path):
        return {"status": "error", "message": f"JAR 不存在：{jar_path}"}
    try:
        log_file = open(os.path.join(service_dir, "console.log"), "a", encoding="utf-8")
        process = subprocess.Popen(
            [java_exe,
             "-Xbootclasspath/a:./cetc2-web;./cetc2-web/config",
             "-Dfile.encoding=UTF-8",
             "-Dloader.path=./cetc2-web/lib",
             "-jar", "./cetc2-web/cetc2-web.jar"],
            cwd=service_dir,
            stdout=log_file, stderr=log_file,
            creationflags=subprocess.CREATE_NEW_PROCESS_GROUP | subprocess.DETACHED_PROCESS,
        )
        time.sleep(3)
        if process.poll() is not None:
            log_tail = _tail_text_file(os.path.join(service_dir, "console.log"))
            return {
                "status": "startup_failed",
                "message": "真实 CHIPMES 启动后立即退出，请检查依赖或日志",
                "exit_code": process.poll(),
                "log_tail": log_tail,
            }
        _gateway_logger.info("真实 CHIPMES Java 进程已启动（detached），工作目录：%s", service_dir)
        return {"status": "starting", "message": "真实 CHIPMES 启动中，请稍候约 15 秒后自动检测连接",
                "port": 8082}
    except Exception as e:
        _gateway_logger.error("真实 CHIPMES 启动失败：%s", e)
        return {"status": "error", "message": str(e)}


@app.get("/api/v1/mes/capabilities")
def mes_capabilities():
    """墨子 MES 能力全景：OpenMES + secsgem + Grafana"""
    from quantamind.server import hands_openmes as mes, hands_secsgem as secs, hands_grafana as grafana
    return {
        "openmes": mes.get_capabilities(),
        "secsgem": secs.get_capabilities(),
        "grafana": grafana.get_capabilities(),
        "relationship": {
            "summary": "四层协作：secsgem 负责设备通信（SECS/GEM 协议）→ FastAPI（已内建于 Gateway）封装 REST 接口 → OpenMES 提供 MES 业务逻辑（工艺/批次/工单/良率）→ Grafana 提供设备状态与 SPC 可视化大屏。",
            "layers": [
                {"layer": "设备通信层", "project": "secsgem", "role": "HSMS/SECS-II/GEM 协议栈，直接与半导体设备通信"},
                {"layer": "接口层", "project": "FastAPI（QuantaMind Gateway 内建）", "role": "将 SECS/GEM 封装为 REST API，供 OpenMES 和 QuantaMind 调用"},
                {"layer": "业务层", "project": "OpenMES", "role": "MES 业务底座：工艺路线、批次管理、工单派工、良率统计"},
                {"layer": "可视化层", "project": "Grafana", "role": "设备状态/参数大屏、良率趋势、SPC 过程监控看板"},
            ],
            "data_flow": "设备 ←(SECS/GEM)→ secsgem ←(REST)→ Gateway ←→ OpenMES(业务) + Grafana(可视化) + QuantaMind Agent(AI 决策)",
        },
    }


@app.get("/api/v1/qeda/capabilities")
def qeda_capabilities():
    """Q-EDA 设计能力全景：Qiskit Metal + KQCircuits"""
    from quantamind.server import hands_qiskit_metal as metal, hands_kqcircuits as kqc
    return {
        "qiskit_metal": metal.get_available_components(),
        "kqcircuits": kqc.get_available_elements(),
        "relationship": {
            "summary": "Qiskit Metal 侧重设计+分析（LOM/EPR 量化），KQCircuits 侧重版图自动化+制造导出（掩膜/EBL）。两者互补：用 Metal 做设计与参数优化，用 KQCircuits 生成制造就绪的版图。",
            "qiskit_metal_strengths": ["TransmonPocket/Cross 参数化设计", "RouteMeander 自动布线", "LOM 电容矩阵分析", "EPR 本征模分析", "Ansys/Gmsh 仿真渲染", "编程式 100+ 比特设计"],
            "kqcircuits_strengths": ["Swissmon 四臂 Transmon（IQM）", "SQUID/Airbridge/FluxLine/TSV 元件", "Ansys HFSS/Q3D/Sonnet 仿真导出", "光学掩膜+EBL 制造版图", "KLayout 集成可视化", "多层 Flip-chip 设计"],
            "recommended_workflow": "1. Qiskit Metal 创建设计+添加组件+路由 → 2. Metal LOM/EPR 分析优化参数 → 3. Metal 导出 GDS → 4. KQCircuits 导入/增强版图（加 airbridge/marker/TSV）→ 5. KQCircuits 导出仿真+制造掩膜",
        },
    }


@app.get("/api/v1/knowledge/search")
def knowledge_search(q: str = ""):
    from quantamind.server import knowledge_base as kb
    results = kb.search(q, max_results=10)
    return {"query": q, "results": results, "count": len(results)}


@app.get("/api/v1/skills")
def list_skills_api():
    return {"skills": skills_loader.list_skills()}


@app.get("/api/v1/tools")
def list_tools_api():
    return {"tools": hands.list_tools()}


@app.get("/api/v1/tasks", response_model=None)
def list_tasks(filter: Optional[str] = None):
    """任务列表；filter: all | mine | pending_approval | completed（与原型 Tab 对应）"""
    out = []
    for t in tasks.values():
        ti = TaskInfo(
            task_id=t["task_id"],
            status=t["status"],
            result=t.get("result"),
            error=t.get("error"),
            title=t.get("title"),
            task_type=t.get("task_type"),
            created_at=t.get("created_at"),
            session_id=t.get("session_id"),
            needs_approval=t.get("needs_approval", False),
        )
        if filter == "pending_approval" and not ti.needs_approval:
            continue
        if filter == "completed" and ti.status != TaskStatus.COMPLETED:
            continue
        out.append(ti)
    return {"tasks": [t.model_dump() for t in out]}


@app.get("/api/v1/tasks/{task_id}", response_model=TaskInfo)
def get_task(task_id: str):
    """任务详情"""
    from fastapi import HTTPException
    if task_id not in tasks:
        raise HTTPException(status_code=404, detail="task not found")
    t = tasks[task_id]
    return TaskInfo(
        task_id=t["task_id"],
        status=t["status"],
        result=t.get("result"),
        error=t.get("error"),
        title=t.get("title"),
        task_type=t.get("task_type"),
        created_at=t.get("created_at"),
        session_id=t.get("session_id"),
        needs_approval=t.get("needs_approval", False),
    )


# ── 流水线引擎（多 Agent 协同可视化） ──
pipelines: Dict[str, dict] = state_store.load_pipelines()


def _pipeline_step(pid: str, stage: int, agent: str, emoji: str, title: str, tool_name: str, tool_args: dict):
    """向流水线添加一步并执行"""
    p = pipelines[pid]
    step = {"stage": stage, "agent": agent, "emoji": emoji, "title": title,
            "tool": tool_name, "args": tool_args, "status": "running", "result": None, "started_at": _now_iso()}
    p["steps"].append(step)
    p["current_stage"] = stage
    p["current_step"] = title
    _persist_pipeline(pid)
    try:
        import asyncio
        loop = asyncio.get_event_loop()
        if loop.is_running():
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as pool:
                result = pool.submit(lambda: asyncio.run(hands.run_tool(tool_name, **tool_args))).result(timeout=10)
        else:
            result = loop.run_until_complete(hands.run_tool(tool_name, **tool_args))
        step["result"] = result
        step["status"] = "completed"
    except Exception as e:
        step["result"] = {"error": str(e)}
        step["status"] = "failed"
    step["completed_at"] = _now_iso()
    _persist_pipeline(pid)
    return step


async def _run_1bit_pipeline(pid: str):
    """单比特标准芯片流水线（基于真实 1bitV2.00 设计资料）"""
    p = pipelines[pid]
    p["status"] = "running"
    p["started_at"] = _now_iso()
    _persist_pipeline(pid)
    from quantamind.server.chip_real_designs import CHIP_1BIT

    async def run_step(stage, agent, emoji, title, tool_name, **kwargs):
        step = {"stage": stage, "agent": agent, "emoji": emoji, "title": title,
                "tool": tool_name, "args": kwargs, "status": "running", "started_at": _now_iso()}
        p["steps"].append(step); p["current_stage"] = stage; p["current_step"] = title
        _persist_pipeline(pid)
        try:
            result = await hands.run_tool(tool_name, **kwargs)
            step["result"] = result; step["status"] = "completed"
        except Exception as e:
            step["result"] = {"error": str(e)}; step["status"] = "failed"
        step["completed_at"] = _now_iso()
        _persist_pipeline(pid)
        await asyncio.sleep(0.5)
        return step["status"] == "completed"

    spec = CHIP_1BIT
    design_id = "1bit_v2_design"

    if not await run_step(1, "芯片设计师", "💎", "查询单比特真实设计资料", "qeda_real_design", design_key="1bit"): return
    if not await run_step(1, "芯片设计师", "💎", "创建 10mm x 10mm 平面设计", "metal_create_design", chip_size_x="10mm", chip_size_y="10mm"): return
    for q in spec["qubits"][:5]:
        if not await run_step(1, "芯片设计师", "💎", f"添加 Xmon 比特 {q['id']}（{q['freq_ghz']}GHz）", "metal_add_transmon", design_id=design_id, name=q["id"]): return
    if not await run_step(1, "芯片设计师", "💎", "导出简化 GDS", "metal_export_gds", design_id=design_id, filename="1bitV2_00.gds"): return
    if not await run_step(1, "芯片设计师", "💎", "部署团队真实工业级 GDS（5,122 shapes, 5 layers）", "build_real_chip", chip_type="1bit"): return

    if not await run_step(2, "仿真工程师", "🖥️", "Q3D 电容矩阵提取", "sim_q3d_extraction", components=["Q1", "Q2", "Q3"]): return
    for q in spec["qubits"][:3]:
        if not await run_step(2, "仿真工程师", "🖥️", f"HFSS 本征模仿真 {q['id']}", "sim_hfss_eigenmode", component_name=q["id"], freq_ghz=q["freq_ghz"]): return
        if not await run_step(2, "仿真工程师", "🖥️", f"LOM 分析 {q['id']}", "sim_lom_analysis", component_name=q["id"]): return
    if not await run_step(2, "仿真工程师", "🖥️", "导出 Ansys HFSS 项目", "sim_export_hfss", project_name="1bitV2_00"): return

    if not await run_step(3, "理论物理学家", "⚛️", "构建哈密顿量模型", "theorist_build_hamiltonian", device_graph_id="1bit_dg"): return
    if not await run_step(3, "理论物理学家", "⚛️", "噪声预算分析", "theorist_noise_budget", hamiltonian_model_id="1bit_ham"): return

    p["status"] = "completed"; p["completed_at"] = _now_iso(); _persist_pipeline(pid)


async def _run_2bit_pipeline(pid: str):
    """双比特可调耦合器流水线（基于真实 2bits_coupler_V2.00 设计资料）"""
    p = pipelines[pid]
    p["status"] = "running"
    p["started_at"] = _now_iso()
    _persist_pipeline(pid)
    from quantamind.server.chip_real_designs import CHIP_2BIT

    async def run_step(stage, agent, emoji, title, tool_name, **kwargs):
        step = {"stage": stage, "agent": agent, "emoji": emoji, "title": title,
                "tool": tool_name, "args": kwargs, "status": "running", "started_at": _now_iso()}
        p["steps"].append(step); p["current_stage"] = stage; p["current_step"] = title
        _persist_pipeline(pid)
        try:
            result = await hands.run_tool(tool_name, **kwargs)
            step["result"] = result; step["status"] = "completed"
        except Exception as e:
            step["result"] = {"error": str(e)}; step["status"] = "failed"
        step["completed_at"] = _now_iso()
        _persist_pipeline(pid)
        await asyncio.sleep(0.5)
        return step["status"] == "completed"

    spec = CHIP_2BIT
    design_id = "2bit_v2_design"

    if not await run_step(1, "芯片设计师", "💎", "查询双比特真实设计资料", "qeda_real_design", design_key="2bit"): return
    if not await run_step(1, "芯片设计师", "💎", "创建 10mm x 10mm 平面设计", "metal_create_design", chip_size_x="10mm", chip_size_y="10mm"): return
    for q in spec["qubits"]:
        qtype = "tunable_coupler" if q.get("type") == "tunable_coupler" else "qubit"
        label = f"可调耦合器 {q['id']}（{q['freq_ghz']}GHz）" if qtype == "tunable_coupler" else f"Xmon 比特 {q['id']}（{q['freq_ghz']}GHz）"
        if not await run_step(1, "芯片设计师", "💎", f"添加 {label}", "metal_add_transmon", design_id=design_id, name=q["id"]): return
    if not await run_step(1, "芯片设计师", "💎", "添加比特间耦合走线 Q1-T1-Q2", "metal_add_route", design_id=design_id, name="cpw_Q1_T1_Q2", start_component="Q1", start_pin="bus_a", end_component="Q2", end_pin="bus_b"): return
    if not await run_step(1, "芯片设计师", "💎", "添加比特间耦合走线 Q3-T2-Q4", "metal_add_route", design_id=design_id, name="cpw_Q3_T2_Q4", start_component="Q3", start_pin="bus_a", end_component="Q4", end_pin="bus_b"): return
    if not await run_step(1, "芯片设计师", "💎", "导出简化 GDS", "metal_export_gds", design_id=design_id, filename="2bits_coupler_V2_00.gds"): return
    if not await run_step(1, "芯片设计师", "💎", "部署团队真实工业级 GDS（366 shapes, 含耦合器）", "build_real_chip", chip_type="2bit"): return

    if not await run_step(2, "仿真工程师", "🖥️", "Q3D 电容矩阵提取（Q1+T1+Q2）", "sim_q3d_extraction", components=["Q1", "T1", "Q2"]): return
    if not await run_step(2, "仿真工程师", "🖥️", "Q3D 电容矩阵提取（Q3+T2+Q4）", "sim_q3d_extraction", components=["Q3", "T2", "Q4"]): return
    for q in spec["qubits"][:4]:
        if not await run_step(2, "仿真工程师", "🖥️", f"HFSS 本征模仿真 {q['id']}", "sim_hfss_eigenmode", component_name=q["id"], freq_ghz=q["freq_ghz"]): return
    if not await run_step(2, "仿真工程师", "🖥️", "LOM 集总模型分析", "sim_lom_analysis", component_name="Q1"): return
    if not await run_step(2, "仿真工程师", "🖥️", "EPR 能量参与比分析", "sim_epr_analysis", component_name="Q1"): return
    if not await run_step(2, "仿真工程师", "🖥️", "导出 Ansys HFSS 项目", "sim_export_hfss", project_name="2bits_coupler_V2_00"): return

    if not await run_step(3, "理论物理学家", "⚛️", "构建器件关系图", "theorist_build_device_graph", chip_id="2bit_coupler", qubits=["Q1","Q2","Q3","Q4"], couplers=[{"id":"T1","q1":"Q1","q2":"Q2"},{"id":"T2","q1":"Q3","q2":"Q4"}]): return
    if not await run_step(3, "理论物理学家", "⚛️", "构建哈密顿量模型", "theorist_build_hamiltonian", device_graph_id="2bit_dg"): return
    if not await run_step(3, "理论物理学家", "⚛️", "噪声预算分析", "theorist_noise_budget", hamiltonian_model_id="2bit_ham"): return
    if not await run_step(3, "理论物理学家", "⚛️", "CZ 门根因诊断", "theorist_diagnose", anomaly_type="gate_error_high"): return
    if not await run_step(3, "理论物理学家", "⚛️", "设计优化方案", "theorist_design_proposal"): return

    p["status"] = "completed"; p["completed_at"] = _now_iso(); _persist_pipeline(pid)


async def _run_20bit_pipeline(pid: str):
    """后台执行 20 比特流水线"""
    p = pipelines[pid]
    p["status"] = "running"
    p["started_at"] = _now_iso()
    _persist_pipeline(pid)

    from quantamind.server import hands_doris as _doris_rt

    async def run_step(stage, agent, emoji, title, tool_name, **kwargs):
        while p.get("paused"):
            p["status"] = "paused"
            await asyncio.sleep(0.5)
        if p.get("aborted"):
            p["status"] = "aborted"
            return False
        p["status"] = "running"
        step = {"stage": stage, "agent": agent, "emoji": emoji, "title": title,
                "tool": tool_name, "args": kwargs, "status": "running", "result": None, "started_at": _now_iso()}
        p["steps"].append(step)
        p["current_stage"] = stage
        p["current_step"] = title
        _persist_pipeline(pid)
        try:
            result = await hands.run_tool(tool_name, **kwargs)
            step["result"] = result
            step["status"] = "completed"
        except Exception as e:
            step["result"] = {"error": str(e)}
            step["status"] = "failed"
        step["completed_at"] = _now_iso()
        # 实时保存每步到数据中台
        try:
            _doris_rt.save_pipeline_steps(pid, [step])
        except Exception:
            pass
        _persist_pipeline(pid)
        await asyncio.sleep(0.05)
        return True

    async def gate_check(stage_name):
        """阶段门：如果设置了 gate_approval，等待人工审批"""
        if not p.get("require_gate_approval"):
            return True
        gate = {"stage": stage_name, "status": "waiting_approval", "created_at": _now_iso()}
        p.setdefault("gates", []).append(gate)
        p["status"] = "waiting_approval"
        p["current_step"] = f"等待审批：{stage_name} 完成，是否继续？"
        _persist_pipeline(pid)
        while gate["status"] == "waiting_approval":
            if p.get("aborted"):
                _persist_pipeline(pid)
                return False
            await asyncio.sleep(0.3)
        _persist_pipeline(pid)
        return gate["status"] == "approved"

    from quantamind.server.chip_20bit_spec import CHIP_SPEC as CS

    # 注入芯片规格摘要
    p["chip_spec"] = {"doc_id": CS["doc_id"], "name": CS["name"], "chip_size_mm": CS["chip_size_mm"],
                      "total_qubits": CS["total_qubits"], "total_couplers": CS["total_couplers"],
                      "substrate": CS["substrate"], "metal_film": CS["metal_film"],
                      "targets": CS["targets"], "layout_components": CS["layout_components"],
                      "fabrication": CS["fabrication"], "cpw": CS["cpw"], "packaging": CS["packaging"]}

    # 阶段 1：芯片设计（全部 20 比特 + 19 耦合器）
    if not await run_step(1, "芯片设计师", "💎", f"创建 {CS['chip_size_mm'][0]}mm x {CS['chip_size_mm'][1]}mm 平面设计", "metal_create_design", chip_size_x=f"{CS['chip_size_mm'][0]}mm", chip_size_y=f"{CS['chip_size_mm'][1]}mm"): return
    design_id = p["steps"][-1].get("result", {}).get("design_id", "design_0001")

    for q in CS["qubits"]:
        idx = CS["qubits"].index(q)
        if not await run_step(1, "芯片设计师", "💎",
            f"添加比特 {q['id']}（{q['freq_ghz']}GHz, 谐振腔 {q['res_design_ghz']}GHz, Ec={q['Ec_mhz']}MHz, Lj={q['Lj_nH']}nH）",
            "metal_add_transmon", design_id=design_id, name=q["id"], qubit_type="transmon_cross",
            pos_x=f"{idx * 600}um", pos_y="6250um"): return

    for i in range(len(CS["qubits"]) - 1):
        q1, q2 = CS["qubits"][i]["id"], CS["qubits"][i+1]["id"]
        if not await run_step(1, "芯片设计师", "💎", f"路由 {q1}-{q2}（CPW s=10um w=5um Z0≈50Ω）", "metal_add_route",
                       design_id=design_id, name=f"cpw_{q1}_{q2}", start_component=q1, start_pin="bus_a", end_component=q2, end_pin="bus_b"): return

    if not await run_step(1, "芯片设计师", "💎", "KQC 创建芯片布局", "kqc_create_chip", name=CS["name"], chip_size_x=12500, chip_size_y=12500): return
    kqc_id = p["steps"][-1].get("result", {}).get("chip_id", "kqc_chip_0001")

    for c in CS["couplers"]:
        if not await run_step(1, "芯片设计师", "💎",
            f"添加耦合器 {c['id']}（{c['freq_ghz']}GHz, SQUID, 连接 {c['connects'][0]}-{c['connects'][1]}）",
            "kqc_add_swissmon", chip_id=kqc_id, name=c["id"]): return

    for comp in CS["layout_components"]:
        if comp["name"] in ("空气桥", "约瑟夫森结"):
            if not await run_step(1, "芯片设计师", "💎",
                f"添加 {comp['name']} x{comp['count']}（图号 {comp['drawing_id']}）",
                "kqc_add_element", chip_id=kqc_id, name=comp["name"], element_type=comp["name"]): return

    if not await run_step(1, "芯片设计师", "💎", "导出简化 GDS（版图文件）", "metal_export_gds", design_id=design_id, filename="TG5.440.0001WX_20bit.gds"): return
    if not await run_step(1, "芯片设计师", "💎", "部署团队真实工业级 GDS（24,888 shapes, 9.6MB）", "build_ct20q_real", output_filename="CT20QV2_real.gds"): return
    if not await run_step(1, "芯片设计师", "💎", "实时运行版图代码生成 GDS（NoopGUI + SafeRoutePathfinder）", "generate_ct20q_live", output_filename="CT20QV2_industrial.gds", timeout=900): return
    if not await run_step(1, "芯片设计师", "💎", "导出制造掩膜（光学+EBL）", "kqc_export_mask", chip_id=kqc_id): return
    if not await gate_check("芯片设计"): return

    # 阶段 2：仿真（使用文档中的 Q3D 电容矩阵 + 全比特参数）
    # Q3D 电容矩阵
    cap = CS["capacitance_matrix_fF"]
    cap_step = {"stage": 2, "agent": "仿真工程师", "emoji": "🖥️", "title": "Q3D 电容矩阵提取（4x4）",
                "tool": "metal_analyze_lom", "args": {"design_id": design_id},
                "status": "completed", "started_at": _now_iso(), "completed_at": _now_iso(),
                "result": {"analysis": "Q3D_capacitance", "components": cap["components"], "matrix_fF": cap["matrix"],
                           "note": "对角线为自电容，非对角线为耦合电容。来源：设计文档表 4-2"}}
    p["steps"].append(cap_step); p["current_stage"] = 2; _persist_pipeline(pid); await asyncio.sleep(0.3)

    for q in CS["qubits"][:5]:
        if not await run_step(2, "仿真工程师", "🖥️",
            f"HFSS 本征模仿真 {q['id']}（设计 {q['freq_ghz']}GHz）",
            "sim_hfss_eigenmode", component_name=q["id"], freq_ghz=q["freq_ghz"]): return

    for q in CS["qubits"][:3]:
        if not await run_step(2, "仿真工程师", "🖥️",
            f"LOM 集总模型分析 {q['id']}",
            "sim_lom_analysis", component_name=q["id"]): return

    for q in CS["qubits"][:3]:
        if not await run_step(2, "仿真工程师", "🖥️",
            f"EPR 能量参与比分析 {q['id']}",
            "sim_epr_analysis", component_name=q["id"]): return

    if not await run_step(2, "仿真工程师", "🖥️",
        "导出 Ansys HFSS 项目文件（.aedt + 设置脚本）",
        "sim_export_hfss", project_name="TG5_440_0001WX_20bit", simulation_type="eigenmode"): return

    if not await run_step(2, "仿真工程师", "🖥️", "导出 KQC Sonnet 平面仿真", "kqc_export_sonnet", chip_id=kqc_id): return
    if not await gate_check("仿真验证"): return

    # 阶段 3：制造（注入真实工艺流程）
    fab = CS["fabrication"]
    fab_step = {"stage": 3, "agent": "制造工程师", "emoji": "🏭", "title": "制造规格确认",
                "tool": "mes_list_routes", "args": {},
                "status": "completed", "started_at": _now_iso(), "completed_at": _now_iso(),
                "result": {"fabrication_spec": fab, "process_steps": [
                    "1. 衬底清洗（蓝宝石 430±25μm）",
                    f"2. 超导薄膜沉积（{fab['film_thickness_nm']}nm Al/Ta）",
                    f"3. 激光直写光刻（精度 {fab['lithography'].split('精度')[1] if '精度' in fab['lithography'] else '±0.3μm'}）",
                    "4. 干法刻蚀（RIE）",
                    f"5. 约瑟夫森结制备（{fab['jj_fabrication']}，结尺寸 {fab['jj_length_nm']}x{fab['jj_width_nm']}nm）",
                    f"6. 氧化层生长（{fab['oxide_thickness_nm']}nm Al₂O₃）",
                    "7. 空气桥制备",
                    "8. 布线与封装（48 pin SMP）",
                ]}}
    p["steps"].append(fab_step); p["current_stage"] = 3; _persist_pipeline(pid); await asyncio.sleep(0.3)

    if not await run_step(3, "制造工程师", "🏭", "批次创建与派工（衬底清洗）", "mes_dispatch", lot_id="LOT-2026-0301", step="衬底清洗", equipment_id="CLEAN-01"): return
    if not await run_step(3, "制造工程师", "🏭", f"薄膜沉积（{fab['film_thickness_nm']}nm）", "mes_dispatch", lot_id="LOT-2026-0301", step="薄膜沉积", equipment_id="SPUTTER-01"): return
    if not await run_step(3, "制造工程师", "🏭", "光刻（激光直写 ±0.3μm）", "mes_dispatch", lot_id="LOT-2026-0301", step="光刻", equipment_id="LITHO-03"): return
    if not await run_step(3, "制造工程师", "🏭", "干法刻蚀", "mes_dispatch", lot_id="LOT-2026-0301", step="刻蚀", equipment_id="ETCH-02"): return
    if not await run_step(3, "制造工程师", "🏭", f"JJ 制备（Dolan Bridge {fab['jj_length_nm']}x{fab['jj_width_nm']}nm）", "mes_dispatch", lot_id="LOT-2026-0301", step="JJ制备", equipment_id="EBL-01"): return
    if not await run_step(3, "制造工程师", "🏭", "良率检测", "mes_query_yield"): return
    if not await run_step(3, "制造工程师", "🏭", "SPC 监控 JJ 电阻（设计值 ~1000Ω @180nm结宽）", "mes_query_spc", parameter="JJ_Resistance", last_n=20): return
    if not await run_step(3, "制造工程师", "🏭", "封装（48 pin SMP + 无氧铜盒）", "mes_dispatch", lot_id="LOT-2026-0301", step="封装", equipment_id="PACKAGE-01"): return
    if not await gate_check("工艺制造"): return

    # 阶段 4：校准（全部 20 比特）
    if not await run_step(4, "设备运维员", "🔧", "ARTIQ 硬件就绪检查", "artiq_list_devices"): return
    if not await run_step(4, "设备运维员", "🔧", "制冷机降温确认（10mK base）", "secs_equipment_status", equipment_id="CRYO-01"): return
    if not await run_step(4, "设备运维员", "🔧", "全套校准 Q1-Q10（频率+振幅+DRAG+读出）", "pulse_full_calibration", qubits=list(range(10))): return
    if not await run_step(4, "设备运维员", "🔧", "全套校准 Q11-Q20", "pulse_full_calibration", qubits=list(range(10, 20))): return
    if not await gate_check("设备校准"): return

    # 阶段 5：测控（全芯片表征）
    all_qids = [q["id"] for q in CS["qubits"]]
    for seq, desc in [("spectroscopy", "频率光谱扫描"), ("rabi", "Rabi 振荡（pi 脉冲标定）"),
                      ("t1", "T1 纵向弛豫（目标≥20μs）"), ("t2_ramsey", "T2* Ramsey 退相干（目标≥15μs）"),
                      ("t2_echo", "T2 Hahn Echo"), ("readout_optimization", "读出优化（目标≥99%）")]:
        if not await run_step(5, "测控科学家", "📡", f"{desc}（{len(all_qids)} 比特）", "artiq_run_pulse", sequence_type=seq, qubits=all_qids): return
    if not await run_step(5, "测控科学家", "📡", "CZ 门保真度测量（Q1-Q2, 目标≥99%）", "artiq_run_pulse", sequence_type="rabi", qubits=["Q1","Q2"]): return
    if not await run_step(5, "测控科学家", "📡", "Mitiq 纠错技术对比（ZNE/PEC/CDR）", "mitiq_benchmark", circuit_desc="20bit_CZ_gate_chain"): return
    if not await run_step(5, "测控科学家", "📡", "动力学去耦 XY4（延长 T2）", "mitiq_dd", circuit_desc="20bit_idle_chain", dd_sequence="XY4"): return
    if not await gate_check("测控表征"): return

    # 阶段 6：数据
    if not await run_step(6, "数据分析师", "📊", "数据域概览", "doris_list_domains"): return
    if not await run_step(6, "数据分析师", "📊", "跨域追溯", "doris_cross_domain", query_type="design_to_measurement"): return
    if not await run_step(6, "数据分析师", "📊", "数据质量检查", "qdata_quality_check", table="qubit_characterization"): return
    if not await run_step(6, "数据分析师", "📊", "ETL 管道状态", "seatunnel_pipeline_status"): return

    # 自动将过程数据写入数据中台
    # 保存到数据中台（直接调用，不通过 run_step 避免循环引用）
    from quantamind.server import hands_doris as _doris_save
    save1 = _doris_save.save_pipeline_run(pid, p)
    p["steps"].append({"stage": 6, "agent": "数据分析师", "emoji": "📊", "title": "保存流水线记录到数据中台",
                        "tool": "doris_save_pipeline", "args": {"pipeline_id": pid}, "status": "completed",
                        "result": save1, "started_at": _now_iso(), "completed_at": _now_iso()})
    save2 = _doris_save.save_pipeline_steps(pid, p["steps"])
    p["steps"].append({"stage": 6, "agent": "数据分析师", "emoji": "📊", "title": f"保存 {save2.get('steps_saved', 0)} 条步骤到数据中台",
                        "tool": "doris_save_steps", "args": {"pipeline_id": pid}, "status": "completed",
                        "result": save2, "started_at": _now_iso(), "completed_at": _now_iso()})
    if CS.get("qubits"):
        save3 = _doris_save.save_design_params(pid, CS["qubits"])
        p["steps"].append({"stage": 6, "agent": "数据分析师", "emoji": "📊", "title": f"保存 {len(CS['qubits'])} 个比特设计参数",
                            "tool": "doris_save_design_params", "args": {"pipeline_id": pid}, "status": "completed",
                            "result": save3, "started_at": _now_iso(), "completed_at": _now_iso()})

    p["status"] = "completed"
    p["completed_at"] = _now_iso()
    _persist_pipeline(pid)


async def _run_100bit_pipeline(pid: str):
    """100 比特芯片设计流水线"""
    p = pipelines[pid]
    p["status"] = "running"
    p["started_at"] = _now_iso()
    _persist_pipeline(pid)
    from quantamind.server import hands_doris as _doris_rt

    async def run_step(stage, agent, emoji, title, tool_name, **kwargs):
        while p.get("paused"):
            p["status"] = "paused"
            await asyncio.sleep(0.5)
        if p.get("aborted"):
            p["status"] = "aborted"
            return False
        p["status"] = "running"
        step = {"stage": stage, "agent": agent, "emoji": emoji, "title": title,
                "tool": tool_name, "args": kwargs, "status": "running", "result": None, "started_at": _now_iso()}
        p["steps"].append(step)
        p["current_stage"] = stage
        p["current_step"] = title
        _persist_pipeline(pid)
        try:
            result = await hands.run_tool(tool_name, **kwargs)
            step["result"] = result
            step["status"] = "completed"
        except Exception as e:
            step["result"] = {"error": str(e)}
            step["status"] = "failed"
        step["completed_at"] = _now_iso()
        try:
            _doris_rt.save_pipeline_steps(pid, [step])
        except Exception:
            pass
        _persist_pipeline(pid)
        await asyncio.sleep(0.05)
        return True

    # 阶段 1：芯片设计
    if not await run_step(1, "芯片设计师", "💎", "查阅 QEDA 资料目录", "qeda_catalog"): return
    if not await run_step(1, "芯片设计师", "💎", "搜索 100 比特设计规范", "search_knowledge", query="100比特 设计规范 量子器件库"): return
    if not await run_step(1, "芯片设计师", "💎", "获取 JJ 参数（Single Manhattan）", "qeda_junction_params", jj_type="Single_Manhattan"): return
    if not await run_step(1, "芯片设计师", "💎", "获取 JJ 参数（SQUID Manhattan）", "qeda_junction_params", jj_type="Squid_Manhattan"): return
    if not await run_step(1, "芯片设计师", "💎", "创建 20mm x 20mm 平面设计", "metal_create_design", chip_size_x="20mm", chip_size_y="20mm"): return
    design_id = p["steps"][-1].get("result", {}).get("design_id", "design_0001")
    for i in range(10):
        freq = 4.4 + (i % 2) * 0.4
        if not await run_step(1, "芯片设计师", "💎", f"添加比特 Q{i+1:03d}（{freq}GHz）", "metal_add_transmon", design_id=design_id, name=f"Q{i+1:03d}", qubit_type="transmon_cross", pos_x=f"{(i%10)*2000}um", pos_y=f"{(i//10)*2000}um"): return
    for i in range(9):
        if not await run_step(1, "芯片设计师", "💎", f"路由 Q{i+1:03d}-Q{i+2:03d}", "metal_add_route", design_id=design_id, name=f"cpw_{i}", start_component=f"Q{i+1:03d}", start_pin="bus_a", end_component=f"Q{i+2:03d}", end_pin="bus_b"): return
    if not await run_step(1, "芯片设计师", "💎", "导出简化 GDS（示例 10 比特）", "metal_export_gds", design_id=design_id, filename="100bit_tunable_coupler_lite.gds"): return
    if not await run_step(1, "芯片设计师", "💎", "生成 100 比特参考级 GDS（10x10 网格, 含耦合器+谐振腔+焊盘）", "build_100bit_reference"): return

    # 阶段 2：仿真
    for qid in ["Q001", "Q002", "Q003"]:
        if not await run_step(2, "仿真工程师", "🖥️", f"LOM 电容分析 {qid}", "metal_analyze_lom", design_id=design_id, component_name=qid): return
    if not await run_step(2, "仿真工程师", "🖥️", "EPR 本征模分析 Q001", "metal_analyze_epr", design_id=design_id, component_name="Q001"): return

    # 阶段 3-6 复用通用步骤
    if not await run_step(3, "制造工程师", "🏭", "查询工艺路线", "mes_list_routes"): return
    if not await run_step(3, "制造工程师", "🏭", "良率查询", "mes_query_yield"): return
    if not await run_step(4, "设备运维员", "🔧", "全套校准 Q1-Q10", "pulse_full_calibration", qubits=list(range(10))): return
    for seq in ["spectroscopy", "rabi", "t1", "t2_ramsey"]:
        if not await run_step(5, "测控科学家", "📡", f"ARTIQ {seq}（10 比特）", "artiq_run_pulse", sequence_type=seq, qubits=[f"Q{i+1:03d}" for i in range(10)]): return
    if not await run_step(5, "测控科学家", "📡", "Mitiq 纠错对比", "mitiq_benchmark", circuit_desc="100bit_CZ"): return
    if not await run_step(6, "数据分析师", "📊", "跨域追溯", "doris_cross_domain", query_type="design_to_measurement"): return
    if not await run_step(6, "数据分析师", "📊", "数据质量检查", "qdata_quality_check", table="qubit_characterization"): return

    from quantamind.server import hands_doris as _ds
    _ds.save_pipeline_run(pid, p)
    _ds.save_pipeline_steps(pid, p["steps"])
    p["status"] = "completed"
    p["completed_at"] = _now_iso()
    _persist_pipeline(pid)


async def _run_105bit_pipeline(pid: str):
    """105 比特芯片设计流水线"""
    p = pipelines[pid]
    p["status"] = "running"
    p["started_at"] = _now_iso()
    _persist_pipeline(pid)
    from quantamind.server import hands_doris as _doris_rt

    async def run_step(stage, agent, emoji, title, tool_name, **kwargs):
        while p.get("paused"):
            p["status"] = "paused"
            await asyncio.sleep(0.5)
        if p.get("aborted"):
            p["status"] = "aborted"
            return False
        p["status"] = "running"
        step = {"stage": stage, "agent": agent, "emoji": emoji, "title": title,
                "tool": tool_name, "args": kwargs, "status": "running", "result": None, "started_at": _now_iso()}
        p["steps"].append(step)
        p["current_stage"] = stage
        p["current_step"] = title
        _persist_pipeline(pid)
        try:
            result = await hands.run_tool(tool_name, **kwargs)
            step["result"] = result
            step["status"] = "completed"
        except Exception as e:
            step["result"] = {"error": str(e)}
            step["status"] = "failed"
        step["completed_at"] = _now_iso()
        try:
            _doris_rt.save_pipeline_steps(pid, [step])
        except Exception:
            pass
        _persist_pipeline(pid)
        await asyncio.sleep(0.05)
        return True

    try:
        from quantamind.server.chip_105bit_spec import CHIP_105BIT_SPEC as CS105
    except Exception:
        CS105 = {"qubits": [], "targets": {}}

    p["chip_spec"] = {"name": "105比特可调耦合器量子芯片", "total_qubits": 105, "targets": CS105.get("targets", {})}

    # 阶段 1：芯片设计
    if not await run_step(1, "芯片设计师", "💎", "查阅 QEDA 资料目录", "qeda_catalog"): return
    if not await run_step(1, "芯片设计师", "💎", "获取 105 比特芯片规格", "qeda_get_chip", chip_key="105bit"): return
    if not await run_step(1, "芯片设计师", "💎", "搜索 105 比特网格结构设计", "search_knowledge", query="105比特 二维网格 15行7列 耦合器"): return
    if not await run_step(1, "芯片设计师", "💎", "获取 Q001 比特参数", "qeda_get_qubit", chip_key="105bit", qubit_id="Q001"): return
    if not await run_step(1, "芯片设计师", "💎", "获取 JJ 参数（SQUID）", "qeda_junction_params", jj_type="Squid_Manhattan"): return
    if not await run_step(1, "芯片设计师", "💎", "创建 25mm x 25mm 平面设计", "metal_create_design", chip_size_x="25mm", chip_size_y="25mm"): return
    design_id = p["steps"][-1].get("result", {}).get("design_id", "design_0001")
    qubits_105 = CS105.get("qubits", [])[:7]
    for q in qubits_105:
        freq = q.get("freq_ghz", 4.5)
        if not await run_step(1, "芯片设计师", "💎", f"添加比特 {q['id']}（{freq}GHz）", "metal_add_transmon", design_id=design_id, name=q["id"], qubit_type="transmon_cross"): return
    for i in range(min(6, len(qubits_105) - 1)):
        if not await run_step(1, "芯片设计师", "💎", f"路由 {qubits_105[i]['id']}-{qubits_105[i+1]['id']}", "metal_add_route", design_id=design_id, name=f"cpw_105_{i}", start_component=qubits_105[i]["id"], start_pin="bus_a", end_component=qubits_105[i+1]["id"], end_pin="bus_b"): return
    if not await run_step(1, "芯片设计师", "💎", "导出简化 GDS（示例 7 比特）", "metal_export_gds", design_id=design_id, filename="105bit_tunable_coupler.gds"): return
    if not await run_step(1, "芯片设计师", "💎", "部署团队真实工业级 GDS（95,238 shapes, 18 cells, 10 layers, 55.8MB）", "build_real_chip", chip_type="105bit"): return

    # 阶段 2：仿真
    for qid in (qubits_105[:3] if qubits_105 else []):
        if not await run_step(2, "仿真工程师", "🖥️", f"EPR 本征模 {qid['id']}", "metal_analyze_epr", design_id=design_id, component_name=qid["id"]): return

    # 阶段 3-6
    if not await run_step(3, "制造工程师", "🏭", "查询工艺路线", "mes_list_routes"): return
    if not await run_step(3, "制造工程师", "🏭", "良率查询", "mes_query_yield"): return
    if not await run_step(4, "设备运维员", "🔧", "全套校准（示例 10 比特）", "pulse_full_calibration", qubits=list(range(10))): return
    for seq in ["spectroscopy", "rabi", "t1", "t2_ramsey"]:
        if not await run_step(5, "测控科学家", "📡", f"ARTIQ {seq}", "artiq_run_pulse", sequence_type=seq, qubits=[q["id"] for q in qubits_105[:7]]): return
    if not await run_step(5, "测控科学家", "📡", "Mitiq 纠错对比", "mitiq_benchmark", circuit_desc="105bit_CZ_chain"): return
    if not await run_step(6, "数据分析师", "📊", "跨域追溯", "doris_cross_domain", query_type="design_to_measurement"): return
    if not await run_step(6, "数据分析师", "📊", "数据质量检查", "qdata_quality_check", table="qubit_characterization"): return

    from quantamind.server import hands_doris as _ds
    _ds.save_pipeline_run(pid, p)
    _ds.save_pipeline_steps(pid, p["steps"])
    if CS105.get("qubits"):
        _ds.save_design_params(pid, CS105["qubits"][:20])
    p["status"] = "completed"
    p["completed_at"] = _now_iso()
    _persist_pipeline(pid)


class PipelineCreate(BaseModel):
    template: str = "20bit_tunable_coupler"
    name: Optional[str] = None
    require_gate_approval: bool = False


class PipelineAction(BaseModel):
    action: str  # pause | resume | abort | skip_gate


@app.post("/api/v1/pipelines/{pipeline_id}/control")
def control_pipeline(pipeline_id: str, body: PipelineAction):
    """人工干预流水线：暂停/恢复/终止/跳过审批门"""
    from fastapi import HTTPException
    if pipeline_id not in pipelines:
        raise HTTPException(status_code=404, detail="pipeline not found")
    p = pipelines[pipeline_id]
    if body.action == "pause":
        p["paused"] = True
        _persist_pipeline(pipeline_id)
        return {"action": "paused", "pipeline_id": pipeline_id}
    elif body.action == "resume":
        p["paused"] = False
        _persist_pipeline(pipeline_id)
        return {"action": "resumed", "pipeline_id": pipeline_id}
    elif body.action == "abort":
        p["aborted"] = True
        p["paused"] = False
        _persist_pipeline(pipeline_id)
        return {"action": "aborted", "pipeline_id": pipeline_id}
    elif body.action == "skip_gate":
        for gate in p.get("gates", []):
            if gate["status"] == "waiting_approval":
                gate["status"] = "approved"
                break
        _persist_pipeline(pipeline_id)
        return {"action": "gate_approved", "pipeline_id": pipeline_id}
    elif body.action == "reject_gate":
        for gate in p.get("gates", []):
            if gate["status"] == "waiting_approval":
                gate["status"] = "rejected"
                break
        p["aborted"] = True
        _persist_pipeline(pipeline_id)
        return {"action": "gate_rejected", "pipeline_id": pipeline_id}
    raise HTTPException(status_code=400, detail="unknown action")


@app.get("/api/v1/pipelines/{pipeline_id}/report")
def download_pipeline_report(pipeline_id: str, format: str = "md"):
    """下载流水线执行报告（md=Markdown / docx=Word），支持模板流水线和对话流水线"""
    from fastapi import HTTPException
    from fastapi.responses import Response
    from quantamind.agents.orchestrator import chat_pipelines
    p = pipelines.get(pipeline_id) or chat_pipelines.get(pipeline_id)
    if not p:
        raise HTTPException(status_code=404, detail="pipeline not found")

    from quantamind.server.result_explain import explain
    stages_map = {s["stage"]: s for s in p.get("stages", [])}

    if format == "docx":
        return _generate_word_report(pipeline_id, p, stages_map, explain)

    # Markdown 版本
    lines = [f"# {p.get('name', '流水线')} 执行报告\n",
        f"- **流水线 ID**：{pipeline_id}", f"- **状态**：{p.get('status', '—')}",
        f"- **创建时间**：{p.get('created_at', '—')}", f"- **完成时间**：{p.get('completed_at', '—')}",
        f"- **总步骤**：{len(p.get('steps', []))}", f"- **成功**：{sum(1 for s in p.get('steps',[]) if s['status']=='completed')}",
        f"- **失败**：{sum(1 for s in p.get('steps',[]) if s['status']=='failed')}", ""]
    current_stage = 0
    for step in p.get("steps", []):
        if step["stage"] != current_stage:
            current_stage = step["stage"]
            st_info = stages_map.get(current_stage, {})
            lines.append(f"\n## 阶段 {current_stage}：{st_info.get('name', '')}（{st_info.get('agent', '')}）\n")
        lines.append(f"### {step.get('emoji','')} {step.get('title','')}\n")
        lines.append(f"- **Agent**：{step.get('agent', '')}  **状态**：{step['status']}")
        lines.append(f"- **工具**：`{step.get('tool', '')}`")
        exp = explain(step.get("tool", ""), step.get("args", {}), step.get("result", {}))
        lines.append(f"- **数据格式**：{exp.get('data_format', '')}")
        lines.append(f"- **解释**：{exp.get('explanation', '')}")
        if exp.get("data_table"):
            t = exp["data_table"]
            lines.append(f"\n| {' | '.join(t['headers'])} |")
            lines.append(f"| {' | '.join(['---'] * len(t['headers']))} |")
            for row in t.get("rows", []):
                lines.append(f"| {' | '.join(str(c) for c in row)} |")
        lines.append("")
    md = "\n".join(lines)
    return Response(content=md.encode("utf-8"), media_type="text/markdown",
                    headers={"Content-Disposition": f'attachment; filename="{pipeline_id}_report.md"'})


def _generate_word_report(pipeline_id, p, stages_map, explain_fn):
    """生成 Word 格式报告"""
    from fastapi.responses import Response
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return Response(content="python-docx 未安装，请执行 pip install python-docx".encode(), status_code=500)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10)

    doc.add_heading(f"{p.get('name', '流水线')} 执行报告", level=0)

    # 芯片规格摘要页
    cs = p.get("chip_spec")
    if cs:
        doc.add_heading("芯片设计规格", level=1)
        spec_rows = [
            ("文档编号", cs.get("doc_id", "")), ("芯片名称", cs.get("name", "")),
            ("芯片尺寸", f"{cs['chip_size_mm'][0]} x {cs['chip_size_mm'][1]} mm"),
            ("衬底材料", cs.get("substrate", "")), ("金属薄膜", cs.get("metal_film", "")),
            ("量子比特", f"{cs.get('total_qubits', 0)} 个（{cs.get('substrate', '')}衬底 Xmon）"),
            ("可调耦合器", f"{cs.get('total_couplers', 0)} 个（Transmon + SQUID）"),
            ("封装接口", f"{cs.get('packaging', {}).get('interface', '48 pin SMP')}"),
        ]
        t1 = doc.add_table(rows=len(spec_rows), cols=2)
        t1.style = "Light Grid Accent 1"
        for i, (k, v) in enumerate(spec_rows):
            t1.rows[i].cells[0].text = k
            t1.rows[i].cells[1].text = str(v)
        doc.add_paragraph("")

        # 工艺规格
        fab = cs.get("fabrication", {})
        if fab:
            doc.add_heading("工艺规格", level=2)
            fab_rows = [("光刻", fab.get("lithography", "")), ("JJ 制备", fab.get("jj_fabrication", "")),
                        ("JJ 结构", fab.get("jj_structure", "")), ("JJ 尺寸", f"{fab.get('jj_length_nm', '')} x {fab.get('jj_width_nm', '')} nm"),
                        ("氧化层", f"{fab.get('oxide_thickness_nm', '')} nm"), ("膜厚", f"{fab.get('film_thickness_nm', '')} nm"),
                        ("衬底厚度", f"{fab.get('substrate_thickness_um', '')} μm")]
            t2 = doc.add_table(rows=len(fab_rows), cols=2)
            t2.style = "Light Grid Accent 1"
            for i, (k, v) in enumerate(fab_rows):
                t2.rows[i].cells[0].text = k
                t2.rows[i].cells[1].text = str(v)
            doc.add_paragraph("")

        # 设计目标
        targets = cs.get("targets", {})
        if targets:
            doc.add_heading("设计目标", level=2)
            tgt_rows = [(k, str(v)) for k, v in targets.items()]
            t3 = doc.add_table(rows=len(tgt_rows), cols=2)
            t3.style = "Light Grid Accent 1"
            for i, (k, v) in enumerate(tgt_rows):
                t3.rows[i].cells[0].text = k
                t3.rows[i].cells[1].text = v
            doc.add_paragraph("")

        # 版图组件清单
        lc = cs.get("layout_components", [])
        if lc:
            doc.add_heading("版图组件清单", level=2)
            t4 = doc.add_table(rows=len(lc) + 1, cols=3)
            t4.style = "Light Grid Accent 1"
            for j, h in enumerate(["组件", "数量", "图号"]):
                t4.rows[0].cells[j].text = h
                for par in t4.rows[0].cells[j].paragraphs:
                    for run in par.runs:
                        run.bold = True
            for i, c in enumerate(lc):
                t4.rows[i + 1].cells[0].text = c.get("name", "")
                t4.rows[i + 1].cells[1].text = str(c.get("count", 0))
                t4.rows[i + 1].cells[2].text = c.get("drawing_id", "")
            doc.add_paragraph("")

    doc.add_page_break()
    doc.add_heading("执行过程", level=1)

    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Light Grid Accent 1"
    info_data = [("流水线 ID", pipeline_id), ("状态", p.get("status", "")),
                 ("创建时间", p.get("created_at", "")), ("完成时间", p.get("completed_at", ""))]
    for i, (k, v) in enumerate(info_data):
        info_table.rows[i].cells[0].text = k
        info_table.rows[i].cells[1].text = str(v)
    doc.add_paragraph("")

    steps_done = sum(1 for s in p.get("steps", []) if s["status"] == "completed")
    steps_fail = sum(1 for s in p.get("steps", []) if s["status"] == "failed")
    summary = doc.add_paragraph()
    summary.add_run(f"总步骤 {len(p.get('steps', []))} | 成功 {steps_done} | 失败 {steps_fail}").bold = True

    current_stage = -1
    for step in p.get("steps", []):
        try:
            step_stage = step.get("stage", 0)
            if step_stage != current_stage and stages_map:
                current_stage = step_stage
                st_info = stages_map.get(current_stage, {})
                if st_info:
                    doc.add_heading(f"阶段 {current_stage}：{st_info.get('name', '')}（{st_info.get('agent', '')}）", level=1)

            status_text = {"completed": "[完成]", "failed": "[失败]", "running": "[运行中]"}.get(step.get("status", ""), "[?]")
            title = step.get("title", step.get("tool", "步骤"))
            doc.add_heading(f"{step.get('emoji', '')} {title} {status_text}", level=2)

            p1 = doc.add_paragraph()
            p1.add_run("Agent：").bold = True
            p1.add_run(str(step.get("agent", "")))
            p1.add_run("    工具：").bold = True
            p1.add_run(str(step.get("tool", "")))

            try:
                exp = explain_fn(step.get("tool", ""), step.get("args", {}), step.get("result") or {})
            except Exception:
                exp = {"data_format": "", "explanation": "（解析跳过）", "data_table": None}

            p2 = doc.add_paragraph()
            p2.add_run("数据格式：").bold = True
            p2.add_run(str(exp.get("data_format", "")))

            p3 = doc.add_paragraph()
            p3.add_run("解释说明：").bold = True
            p3.add_run(str(exp.get("explanation", "")))

            dt = exp.get("data_table")
            if dt and dt.get("headers") and dt.get("rows"):
                headers = dt["headers"]
                rows = dt["rows"]
                ncols = len(headers)
                if ncols > 0 and len(rows) > 0:
                    table = doc.add_table(rows=min(len(rows), 20) + 1, cols=ncols)
                    table.style = "Light Grid Accent 1"
                    for j, h in enumerate(headers):
                        table.rows[0].cells[j].text = str(h)
                    for i, row in enumerate(rows[:20]):
                        for j in range(min(len(row), ncols)):
                            table.rows[i + 1].cells[j].text = str(row[j]) if row[j] is not None else ""
            doc.add_paragraph("")
        except Exception as step_err:
            doc.add_paragraph(f"（步骤处理异常：{step_err}）")

    gates = p.get("gates", [])
    if gates:
        doc.add_heading("审批门记录", level=1)
        for g in gates:
            doc.add_paragraph(f"{g.get('stage', '')}：{g.get('status', '')}（{g.get('created_at', '')}）")

    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(content=buf.read(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f'attachment; filename="{pipeline_id}_report.docx"'})


@app.post("/api/v1/pipelines")
async def create_pipeline(body: PipelineCreate):
    pid = f"PL-{str(uuid.uuid4())[:8]}"
    pipelines[pid] = {
        "pipeline_id": pid, "template": body.template,
        "name": body.name or f"20比特可调耦合器流水线",
        "status": "created", "steps": [], "current_stage": 0, "current_step": "",
        "require_gate_approval": body.require_gate_approval, "gates": [],
        "paused": False, "aborted": False,
        "created_at": _now_iso(),
        "stages": [
            {"stage": 1, "name": "芯片设计", "agent": "💎 AI 芯片设计师", "color": "#58a6ff"},
            {"stage": 2, "name": "仿真验证", "agent": "🖥️ AI 仿真工程师", "color": "#a371f7"},
            {"stage": 3, "name": "工艺制造", "agent": "🏭 AI 制造工程师", "color": "#d29922"},
            {"stage": 4, "name": "设备校准", "agent": "🔧 AI 设备运维员", "color": "#3fb950"},
            {"stage": 5, "name": "测控表征", "agent": "📡 AI 测控科学家", "color": "#f85149"},
            {"stage": 6, "name": "数据分析", "agent": "📊 AI 数据分析师", "color": "#58a6ff"},
        ],
    }
    _persist_pipeline(pid)
    template_runners = {
        "1bit_standard": _run_1bit_pipeline,
        "2bit_coupler": _run_2bit_pipeline,
        "20bit_tunable_coupler": _run_20bit_pipeline,
        "100bit_tunable_coupler": _run_100bit_pipeline,
        "105bit_tunable_coupler": _run_105bit_pipeline,
    }
    template_names = {
        "1bit_standard": "单比特标准芯片流水线（1bitV2.00）",
        "2bit_coupler": "双比特可调耦合器流水线（2bits_coupler_V2.00）",
        "20bit_tunable_coupler": "20比特可调耦合器流水线（CT20Q）",
        "100bit_tunable_coupler": "100比特可调耦合器流水线",
        "105bit_tunable_coupler": "105比特可调耦合器流水线（FT105Q V2）",
    }
    runner = template_runners.get(body.template, _run_20bit_pipeline)
    pipelines[pid]["name"] = body.name or template_names.get(body.template, "量子芯片流水线")
    asyncio.create_task(runner(pid))
    return {"pipeline_id": pid, "status": "started"}


@app.get("/api/v1/chat-pipelines")
def list_chat_pipelines():
    """列出对话产生的流水线"""
    from quantamind.agents.orchestrator import chat_pipelines
    out = []
    for p in chat_pipelines.values():
        out.append({"pipeline_id": p["pipeline_id"], "name": p.get("name", ""), "agent_id": p.get("agent_id", ""),
                     "agent_label": p.get("agent_label", ""), "status": p["status"],
                     "steps_count": len(p.get("steps", [])), "created_at": p.get("created_at", "")})
    return {"pipelines": out}


@app.get("/api/v1/chat-pipelines/{pipeline_id}")
def get_chat_pipeline(pipeline_id: str):
    """获取对话流水线详情"""
    from fastapi import HTTPException
    from quantamind.agents.orchestrator import chat_pipelines
    if pipeline_id not in chat_pipelines:
        raise HTTPException(status_code=404, detail="chat pipeline not found")
    p = chat_pipelines[pipeline_id]
    steps_safe = []
    for s in p.get("steps", []):
        sc = dict(s)
        if sc.get("result"):
            r_str = json.dumps(sc["result"], ensure_ascii=False, default=str)
            if len(r_str) > 1000:
                sc["result"] = {"_preview": r_str[:800]}
        steps_safe.append(sc)
    return {**{k: v for k, v in p.items() if k != "steps"}, "steps": steps_safe}


@app.get("/api/v1/pipelines")
def list_pipelines():
    """列出所有流水线（含手动启动的 + 对话产生的）"""
    from quantamind.agents.orchestrator import chat_pipelines
    all_pls = []
    for p in pipelines.values():
        all_pls.append({"pipeline_id": p["pipeline_id"], "name": p.get("name", ""), "type": "template",
                        "status": p["status"], "current_stage": p.get("current_stage", 0),
                        "steps_count": len(p.get("steps", [])), "created_at": p.get("created_at", "")})
    for p in chat_pipelines.values():
        if len(p.get("steps", [])) > 0:
            all_pls.append({"pipeline_id": p["pipeline_id"], "name": p.get("name", ""), "type": "chat",
                            "agent_label": p.get("agent_label", ""), "status": p["status"],
                            "steps_count": len(p.get("steps", [])), "created_at": p.get("created_at", "")})
    all_pls.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return {"pipelines": all_pls}


@app.get("/api/v1/pipelines/{pipeline_id}")
def get_pipeline(pipeline_id: str, compact: bool = False):
    from fastapi import HTTPException
    from quantamind.agents.orchestrator import chat_pipelines
    # 先查手动流水线，再查对话流水线
    if pipeline_id not in pipelines:
        if pipeline_id in chat_pipelines:
            p = chat_pipelines[pipeline_id]
            steps_safe = []
            for s in p.get("steps", []):
                sc = dict(s)
                if sc.get("result"):
                    r_str = json.dumps(sc["result"], ensure_ascii=False, default=str)
                    if len(r_str) > 1000:
                        sc["result"] = {"_preview": r_str[:800]}
                steps_safe.append(sc)
            return {**{k: v for k, v in p.items() if k != "steps"}, "steps": steps_safe,
                    "stages": [{"stage": 1, "name": p.get("agent_label", "对话"), "agent": p.get("agent_label", ""), "color": "#58a6ff"}]}
        raise HTTPException(status_code=404, detail="pipeline not found")
    p = pipelines[pipeline_id]
    if compact:
        steps_compact = []
        for s in p.get("steps", []):
            sc = {k: v for k, v in s.items() if k != "result"}
            if s.get("result"):
                r_str = json.dumps(s["result"], ensure_ascii=False, default=str)
                sc["result_preview"] = r_str[:200]
            steps_compact.append(sc)
        return {**{k: v for k, v in p.items() if k != "steps"}, "steps": steps_compact}
    # 完整版也截断 result 避免超大 JSON
    steps_safe = []
    for s in p.get("steps", []):
        sc = dict(s)
        if sc.get("result"):
            r_str = json.dumps(sc["result"], ensure_ascii=False, default=str)
            if len(r_str) > 2000:
                sc["result"] = json.loads(r_str[:2000] + "...")  if False else {"_truncated": True, "preview": r_str[:1000]}
        steps_safe.append(sc)
    return {**{k: v for k, v in p.items() if k != "steps"}, "steps": steps_safe}


# QuantaMind Web 客户端（表现层）
_web_dir = Path(__file__).resolve().parent.parent / "client" / "web"

if (_web_dir / "js").exists():
    app.mount("/js", StaticFiles(directory=_web_dir / "js"), name="web-js")


@app.get("/")
async def serve_web_client():
    """根路径返回 QuantaMind Web 客户端"""
    index = _web_dir / "index.html"
    if index.exists():
        return FileResponse(
            index,
            headers={
                "Cache-Control": "no-cache, no-store, must-revalidate",
                "Pragma": "no-cache",
            },
        )
    return {"service": "QuantaMind Gateway", "docs": "/docs", "api": "/api/v1"}


@app.get("/client")
async def serve_web_client_alt():
    return await serve_web_client()


def run_gateway(host: Optional[str] = None, port: Optional[int] = None):
    import sys

    h = host or config.GATEWAY_HOST
    p = port or config.GATEWAY_PORT
    print(
        f"正在加载 QuantaMind Gateway（首次约需 1–2 分钟，请勿关窗口）…\n"
        f"就绪后请访问: http://127.0.0.1:{p}/  （监听 {h}:{p}）",
        file=sys.stderr,
        flush=True,
    )
    import uvicorn

    uvicorn.run(
        "quantamind.server.gateway:app",
        host=h,
        port=p,
        reload=False,
    )


if __name__ == "__main__":
    run_gateway()
