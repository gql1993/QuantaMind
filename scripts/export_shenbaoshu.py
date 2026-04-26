"""
将申报书 Markdown 导出为符合国家重点研发计划申报格式的 Word 文件。
排版规范：
  封面          黑体一号居中
  一级标题      黑体三号
  正文           宋体小四（12pt），首行缩进2字符，1.5倍行距
  表格表头       黑体小四，正文宋体小四
  页眉           宋体五号居中
  页脚           页码居中
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

FIGURES_DIR = Path(__file__).resolve().parent.parent / "docs" / "_generated_figures"

# ── 字体常量 ──────────────────────────────────────────────
FT_SONG = "宋体"
FT_HEI  = "黑体"
FT_YAHE = "Microsoft YaHei"

# ── 字号辅助 (Word Pt == 磅) ─────────────────────────────
PT_YI   = Pt(26)   # 一号
PT_ER   = Pt(22)   # 二号
PT_XIAO_ER = Pt(18)
PT_SAN  = Pt(16)   # 三号
PT_XIAO_SAN = Pt(15)
PT_SI   = Pt(14)   # 四号
PT_XIAO_SI = Pt(12)  # 小四 ★ 正文
PT_WU   = Pt(10.5)  # 五号


def _rpr(run):
    """Ensure rPr element exists and return it."""
    if run._r.rPr is None:
        rpr = OxmlElement("w:rPr")
        run._r.insert(0, rpr)
    return run._r.rPr


def set_run_font(run, name_latin: str, name_east: str, size: Pt,
                 bold: bool = False, color: str | None = None) -> None:
    run.font.name = name_latin
    _rpr(run).rFonts.set(qn("w:eastAsia"), name_east)
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field_run(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    for tag, attr, val in [
        ("w:fldChar",   "w:fldCharType", "begin"),
        ("w:instrText", "xml:space",     "preserve"),
        ("w:fldChar",   "w:fldCharType", "separate"),
        ("w:fldChar",   "w:fldCharType", "end"),
    ]:
        el = OxmlElement(tag)
        if attr == "xml:space":
            el.text = instruction
        else:
            el.set(qn(attr), val)
        run._r.append(el)


# ── 页面布局 ──────────────────────────────────────────────
def set_page(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin    = Cm(3.7)
    sec.bottom_margin = Cm(3.5)
    sec.left_margin   = Cm(2.8)
    sec.right_margin  = Cm(2.6)
    sec.header_distance = Cm(2.5)
    sec.footer_distance = Cm(1.5)
    sec.different_first_page_header_footer = True


def set_header_footer(doc: Document, header_text: str) -> None:
    sec = doc.sections[0]
    hdr = sec.header
    hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 分隔线
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom); pPr.append(pBdr)

    run = hp.add_run(header_text)
    set_run_font(run, FT_SONG, FT_SONG, PT_WU)

    ftr = sec.footer
    fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for txt in ["第 ", " 页"]:
        if txt == " 页":
            add_field_run(fp, " PAGE ")
        r = fp.add_run(txt)
        set_run_font(r, FT_SONG, FT_SONG, PT_WU)


# ── 样式默认值 ────────────────────────────────────────────
def init_styles(doc: Document) -> None:
    # Normal
    n = doc.styles["Normal"]
    n.font.name = FT_SONG
    n._element.rPr.rFonts.set(qn("w:eastAsia"), FT_SONG)
    n.font.size = PT_XIAO_SI
    n.paragraph_format.line_spacing    = Pt(20)   # ≈ 1.5 倍
    n.paragraph_format.first_line_indent = Cm(0.84)
    n.paragraph_format.space_after  = Pt(0)
    n.paragraph_format.space_before = Pt(0)

    heading_cfg = [
        ("Title",     PT_YI,       True,  WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 1", PT_SAN,      True,  WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 2", PT_SI,       True,  WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", PT_XIAO_SI,  True,  WD_ALIGN_PARAGRAPH.LEFT),
    ]
    for sname, sz, bold, align in heading_cfg:
        st = doc.styles[sname]
        st.font.name = FT_HEI
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FT_HEI)
        st.font.size = sz
        st.font.bold = bold
        st.paragraph_format.line_spacing  = Pt(22)
        st.paragraph_format.first_line_indent = Cm(0)
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after  = Pt(4)
        st.paragraph_format.alignment    = align

    for sname in ("List Bullet", "List Number"):
        st = doc.styles[sname]
        st.font.name = FT_SONG
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FT_SONG)
        st.font.size = PT_XIAO_SI
        st.paragraph_format.line_spacing = Pt(20)
        st.paragraph_format.first_line_indent = Cm(0)


# ── 封面页 ────────────────────────────────────────────────
def add_cover(doc: Document, proj_title: str, org: str, date_str: str) -> None:
    def cp(text: str, size: Pt, bold: bool = False, sp_before: int = 0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(sp_before)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(text)
        set_run_font(r, FT_HEI, FT_HEI, size, bold)

    cp("国家重点研发计划", PT_XIAO_ER, True, sp_before=60)
    cp("项  目  申  报  书", PT_YI, True, sp_before=20)

    doc.add_paragraph().paragraph_format.space_before = Pt(40)

    # 表格式封面信息
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    meta = [
        ("项目名称", proj_title),
        ("申报单位", org),
        ("项目周期", "2026年1月—2030年12月"),
        ("申请经费", "（万元）"),
        ("填报日期", date_str),
    ]
    for i, (k, v) in enumerate(meta):
        if i >= len(tbl.rows):
            tbl.add_row()
        c0, c1 = tbl.cell(i, 0), tbl.cell(i, 1)
        c0.text = k; c1.text = v
        for cell, bold in [(c0, True), (c1, False)]:
            for para in cell.paragraphs:
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after  = Pt(4)
                for run in para.runs:
                    set_run_font(run, FT_SONG, FT_SONG, PT_XIAO_SI, bold)

    doc.add_page_break()


# ── 目录页 ────────────────────────────────────────────────
def add_toc(doc: Document) -> None:
    ph = doc.add_paragraph(style="Heading 1")
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ph.paragraph_format.first_line_indent = Cm(0)
    r = ph.add_run("目  录")
    set_run_font(r, FT_HEI, FT_HEI, PT_SAN, True)

    tp = doc.add_paragraph()
    tp.paragraph_format.first_line_indent = Cm(0)
    add_field_run(tp, r' TOC \o "1-3" \h \z \u ')

    note = doc.add_paragraph()
    note.paragraph_format.first_line_indent = Cm(0)
    nr = note.add_run("（请在 Word 中右键目录区域 → 更新域，以刷新页码）")
    set_run_font(nr, FT_SONG, FT_SONG, PT_WU, color="888888")

    doc.add_page_break()


# ── 行内富文本 ────────────────────────────────────────────
def add_inline(para, text: str, base_size: Pt = PT_XIAO_SI) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        bold   = part.startswith("**") and part.endswith("**")
        mono   = part.startswith("`") and part.endswith("`")
        clean  = part.strip("*").strip("`")
        run    = para.add_run(clean)
        fname  = "Consolas" if mono else FT_SONG
        feast  = "Consolas" if mono else FT_SONG
        set_run_font(run, fname, feast, base_size, bold)


# ── 表格 ─────────────────────────────────────────────────
def parse_table(lines: list[str], start: int):
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(row)
        i += 1
    if len(rows) >= 2 and all(set(c.strip()) <= {"-", ":"} for c in rows[1]):
        rows.pop(1)
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=cols)
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = tbl.cell(ri, ci)
            cell.text = row[ci] if ci < len(row) else ""
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                p.paragraph_format.line_spacing = Pt(18)
                for run in p.runs:
                    set_run_font(run, FT_SONG, FT_SONG, PT_WU, bold=(ri == 0))
    # 表头底纹
    from docx.oxml import OxmlElement as OE
    for cell in tbl.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OE("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D6E4F5")
        tcPr.append(shd)


# ── 插图 ─────────────────────────────────────────────────
def try_insert_figure(doc: Document, caption: str) -> bool:
    m = re.search(r"图\s*(\d+)", caption)
    if not m:
        return False
    num  = m.group(1)
    path = FIGURES_DIR / f"figure_{num}.png"
    if not path.exists():
        return False

    # 图片段落（居中，无首行缩进）
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.first_line_indent = Cm(0)
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after  = Pt(0)
    run = p_img.add_run()
    run.add_picture(str(path), width=Inches(6.0))

    # 图注段落
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.first_line_indent = Cm(0)
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after  = Pt(8)
    cap_run = p_cap.add_run(caption)
    set_run_font(cap_run, FT_SONG, FT_SONG, PT_WU, bold=True)
    return True


# ── 正文行解析 ────────────────────────────────────────────
MAJOR_SECTION_RE = re.compile(r"^##\s*[一二三四五六七八九十]+[、．.]")
COVER_SKIP_RE    = re.compile(r"^##\s*(封面信息|目录|中心一句话总述|模块正文)")
FIGURE_RE        = re.compile(r"^\*\*(图\s*\d+[^*]*)\*\*$")


def export(md_path: Path, docx_path: Path) -> None:
    lines  = md_path.read_text(encoding="utf-8").splitlines()
    doc    = Document()
    init_styles(doc)
    set_page(doc)

    # 提取标题
    proj_title = "量子科技自主科研人工智能中台关键技术研究与系统研制"
    for line in lines[:10]:
        s = line.strip()
        if s.startswith("# ") and "申报书" not in s and "国家" not in s:
            proj_title = s[2:].strip(); break

    set_header_footer(doc, "国家重点研发计划项目申报书")
    add_cover(doc, proj_title, "量子计算工业软件工程技术研发中心", "2026年4月")
    add_toc(doc)

    # 记录是否已过封面/目录行
    body_started = False
    seen_major   = False
    i = 0

    while i < len(lines):
        raw     = lines[i]
        stripped = raw.strip()
        i += 1

        if not stripped or stripped == "---":
            continue

        # 跳过封面信息表行（在目录之前）
        if COVER_SKIP_RE.match(stripped):
            continue
        if not body_started:
            if re.match(r"^##\s*封面信息", stripped):
                continue
            if re.match(r"^##\s*目录", stripped):
                continue
            if re.match(r"^-\s+", stripped) and not body_started:
                continue  # 目录列表行
            # 遇到真正内容行则开始
            if re.match(r"^#", stripped) and "申报书" not in stripped and "国家" not in stripped:
                body_started = True
            elif not re.match(r"^#", stripped):
                body_started = True

        # ── 表格 ─────────────────────────────────────────
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i - 1)
            i = i  # already advanced
            add_table(doc, rows)
            continue

        # ── 插图 ─────────────────────────────────────────
        fm = FIGURE_RE.match(stripped)
        if fm:
            caption = fm.group(1)
            inserted = try_insert_figure(doc, caption)
            if not inserted:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                add_inline(p, f"【{caption}】", PT_WU)
            # 跳过紧跟的括号说明行
            if i < len(lines) and lines[i].strip().startswith("（"):
                i += 1
            continue

        # ── 标题层级 ─────────────────────────────────────
        if stripped.startswith("# "):
            # 文档一级标题——只在封面用，正文忽略
            continue

        if stripped.startswith("## "):
            text = stripped[3:].strip()
            if MAJOR_SECTION_RE.match(stripped):
                if seen_major:
                    doc.add_page_break()
                seen_major = True
            p = doc.add_paragraph(style="Heading 1")
            p.paragraph_format.first_line_indent = Cm(0)
            add_inline(p, text, PT_SAN)
            continue

        if stripped.startswith("### "):
            text = stripped[4:].strip()
            p = doc.add_paragraph(style="Heading 2")
            p.paragraph_format.first_line_indent = Cm(0)
            add_inline(p, text, PT_SI)
            continue

        if stripped.startswith("#### "):
            text = stripped[5:].strip()
            p = doc.add_paragraph(style="Heading 3")
            p.paragraph_format.first_line_indent = Cm(0)
            add_inline(p, text, PT_XIAO_SI)
            continue

        # ── 列表 ─────────────────────────────────────────
        bm = re.match(r"^[-*]\s+(.+)$", stripped)
        if bm:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.74)
            add_inline(p, bm.group(1), PT_XIAO_SI)
            continue

        nm = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if nm:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.74)
            add_inline(p, nm.group(2), PT_XIAO_SI)
            continue

        # ── 分隔线 / 签名尾行 ────────────────────────────
        if stripped.startswith("---"):
            continue
        if stripped.startswith("**申报单位") or stripped.startswith("**项目负责人") or stripped.startswith("**填报日期"):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(20)
            add_inline(p, stripped, PT_XIAO_SI)
            continue

        # ── 普通正文 ─────────────────────────────────────
        if not body_started:
            continue
        p = doc.add_paragraph()
        add_inline(p, stripped, PT_XIAO_SI)

    doc.save(docx_path)
    print(f"Word 已生成：{docx_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: python export_shenbaoshu.py <input.md> <output.docx>")
    export(Path(sys.argv[1]), Path(sys.argv[2]))
