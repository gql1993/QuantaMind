from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


def set_default_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    for style_name, size in [("Title", 18), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)


def set_page_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)
    section.different_first_page_header_footer = True


def add_field_run(paragraph, instruction: str):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(text)
    run._r.append(fld_char_end)
    return run


def set_header_footer(doc: Document, header_text: str) -> None:
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(header_text)
    hr.font.name = "宋体"
    hr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    hr.font.size = Pt(10.5)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr1 = fp.add_run("第 ")
    fr1.font.name = "宋体"
    fr1._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    fr1.font.size = Pt(10.5)
    add_field_run(fp, " PAGE ")
    fr2 = fp.add_run(" 页")
    fr2.font.name = "宋体"
    fr2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    fr2.font.size = Pt(10.5)


def add_cover(doc: Document, center_title: str, module_title: str, version: str, product_name: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run(center_title)
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(22)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(28)
    r2 = p2.add_run(module_title)
    r2.bold = True
    r2.font.name = "黑体"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r2.font.size = Pt(20)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(18)
    r3 = p3.add_run(product_name)
    r3.font.name = "宋体"
    r3._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r3.font.size = Pt(16)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    meta_rows = [
        ("对应中心", center_title),
        ("模块名称", module_title),
        ("版本", version),
        ("日期", "2026年4月"),
    ]
    for idx, (label, value) in enumerate(meta_rows):
        table.cell(idx, 0).text = label
        table.cell(idx, 1).text = value
    for row in table.rows:
        for cidx, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if cidx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.first_line_indent = Cm(0)
                for run in p.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(12)
                    if cidx == 0:
                        run.bold = True

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(20)
    r4 = p4.add_run("文档属性：申报材料正式稿")
    r4.font.name = "宋体"
    r4._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r4.font.size = Pt(14)

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    add_runs(p, "目录")

    toc_p = doc.add_paragraph()
    add_field_run(toc_p, r'TOC \o "1-3" \h \z \u')

    tip = doc.add_paragraph()
    tip.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tip.paragraph_format.first_line_indent = Cm(0)
    tip_run = tip.add_run("提示：若目录未自动刷新，可在 Word 中右键目录选择“更新域”。")
    tip_run.font.name = "宋体"
    tip_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    tip_run.font.size = Pt(10.5)

    doc.add_page_break()


def get_font(size: int, bold: bool = False) -> "ImageFont.FreeTypeFont | ImageFont.ImageFont":
    candidates = [
        ("C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc"),
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fill: str = "#F4F7FB") -> None:
    draw.rounded_rectangle(box, radius=14, outline="#4A6FA5", width=3, fill=fill)
    font = get_font(24, bold=True)
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=6, align="center")
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    x = (box[0] + box[2] - w) / 2
    y = (box[1] + box[3] - h) / 2
    draw.multiline_text((x, y), text, font=font, fill="#1F1F1F", spacing=6, align="center")


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill="#4A6FA5", width=4)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - 16 * direction, ey - 8), (ex - 16 * direction, ey + 8)]
    else:
        direction = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 8, ey - 16 * direction), (ex + 8, ey - 16 * direction)]
    draw.polygon(pts, fill="#4A6FA5")


def build_figure_image(caption: str, output_path: Path) -> None:
    width, height = 1600, 900
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(30, bold=True)
    draw.text((60, 40), caption, font=title_font, fill="#0E2A47")

    # Default layout
    if "技术体系架构图" in caption:
        labels = [
            ("客户端\n桌面/Web/CLI", (80, 150, 360, 280)),
            ("Gateway", (470, 150, 690, 280)),
            ("Orchestrator\n多智能体", (800, 150, 1080, 280)),
            ("Brain", (1190, 150, 1410, 280)),
            ("Hands", (470, 420, 690, 550)),
            ("Memory / Knowledge", (800, 420, 1080, 550)),
            ("Heartbeat / Skills", (1190, 420, 1490, 550)),
            ("外部系统\n设计/仿真/制造/数据", (800, 680, 1120, 810)),
        ]
        for text, box in labels:
            draw_box(draw, box, text)
        for s, e in [
            ((360, 215), (470, 215)),
            ((690, 215), (800, 215)),
            ((1080, 215), (1190, 215)),
            ((580, 280), (580, 420)),
            ((940, 280), (940, 420)),
            ((1340, 280), (1340, 420)),
            ((940, 550), (940, 680)),
        ]:
            draw_arrow(draw, s, e)
    elif "Gateway 与客户端接入架构图" in caption:
        for text, box in [
            ("桌面客户端", (100, 220, 340, 350)),
            ("Web 客户端", (100, 420, 340, 550)),
            ("CLI", (100, 620, 340, 750)),
            ("Gateway", (620, 360, 920, 520)),
            ("Orchestrator /\n系统路由", (1180, 360, 1480, 520)),
        ]:
            draw_box(draw, box, text)
        for s, e in [((340, 285), (620, 420)), ((340, 485), (620, 440)), ((340, 685), (620, 460)), ((920, 440), (1180, 440))]:
            draw_arrow(draw, s, e)
    elif "工具循环示意图" in caption:
        for text, box in [
            ("System Prompt /\nTool Definitions", (80, 350, 380, 500)),
            ("Brain / LLM", (650, 160, 950, 310)),
            ("Tool Call", (650, 380, 950, 530)),
            ("Hands 执行结果", (650, 600, 950, 750)),
            ("结果回注对话", (1220, 350, 1520, 500)),
        ]:
            draw_box(draw, box, text)
        for s, e in [((380, 425), (650, 235)), ((800, 310), (800, 380)), ((800, 530), (800, 600)), ((950, 425), (1220, 425)), ((1220, 455), (950, 235))]:
            draw_arrow(draw, s, e)
    elif "多智能体路由与流水线结构图" in caption:
        for text, box in [
            ("用户请求", (100, 360, 320, 500)),
            ("Route", (460, 360, 660, 500)),
            ("Agent Registry", (800, 160, 1100, 300)),
            ("专业 Agent", (800, 360, 1100, 500)),
            ("Pipeline", (800, 560, 1100, 700)),
            ("Hands", (1260, 360, 1460, 500)),
        ]:
            draw_box(draw, box, text)
        for s, e in [((320, 430), (460, 430)), ((660, 430), (800, 230)), ((660, 430), (800, 430)), ((660, 430), (800, 630)), ((1100, 430), (1260, 430))]:
            draw_arrow(draw, s, e)
    elif "工具注册与外部系统映射图" in caption:
        for text, box in [
            ("Tools Registry", (120, 360, 380, 520)),
            ("设计适配器", (560, 150, 860, 270)),
            ("仿真适配器", (560, 310, 860, 430)),
            ("MES/测控适配器", (560, 470, 860, 590)),
            ("数据/情报适配器", (560, 630, 860, 750)),
            ("外部系统", (1120, 360, 1420, 520)),
        ]:
            draw_box(draw, box, text)
        for y in [210, 370, 530, 690]:
            draw_arrow(draw, (380, 440), (560, y))
            draw_arrow(draw, (860, y), (1120, 440))
    elif "记忆—知识—情报数据通路图" in caption:
        for text, box in [
            ("对话与项目", (80, 360, 320, 500)),
            ("Memory", (460, 190, 700, 330)),
            ("Knowledge / Library", (460, 390, 700, 530)),
            ("arXiv Intel", (460, 590, 700, 730)),
            ("向量检索 / 报表", (980, 290, 1280, 430)),
            ("推送通道", (980, 530, 1280, 670)),
        ]:
            draw_box(draw, box, text)
        for s, e in [((320, 430), (460, 260)), ((320, 430), (460, 460)), ((320, 430), (460, 660)), ((700, 260), (980, 360)), ((700, 460), (980, 360)), ((700, 660), (980, 600))]:
            draw_arrow(draw, s, e)
    elif "Heartbeat + Skills + Sidecar 部署与依赖关系图" in caption:
        for text, box in [
            ("Heartbeat", (220, 220, 500, 360)),
            ("Skills Loader", (220, 500, 500, 640)),
            ("Gateway", (690, 360, 950, 500)),
            ("Sidecar APIs", (1150, 220, 1430, 360)),
            ("领域服务", (1150, 500, 1430, 640)),
        ]:
            draw_box(draw, box, text)
        for s, e in [((500, 290), (690, 430)), ((500, 570), (690, 430)), ((950, 430), (1150, 290)), ((950, 430), (1150, 570))]:
            draw_arrow(draw, s, e)
    else:
        draw_box(draw, (320, 260, 1280, 640), caption.replace("图 ", ""))

    img.save(output_path)


def ensure_figure_image(figures_dir: Path, caption: str) -> Path:
    figures_dir.mkdir(parents=True, exist_ok=True)
    match = re.search(r"图\s*(\d+)", caption)
    num = match.group(1) if match else "x"
    out = figures_dir / f"figure_{num}.png"
    # 优先使用已由 generate_figures.py 生成的高质量 matplotlib 图
    if not out.exists():
        build_figure_image(caption, out)
    return out


def add_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part.replace("**", "").replace("`", ""))
        if part.startswith("**") and part.endswith("**"):
            run.bold = True
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def get_metadata(lines: list[str]) -> tuple[str, str, str, str]:
    center_title = "量子计算工业软件工程技术研发中心"
    module_title = "第一模块：量智大脑（QuantaMind）"
    version = "V1.0"
    product_name = "QuantaMind（量智大脑）"
    for line in lines[:20]:
        stripped = line.strip()
        if stripped.startswith("# "):
            center_title = stripped[2:].strip()
        elif stripped.startswith("## 第一模块："):
            module_title = stripped[3:].strip()
        elif stripped.startswith("**版本：**"):
            version = stripped.replace("**版本：**", "").strip()
        elif stripped.startswith("**对应产品代号：**"):
            product_name = stripped.replace("**对应产品代号：**", "").strip()
    return center_title, module_title, version, product_name


def find_body_start(lines: list[str]) -> int:
    for i, line in enumerate(lines):
        if line.strip().startswith("## 中心一句话总述"):
            return i
    return 0


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        rows.append(row)
        i += 1
    if len(rows) >= 2 and all(set(cell) <= {"-"} for cell in rows[1]):
        rows.pop(1)
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = row[c_idx] if c_idx < len(row) else ""
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing = 1.25
                for run in p.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(10.5)
    if rows:
        for run in table.rows[0].cells[0].paragraphs[0].runs:
            run.bold = True


def add_body_title(doc: Document, module_title: str) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    add_runs(p, module_title)


def is_major_section_title(text: str) -> bool:
    return bool(re.match(r"^##\s*[一二三四五六七八九十]+、", text.strip()))


def export(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_default_style(doc)
    set_page_layout(doc)
    center_title, module_title, version, product_name = get_metadata(lines)
    set_header_footer(doc, center_title)
    add_cover(doc, center_title, module_title, version, product_name)
    add_toc(doc)
    add_body_title(doc, module_title)
    figures_dir = docx_path.parent / "_generated_figures"
    seen_major_section = False

    i = find_body_start(lines)
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped == "## 模块正文":
            i += 1
            continue

        fig_match = re.match(r"^\*\*(图\s*\d+.+?)\*\*$", stripped)
        if fig_match:
            caption = fig_match.group(1)
            image_path = ensure_figure_image(figures_dir, caption)
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.first_line_indent = Cm(0)
            p_img.add_run().add_picture(str(image_path), width=Inches(6.3))

            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.first_line_indent = Cm(0)
            add_runs(p_cap, caption)
            i += 1
            if i < len(lines) and lines[i].strip().startswith("（"):
                i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_runs(p, stripped[2:].strip())
            i += 1
            continue
        if stripped.startswith("## "):
            if is_major_section_title(stripped):
                if seen_major_section:
                    doc.add_page_break()
                seen_major_section = True
            p = doc.add_paragraph(style="Heading 1")
            add_runs(p, stripped[3:].strip())
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_runs(p, stripped[4:].strip())
            i += 1
            continue
        if stripped.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_runs(p, stripped[5:].strip())
            i += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, bullet_match.group(1))
            i += 1
            continue

        num_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if num_match:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, num_match.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, stripped.replace("  ", " "))
        i += 1

    doc.save(docx_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: python export_word_from_markdown.py <input.md> <output.docx>")
    export(Path(sys.argv[1]), Path(sys.argv[2]))
